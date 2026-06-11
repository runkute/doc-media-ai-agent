# -*- coding: utf-8 -*-
import os
import json
import sqlite3
import io
import asyncio
import uuid
import re
from datetime import datetime
import urllib.request
import urllib.error
import urllib.parse
from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import StreamingResponse, HTMLResponse, JSONResponse
from contextlib import asynccontextmanager
from pydantic import BaseModel
from typing import AsyncGenerator
import itertools
import anthropic
from apscheduler.schedulers.asyncio import AsyncIOScheduler
from google import genai as google_genai
from google.genai import types as genai_types
from dotenv import load_dotenv

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
load_dotenv(os.path.join(BASE_DIR, ".env"))

scheduler = AsyncIOScheduler(timezone="Asia/Ho_Chi_Minh")

@asynccontextmanager
async def lifespan(app):
    scheduler.start()
    yield
    scheduler.shutdown()

app = FastAPI(lifespan=lifespan)
# ── Anthropic multi-key pool (round-robin) ──────────────────────────
def _load_anthropic_pool():
    keys = []
    for k in sorted(os.environ.keys()):
        if (k == "ANTHROPIC_API_KEY" or k.startswith("ANTHROPIC_API_KEY_")) and os.environ[k].strip():
            keys.append(os.environ[k].strip())
    return [anthropic.Anthropic(api_key=k) for k in keys]

_anthropic_pool = _load_anthropic_pool()
_anthropic_cycle = itertools.cycle(range(len(_anthropic_pool))) if _anthropic_pool else None
client = _anthropic_pool[0] if _anthropic_pool else anthropic.Anthropic()  # backward-compat

def get_anthropic_client() -> anthropic.Anthropic:
    """Return next Claude client in round-robin rotation."""
    if not _anthropic_pool:
        return client
    return _anthropic_pool[next(_anthropic_cycle)]

# ── Gemini multi-key pool (round-robin) ────────────────────────────
def _load_gemini_pool():
    keys = []
    env = os.environ
    for k in sorted(env.keys()):
        if (k == "GEMINI_API_KEY" or k.startswith("GEMINI_API_KEY_")) and env[k].strip():
            keys.append(env[k].strip())
    return [google_genai.Client(api_key=k) for k in keys]

_gemini_pool = _load_gemini_pool()
_gemini_cycle = itertools.cycle(range(len(_gemini_pool))) if _gemini_pool else None
gemini_client = _gemini_pool[0] if _gemini_pool else None  # backward-compat for None checks

def get_gemini_client() -> google_genai.Client | None:
    """Return next Gemini client in round-robin rotation."""
    if not _gemini_pool:
        return None
    return _gemini_pool[next(_gemini_cycle)]
DB_PATH = os.path.join(BASE_DIR, "chat.db")

# ── Database ───────────────────────────────────────────────────────

def init_db():
    with sqlite3.connect(DB_PATH) as c:
        c.execute("""CREATE TABLE IF NOT EXISTS messages (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            session_id TEXT NOT NULL, role TEXT NOT NULL,
            content TEXT NOT NULL,
            ts TEXT DEFAULT (datetime('now','localtime')))""")
        c.execute("""CREATE TABLE IF NOT EXISTS alerts (
            id TEXT PRIMARY KEY,
            type TEXT NOT NULL,
            title TEXT NOT NULL,
            body TEXT NOT NULL,
            read INTEGER DEFAULT 0,
            ts TEXT DEFAULT (datetime('now','localtime')))""")
        c.execute("""CREATE TABLE IF NOT EXISTS content_library (
            id TEXT PRIMARY KEY, type TEXT NOT NULL,
            title TEXT NOT NULL, content TEXT NOT NULL,
            client TEXT DEFAULT '',
            created_at TEXT DEFAULT (datetime('now','localtime')))""")

def db_add_alert(alert_type: str, title: str, body: str):
    with sqlite3.connect(DB_PATH) as c:
        c.execute("INSERT INTO alerts(id,type,title,body) VALUES(?,?,?,?)",
                  (str(uuid.uuid4()), alert_type, title, body))

def db_get_alerts():
    with sqlite3.connect(DB_PATH) as c:
        rows = c.execute(
            "SELECT id,type,title,body,read,ts FROM alerts ORDER BY ts DESC LIMIT 30"
        ).fetchall()
    return [{"id":r[0],"type":r[1],"title":r[2],"body":r[3],"read":bool(r[4]),"ts":r[5]} for r in rows]

def db_mark_alert_read(alert_id: str):
    with sqlite3.connect(DB_PATH) as c:
        c.execute("UPDATE alerts SET read=1 WHERE id=?", (alert_id,))

def db_unread_count():
    with sqlite3.connect(DB_PATH) as c:
        return c.execute("SELECT COUNT(*) FROM alerts WHERE read=0").fetchone()[0]

# ── Content Library DB ─────────────────────────────────────────────
_LIB_TYPE_LABELS = {'content':'Bài viết','script':'Kịch bản','brief':'Brief','quote':'Báo giá'}

def db_lib_save(lib_type: str, title: str, content: str, client: str = "") -> str:
    item_id = str(uuid.uuid4())
    with sqlite3.connect(DB_PATH) as c:
        c.execute("INSERT INTO content_library(id,type,title,content,client) VALUES(?,?,?,?,?)",
                  (item_id, lib_type, title, content, client))
    return item_id

def db_lib_list(type_filter: str = "", search: str = "") -> list:
    with sqlite3.connect(DB_PATH) as c:
        q = "SELECT id,type,title,content,client,created_at FROM content_library"
        params: list = []
        conds: list = []
        if type_filter:
            conds.append("type=?"); params.append(type_filter)
        if search:
            conds.append("(title LIKE ? OR content LIKE ? OR client LIKE ?)")
            params += [f"%{search}%", f"%{search}%", f"%{search}%"]
        if conds:
            q += " WHERE " + " AND ".join(conds)
        q += " ORDER BY created_at DESC LIMIT 100"
        rows = c.execute(q, params).fetchall()
    return [{"id":r[0],"type":r[1],"type_label":_LIB_TYPE_LABELS.get(r[1],r[1]),
             "title":r[2],"content":r[3],"client":r[4],"created_at":r[5]} for r in rows]

def db_lib_delete(item_id: str):
    with sqlite3.connect(DB_PATH) as c:
        c.execute("DELETE FROM content_library WHERE id=?", (item_id,))

# ── Export helpers ─────────────────────────────────────────────────
def _md_to_plain(text: str) -> str:
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
    text = re.sub(r'`(.+?)`', r'\1', text)
    return text.strip()

def export_library_word(items: list) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor
    doc = Document()
    h0 = doc.add_heading('Thư viện Nội dung — Độc Media', 0)
    h0.runs[0].font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)
    meta_p = doc.add_paragraph(f'Xuất ngày: {datetime.now().strftime("%d/%m/%Y %H:%M")} | Tổng: {len(items)} mục')
    meta_p.runs[0].font.size = Pt(10)
    doc.add_paragraph('─' * 60)
    for item in items:
        h1 = doc.add_heading(item['title'], 1)
        info = f"Loại: {item['type_label']}"
        if item['client']: info += f" | Client: {item['client']}"
        info += f" | Ngày: {item['created_at']}"
        ip = doc.add_paragraph(info)
        ip.runs[0].font.size = Pt(9)
        ip.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x99)
        doc.add_paragraph(_md_to_plain(item['content']))
        doc.add_paragraph()
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

def export_library_excel(items: list) -> bytes:
    import openpyxl
    from openpyxl.styles import Font as XFont, PatternFill, Alignment
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Thu vien'
    headers = ['#', 'Loại', 'Tiêu đề', 'Client', 'Nội dung', 'Ngày tạo']
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = XFont(bold=True, color='FFFFFF', name='Arial')
        cell.fill = PatternFill('solid', fgColor='7C3AED')
        cell.alignment = Alignment(horizontal='center', vertical='center')
    ws.row_dimensions[1].height = 28
    for i, item in enumerate(items, 2):
        ws.cell(row=i, column=1, value=i-1)
        ws.cell(row=i, column=2, value=item['type_label'])
        ws.cell(row=i, column=3, value=item['title'])
        ws.cell(row=i, column=4, value=item['client'] or '')
        c5 = ws.cell(row=i, column=5, value=_md_to_plain(item['content']))
        c5.alignment = Alignment(wrap_text=True, vertical='top')
        ws.cell(row=i, column=6, value=item['created_at'])
        ws.row_dimensions[i].height = 80
    ws.column_dimensions['A'].width = 5
    ws.column_dimensions['B'].width = 13
    ws.column_dimensions['C'].width = 35
    ws.column_dimensions['D'].width = 20
    ws.column_dimensions['E'].width = 85
    ws.column_dimensions['F'].width = 18
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.getvalue()

def export_library_pdf(items: list) -> bytes:
    from fpdf import FPDF
    FONT_R = 'C:/Windows/Fonts/arial.ttf'
    FONT_B = 'C:/Windows/Fonts/arialbd.ttf'
    pdf = FPDF()
    try:
        pdf.add_font('F', '', FONT_R)
        pdf.add_font('F', 'B', FONT_B)
        font_name = 'F'
    except Exception:
        font_name = 'Helvetica'
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font(font_name, 'B', 18)
    pdf.set_text_color(124, 58, 237)
    pdf.cell(0, 12, 'Thu vien Noi dung - Doc Media', ln=True, align='C')
    pdf.set_font(font_name, '', 10)
    pdf.set_text_color(120, 120, 140)
    pdf.cell(0, 8, f'Xuat ngay: {datetime.now().strftime("%d/%m/%Y %H:%M")} | Tong: {len(items)} muc', ln=True, align='C')
    pdf.ln(6)
    for item in items:
        pdf.set_font(font_name, 'B', 13)
        pdf.set_text_color(40, 40, 60)
        pdf.multi_cell(0, 8, item['title'])
        pdf.set_font(font_name, '', 9)
        pdf.set_text_color(120, 120, 140)
        info = f"Loai: {item['type_label']}"
        if item['client']: info += f" | Client: {item['client']}"
        info += f" | Ngay: {item['created_at']}"
        pdf.cell(0, 6, info, ln=True)
        pdf.ln(2)
        pdf.set_font(font_name, '', 10)
        pdf.set_text_color(50, 50, 70)
        pdf.multi_cell(0, 6, _md_to_plain(item['content'])[:4000])
        pdf.ln(3)
        pdf.set_draw_color(180, 180, 200)
        pdf.line(15, pdf.get_y(), 195, pdf.get_y())
        pdf.ln(5)
    return bytes(pdf.output())

init_db()

def db_get_history(sid, limit=20):
    with sqlite3.connect(DB_PATH) as c:
        rows = c.execute(
            "SELECT role,content FROM messages WHERE session_id=? ORDER BY id DESC LIMIT ?",
            (sid, limit)).fetchall()
    return [{"role": r[0], "content": r[1]} for r in reversed(rows)]

def db_save_turn(sid, user_content, assistant_content):
    with sqlite3.connect(DB_PATH) as c:
        c.execute("INSERT INTO messages(session_id,role,content) VALUES(?,?,?)", (sid,"user",user_content))
        c.execute("INSERT INTO messages(session_id,role,content) VALUES(?,?,?)", (sid,"assistant",assistant_content))

def db_get_sessions():
    with sqlite3.connect(DB_PATH) as c:
        rows = c.execute("""
            SELECT session_id,
                   (SELECT content FROM messages WHERE session_id=m.session_id AND role='user' ORDER BY id LIMIT 1),
                   MAX(ts)
            FROM messages m GROUP BY session_id ORDER BY MAX(id) DESC LIMIT 15""").fetchall()
    return [{"id": r[0], "preview": (r[1] or "Chat moi")[:45], "time": (r[2] or "")[:16]} for r in rows]

def db_clear_session(sid):
    with sqlite3.connect(DB_PATH) as c:
        c.execute("DELETE FROM messages WHERE session_id=?", (sid,))

# ── Chat Agent ─────────────────────────────────────────────────────

CHAT_SYSTEM = (
    "Ban la tro ly marketing chuyen nganh kien truc/noi that Viet Nam, ho tro agency Doc Media.\n"
    "Chuyen ve: Facebook Ads, TikTok Ads, phan tich doi thu, benchmark CPM/CPC, content/hook, toi uu ngan sach.\n"
    "BENCHMARK THUC TE nganh Thiet ke & Thi cong Noi that VN (data RECO 2026):\n"
    "- CPM muc tieu Tin nhan (Inbox): 150.000 – 230.000 VND | CPM Tuong tac: 21.000 – 34.000 VND\n"
    "- CPL (chi phi/inbox): Tot <230K | Trung binh 230-315K | Kem >350K | Best case: 166.803 VND\n"
    "- CTR Tin nhan: 1.26-2.07% | CTR Tuong tac: 2.87-4.52% | CPE: 500-1.100 VND\n"
    "- Leads voi 3 trieu/10 ngay: 10-18 inbox (CPL avg 250-300K)\n"
    "Khi can so lieu moi nhat: su dung web search. Tra loi thuc te, co so lieu cu the, viet tieng Viet."
)
WEB_SEARCH_TOOL = [{"type": "web_search_20260209", "name": "web_search"}]

AGENT_TOOLS = [
    {
        "name": "generate_content",
        "description": "Viết bài quảng cáo Facebook theo công thức copywriting (PAS/AIDA/BAB/4U) cho sản phẩm nội thất/kiến trúc. Dùng khi người dùng muốn tạo bài viết, content, copy quảng cáo.",
        "input_schema": {
            "type": "object",
            "properties": {
                "client": {"type": "string", "description": "Tên khách hàng/doanh nghiệp"},
                "product": {"type": "string", "description": "Sản phẩm hoặc dịch vụ cần quảng cáo"},
                "objective": {"type": "string", "description": "Mục tiêu: Tin nhắn, Tương tác, Brand Awareness"},
                "target": {"type": "string", "description": "Đối tượng khách hàng mục tiêu"},
                "key_message": {"type": "string", "description": "Thông điệp chính muốn truyền tải"},
                "formula": {"type": "string", "enum": ["PAS", "AIDA", "BAB", "4U"], "description": "Công thức copywriting. Mặc định: PAS"},
                "tone": {"type": "string", "description": "Giọng văn. Mặc định: Chuyên nghiệp, uy tín, sang trọng"}
            },
            "required": ["product", "target"]
        }
    },
    {
        "name": "generate_brief",
        "description": "Phân tích và tạo brief chiến lược marketing cho client mới theo chuẩn CMO. Dùng khi người dùng muốn onboard client mới, tạo plan, brief, hoặc phân tích chiến lược.",
        "input_schema": {
            "type": "object",
            "properties": {
                "ten_dn": {"type": "string", "description": "Tên doanh nghiệp/client"},
                "linh_vuc": {"type": "string", "description": "Lĩnh vực kinh doanh"},
                "mo_ta_sp": {"type": "string", "description": "Mô tả sản phẩm/dịch vụ"},
                "doi_tuong": {"type": "string", "description": "Đối tượng khách hàng mục tiêu"},
                "usp": {"type": "string", "description": "Điểm khác biệt, lợi thế cạnh tranh"},
                "ngan_sach_ads": {"type": "string", "description": "Ngân sách quảng cáo/tháng (VND)"},
                "muc_tieu_chinh": {"type": "string", "description": "Mục tiêu chính của campaign"},
                "thoi_gian_hd": {"type": "string", "description": "Thời gian hợp đồng, VD: 3 thang"}
            },
            "required": ["ten_dn"]
        }
    },
    {
        "name": "generate_quote",
        "description": "Tạo báo giá dịch vụ Facebook Ads chi tiết với KPI cam kết. Dùng khi người dùng muốn tạo báo giá, proposal cho client.",
        "input_schema": {
            "type": "object",
            "properties": {
                "ten_client": {"type": "string", "description": "Tên khách hàng"},
                "ngan_sach_ads": {"type": "string", "description": "Ngân sách ads/tháng (VND, chỉ số)"},
                "thoi_gian": {"type": "string", "description": "Thời gian hợp đồng, VD: 3 tháng"},
                "linh_vuc": {"type": "string", "description": "Lĩnh vực"},
                "muc_tieu": {"type": "string", "description": "Mục tiêu chiến dịch"},
                "ghi_chu": {"type": "string", "description": "Ghi chú thêm"}
            },
            "required": ["ten_client"]
        }
    },
    {
        "name": "get_ads_summary",
        "description": "Lấy tóm tắt hiệu suất quảng cáo Facebook Ads thực tế: chi tiêu, reach, CPM, tin nhắn, leads. Dùng khi người dùng hỏi về kết quả ads, hiệu suất, số liệu chiến dịch.",
        "input_schema": {
            "type": "object",
            "properties": {
                "date_preset": {
                    "type": "string",
                    "enum": ["today", "yesterday", "last_7d", "last_14d", "last_30d", "this_month"],
                    "description": "Khoảng thời gian. Mặc định: last_7d"
                }
            },
            "required": []
        }
    }
]

ALL_TOOLS = WEB_SEARCH_TOOL + AGENT_TOOLS

TOOL_LABELS = {
    "generate_content": "Đang viết bài quảng cáo...",
    "generate_brief": "Đang tạo brief chiến lược...",
    "generate_quote": "Đang tạo báo giá...",
    "get_ads_summary": "Đang lấy dữ liệu Facebook Ads...",
    "web_search": "Đang tìm kiếm thông tin...",
}

async def execute_tool(name: str, inputs: dict) -> str:
    if name == "generate_content":
        brief = {
            "client": inputs.get("client", ""),
            "product": inputs.get("product", ""),
            "objective": inputs.get("objective", "Tin nhắn"),
            "target": inputs.get("target", ""),
            "key_message": inputs.get("key_message", ""),
            "tone": inputs.get("tone", "Chuyên nghiệp, uy tín, sang trọng"),
            "formula": inputs.get("formula", "PAS"),
            "budget": "",
            "ai_model": "claude-sonnet-4-6",
        }
        text = ""
        async for chunk in content_stream(brief):
            if chunk.startswith("data: ") and "[DONE]" not in chunk:
                try: text += json.loads(chunk[6:]).get("text", "")
                except: pass
        return text or "Không tạo được content."

    if name == "generate_brief":
        body = {
            "ten_dn": inputs.get("ten_dn", ""),
            "linh_vuc": inputs.get("linh_vuc", "Nội thất"),
            "mo_ta_sp": inputs.get("mo_ta_sp", ""),
            "doi_tuong": inputs.get("doi_tuong", ""),
            "usp": inputs.get("usp", ""),
            "ngan_sach_ads": inputs.get("ngan_sach_ads", ""),
            "muc_tieu_chinh": inputs.get("muc_tieu_chinh", "Thu leads"),
            "thoi_gian_hd": inputs.get("thoi_gian_hd", "3 thang"),
            "fanpage_url": "", "nguoi_lien_he": "",
            "fanpage_status": "Chua co fanpage", "followers": "0",
            "da_chay_ads": "Chua chay bao gio", "cpl_cu": "",
            "leads_mong_muon": "", "doanh_thu_mong_muon": "",
            "timeline_bat_dau": "", "assets": "", "ngan_sach_dv": "",
            "doi_thu": "", "ai_model": "claude-sonnet-4-6",
        }
        text = ""
        async for chunk in intake_stream(body):
            if chunk.startswith("data: ") and "[DONE]" not in chunk:
                try: text += json.loads(chunk[6:]).get("text", "")
                except: pass
        return text or "Không tạo được brief."

    if name == "generate_quote":
        body = {
            "ten_client": inputs.get("ten_client", ""),
            "linh_vuc": inputs.get("linh_vuc", "Nội thất / Thiết kế & Thi công"),
            "muc_tieu": inputs.get("muc_tieu", "Tin nhắn (Lead) là chính"),
            "ngan_sach_ads": inputs.get("ngan_sach_ads", ""),
            "thoi_gian": inputs.get("thoi_gian", "3 tháng"),
            "dich_vu": "", "ghi_chu": inputs.get("ghi_chu", ""),
        }
        text = ""
        async for chunk in quote_stream(body):
            if chunk.startswith("data: ") and "[DONE]" not in chunk:
                try: text += json.loads(chunk[6:]).get("text", "")
                except: pass
        return text or "Không tạo được báo giá."

    if name == "get_ads_summary":
        date_preset = inputs.get("date_preset", "last_7d")
        try:
            account_id = os.getenv("FB_AD_ACCOUNT_ID", "")
            data = await asyncio.to_thread(_get_fb_data, account_id, date_preset)
            rows = data.get("insights", {}).get("data", [])
            if not rows:
                return "Không có dữ liệu quảng cáo trong khoảng thời gian này."
            total_spend = sum(float(r.get("spend", 0)) for r in rows)
            total_reach = sum(int(r.get("reach", 0)) for r in rows)
            total_imp = sum(int(r.get("impressions", 0)) for r in rows)
            total_msgs = total_leads = 0
            for r in rows:
                for a in r.get("actions", []):
                    if a["action_type"] == "onsite_conversion.messaging_conversation_started_7d":
                        total_msgs += int(a.get("value", 0))
                    elif a["action_type"] == "lead":
                        total_leads += int(a.get("value", 0))
            cpm = total_spend / total_imp * 1000 if total_imp else 0
            cp_msg = total_spend / total_msgs if total_msgs else 0
            cp_lead = total_spend / total_leads if total_leads else 0
            return (
                f"KẾT QUẢ FB ADS ({date_preset}):\n"
                f"- Chi tiêu: {total_spend:,.0f} VND\n"
                f"- Reach: {total_reach:,} | Impressions: {total_imp:,}\n"
                f"- CPM: {cpm:,.0f} VND\n"
                f"- Tin nhắn: {total_msgs} | CP/Tin nhắn: {cp_msg:,.0f} VND"
                f"{' ✅' if cp_msg < 230000 else ' ⚠️' if cp_msg > 350000 else ''}\n"
                f"- KHTN (leads): {total_leads} | CP/KHTN: {cp_lead:,.0f} VND\n"
                f"- Số chiến dịch: {len(rows)}"
            )
        except Exception as e:
            return f"Lỗi lấy dữ liệu ads: {str(e)[:200]}"

    return f"Công cụ '{name}' chưa được hỗ trợ."

async def agent_stream(messages, session_id, user_msg):
    all_text = ""
    loop_msgs = messages[:]
    while True:
        try:
            with get_anthropic_client().messages.stream(
                model="claude-sonnet-4-6", max_tokens=4096,
                system=CHAT_SYSTEM, messages=loop_msgs, tools=ALL_TOOLS,
            ) as stream:
                for text in stream.text_stream:
                    all_text += text
                    yield f"data: {json.dumps({'text': text})}\n\n"
                final = stream.get_final_message()
        except Exception as e:
            yield f"data: {json.dumps({'text': f'[Lỗi: {str(e)[:150]}]'})}\n\n"
            break

        if final.stop_reason == "pause_turn":
            # web_search đang chạy
            loop_msgs.append({"role": "assistant", "content": final.content})
            continue

        if final.stop_reason == "tool_use":
            tool_results = []
            for block in final.content:
                if block.type == "tool_use":
                    label = TOOL_LABELS.get(block.name, block.name)
                    yield f"data: {json.dumps({'tool': block.name, 'label': label})}\n\n"
                    try:
                        result = await execute_tool(block.name, block.input)
                    except Exception as te:
                        result = f"Lỗi: {str(te)[:200]}"
                    tool_results.append({
                        "type": "tool_result",
                        "tool_use_id": block.id,
                        "content": result,
                    })
            loop_msgs.append({"role": "assistant", "content": final.content})
            loop_msgs.append({"role": "user", "content": tool_results})
            continue

        break  # end_turn

    if all_text:
        db_save_turn(session_id, user_msg, all_text)
    yield "data: [DONE]\n\n"

# ── File Parsing ───────────────────────────────────────────────────

async def parse_upload(file):
    try:
        import pandas as pd
    except ImportError:
        raise HTTPException(500, "pip install pandas openpyxl")
    content = await file.read()
    fname = (file.filename or "").lower()
    try:
        df = pd.read_csv(io.BytesIO(content)) if fname.endswith(".csv") else pd.read_excel(io.BytesIO(content))
    except Exception as e:
        raise HTTPException(400, f"Loi doc file: {str(e)[:100]}")
    rows, cols = len(df), len(df.columns)
    summary = f"File: {file.filename} | {rows} dong x {cols} cot\nCac cot: {', '.join(str(c) for c in df.columns)}\n\n"
    summary += df.head(50).to_string(index=False)
    return {"summary": summary, "filename": file.filename, "rows": rows, "cols": cols}

# ── Image Generation ───────────────────────────────────────────────

def _enhance_prompt(description, style):
    resp = get_anthropic_client().messages.create(
        model="claude-sonnet-4-6", max_tokens=150,
        system="Convert Vietnamese interior design descriptions to English Flux AI image generation prompts.",
        messages=[{"role": "user", "content":
            f"Style: {style}. Description: {description}. "
            "Output ONLY English prompt, max 100 words. "
            "Include: room type, style, materials, lighting, photorealistic architectural visualization."}])
    return resp.content[0].text.strip()

def _call_fal(prompt):
    key = os.getenv("FAL_KEY", "")
    if not key or key.startswith("your_"):
        raise ValueError("FAL_KEY chua duoc cai dat trong .env")
    payload = json.dumps({"prompt": prompt, "image_size": "landscape_16_9",
        "num_inference_steps": 28, "guidance_scale": 3.5,
        "num_images": 1, "enable_safety_checker": True}).encode()
    req = urllib.request.Request("https://fal.run/fal-ai/flux/dev", data=payload,
        method="POST", headers={"Authorization": f"Key {key}", "Content-Type": "application/json"})
    try:
        with urllib.request.urlopen(req, timeout=120) as r:
            return json.loads(r.read().decode())
    except urllib.error.HTTPError as e:
        raise ValueError(f"fal.ai loi {e.code}: {e.read().decode()[:200]}")

# ── Content Studio ─────────────────────────────────────────────────

CONTENT_SYSTEM = """Bạn là Copywriter thực chiến và Chuyên gia Performance Marketing Facebook Ads với 14 năm kinh nghiệm, làm việc cho agency Độc Media. Chuyên môn: viết quảng cáo hard-sell, tối ưu tỷ lệ chuyển đổi, thu lead chất lượng cho dịch vụ thiết kế & thi công nhà phố cao cấp.

NGUYÊN TẮC BẮT BUỘC:
- Định dạng Markdown, phân chia Heading rõ ràng, Bullet point, in đậm Key Message.
- Ưu tiên tính chính xác kỹ thuật, công năng thực tế. KHÔNG dùng ngôn từ bay bổng, mông lung, phi thực tế.
- Văn phong: mạnh mẽ, dứt khoát, chuyên nghiệp, đi thẳng vào vấn đề, không vòng vo.
- Luôn có CTA rõ ràng, mang tính thúc giục (Urgency), chỉ 1 hành động duy nhất.
- TUYỆT ĐỐI CẤM dùng cụm từ AI sáo rỗng: "Trong thế giới hiện đại ngày nay", "Bạn đang tìm kiếm...", "Hãy tưởng tượng...", "Hơn bao giờ hết", "Giải pháp hoàn hảo".
- Dùng ngôn ngữ chuyên ngành tự nhiên: bóc tách dự toán, chốt phương án mặt bằng, đổ mê, lún nứt, bản vẽ kỹ thuật thi công, suất đầu tư.
- Câu ngắn, dứt khoát, ngắt đoạn tạo nhịp điệu nhanh. Không giải thích dông dài.
- Kết hợp thấu cảm (hiểu khách sợ phát sinh chi phí, sợ nhà thầu bỏ con giữa chợ) và uy quyền chuyên gia.
- Không dùng mồi nhử giá rẻ mạt phi thực tế. Luôn định hướng bằng suất đầu tư tiêu chuẩn thực tế.

CẤU TRÚC OUTPUT BẮT BUỘC cho mỗi bài viết:
1. **Target Audience:** (Xác định nhanh tệp khách hàng nhắm tới)
2. **Angle/Góc tiếp cận:** (Lý do chọn góc viết này)
3. **Nội dung bài viết:** (Áp dụng đúng công thức được chỉ định, tối ưu 3 dòng đầu hiển thị trên Facebook)
4. **Gợi ý Visual:** (Đề xuất hình ảnh/video: bản vẽ kỹ thuật, mặt bằng công năng, video tiến độ thi công thực tế)"""

FORMULA_GUIDES = {
    "PAS": """CÔNG THỨC PAS — Tối ưu đánh vào Nỗi đau/Rủi ro:
- **P (Problem):** Chỉ ra rủi ro/nỗi đau lớn nhất khi xây nhà phố (phát sinh chi phí, lún nứt, sai bản vẽ, thiết kế bí/thiếu sáng)
- **A (Agitate):** Xát muối vào nỗi đau — hậu quả tài chính & tinh thần nghiêm trọng nếu không giải quyết triệt để từ đầu
- **S (Solve):** Giải pháp là dịch vụ của client với quy trình chuẩn kỹ thuật, cam kết rõ ràng bằng hợp đồng""",

    "AIDA": """CÔNG THỨC AIDA — Tối ưu ra mắt mẫu thiết kế/Chương trình ưu đãi:
- **A (Attention):** Tiêu đề giật tít bằng con số, USP độc bản hoặc ưu đãi giới hạn thời gian
- **I (Interest):** Thông số kỹ thuật ấn tượng, giải pháp tối ưu không gian/ánh sáng/vật liệu thi công
- **D (Desire):** Cam kết chất lượng, bảo hành kết cấu, hoặc quà tặng kèm để tạo khao khát sở hữu
- **A (Action):** CTA khẩn cấp để nhận ưu đãi hoặc tư vấn miễn phí ngay""",

    "BAB": """CÔNG THỨC BAB — Tối ưu showcase dự án/portfolio:
- **B (Before):** Tình trạng hiện tại của khách hàng — căn nhà chật, tối, bất tiện, lãng phí diện tích
- **A (After):** Kết quả sau khi dùng dịch vụ — không gian thoáng, sáng, công năng tối ưu từng m²
- **B (Bridge):** Dịch vụ của client là cây cầu — quy trình chuẩn, cam kết timeline, bảo hành kết cấu""",

    "4U": """CÔNG THỨC 4U — Tối ưu quảng cáo promotion/offer giới hạn:
- **Urgent:** Tạo tính khẩn cấp — thời gian ưu đãi có hạn, số lượng dự án nhận trong tháng giới hạn
- **Unique:** Điểm khác biệt độc nhất so với đối thủ — không ai làm được điều này ngoài client
- **Ultra-specific:** Số liệu cụ thể — diện tích, suất đầu tư/m², timeline thi công, số công trình đã bàn giao
- **Useful:** Giá trị thực tế — khách hàng nhận được gì cụ thể (bản vẽ, dự toán, bảo hành bao lâu)""",
}

def _build_content_prompt(brief: dict) -> str:
    formula = brief.get('formula', 'PAS')
    formula_guide = FORMULA_GUIDES.get(formula, FORMULA_GUIDES['PAS'])
    return (
        f"Viết content quảng cáo Facebook chất lượng cao cho brief sau:\n\n"
        f"**Client/Brand:** {brief['client']}\n"
        f"**Sản phẩm/Dịch vụ:** {brief['product']}\n"
        f"**Mục tiêu campaign:** {brief['objective']}\n"
        f"**Target audience:** {brief['target']}\n"
        f"**Key message:** {brief['key_message']}\n"
        f"**Tone of voice:** {brief['tone']}\n"
        f"**Ngân sách ads:** {brief.get('budget', 'Chưa xác định')}\n\n"
        f"---\n"
        f"ÁP DỤNG CÔNG THỨC: **{formula}**\n\n"
        f"{formula_guide}\n\n"
        f"---\n\n"
        f"Tạo ra **3 bài viết Facebook** hoàn chỉnh, mỗi bài theo đúng cấu trúc output:\n"
        f"1. **Target Audience:** (tệp cụ thể cho bài này)\n"
        f"2. **Angle/Góc tiếp cận:** (lý do chọn góc này)\n"
        f"3. **Nội dung bài viết:** (áp dụng đúng {formula}, tối ưu 3 dòng đầu hiển thị trên Facebook)\n"
        f"4. **Gợi ý Visual:** (hình ảnh/video cụ thể đề xuất)\n\n"
        f"Phân cách rõ ràng giữa các bài bằng `---`."
    )

async def content_stream(brief: dict) -> AsyncGenerator[str, None]:
    prompt = _build_content_prompt(brief)
    ai_model = brief.get('ai_model', 'claude')

    if ai_model != 'claude':
        if not _gemini_pool:
            yield f"data: {json.dumps({'text': '❌ Chưa cấu hình GEMINI_API_KEY trong file .env'})}\n\n"
            yield "data: [DONE]\n\n"
            return
        gemini_model = ai_model  # e.g. "gemini-3.5-flash", "gemini-3.1-pro"
        try:
            async for chunk in await get_gemini_client().aio.models.generate_content_stream(
                model=gemini_model,
                contents=prompt,
                config=genai_types.GenerateContentConfig(
                    system_instruction=CONTENT_SYSTEM,
                    max_output_tokens=3000,
                )
            ):
                if chunk.text:
                    yield f"data: {json.dumps({'text': chunk.text})}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'text': f'❌ Gemini lỗi: {str(e)[:200]}'})}\n\n"
        yield "data: [DONE]\n\n"
    else:
        claude_model = ai_model if ai_model.startswith('claude-') else 'claude-sonnet-4-6'
        try:
            with get_anthropic_client().messages.stream(
                model=claude_model, max_tokens=3000,
                system=CONTENT_SYSTEM,
                messages=[{"role": "user", "content": prompt}]
            ) as stream:
                for text in stream.text_stream:
                    yield f"data: {json.dumps({'text': text})}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'text': f'❌ Claude lỗi: {str(e)[:200]}'})}\n\n"
        yield "data: [DONE]\n\n"

# ── Video Script ───────────────────────────────────────────────────

SCRIPT_SYSTEM = (
    "Ban la chuyen gia san xuat video content nganh noi that Viet Nam.\n"
    "Hieu sau ve cac format: TikTok, Instagram Reels, Facebook Video.\n"
    "Viet kich ban thuc te, co the quay duoc ngay, dung tam ly nguoi xem.\n"
    "Format ro rang theo tung scene, co loi thoai, huong dan visual, thoi luong."
)

async def script_stream(brief: dict) -> AsyncGenerator[str, None]:
    dur = brief['duration']
    prompt = (
        f"Viet kich ban video {brief['platform']} {dur} giay:\n\n"
        f"**San pham:** {brief['product']}\n"
        f"**Phong cach hook:** {brief['hook_style']}\n"
        f"**Target:** {brief['target']}\n"
        f"**Ghi chu them:** {brief.get('extra', 'Khong co')}\n\n"
        f"## KICH BAN {brief['platform'].upper()} — {dur}S\n\n"
        "**HOOK (0–3s):** [Loi thoai + mo ta visual + cam xuc can tao]\n\n"
        "**SCENE CHINH (4s–...s):** Chia tung scene 3-5 giay, moi scene co:\n"
        "- Loi thoai/caption on man hinh\n"
        "- Mo ta hinh anh/goc quay\n"
        "- Thoi luong\n\n"
        "**CTA (gio cuoi):** [Loi thoai + hanh dong nguoi xem can lam]\n\n"
        "**GHI CHU SAN XUAT:**\n"
        "- Goc quay de xuat\n"
        "- Anh sang\n"
        "- Nhac nen (phong cach/the loai)\n"
        "- Text overlay chinh\n"
        "- Caption va hashtag de dang kem"
    )
    ai_model = brief.get('ai_model', 'claude-sonnet-4-6')
    if ai_model != 'claude' and not ai_model.startswith('claude-'):
        if not _gemini_pool:
            yield f"data: {json.dumps({'text': '❌ Chưa cấu hình GEMINI_API_KEY trong file .env'})}\n\n"
            yield "data: [DONE]\n\n"
            return
        try:
            async for chunk in await get_gemini_client().aio.models.generate_content_stream(
                model=ai_model,
                contents=prompt,
                config=genai_types.GenerateContentConfig(
                    system_instruction=SCRIPT_SYSTEM,
                    max_output_tokens=3000,
                )
            ):
                if chunk.text:
                    yield f"data: {json.dumps({'text': chunk.text})}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'text': f'❌ Gemini lỗi: {str(e)[:200]}'})}\n\n"
        yield "data: [DONE]\n\n"
    else:
        claude_model = ai_model if ai_model.startswith('claude-') else 'claude-sonnet-4-6'
        try:
            with get_anthropic_client().messages.stream(
                model=claude_model, max_tokens=3000,
                system=SCRIPT_SYSTEM,
                messages=[{"role": "user", "content": prompt}]
            ) as stream:
                for text in stream.text_stream:
                    yield f"data: {json.dumps({'text': text})}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'text': f'❌ Claude lỗi: {str(e)[:200]}'})}\n\n"
        yield "data: [DONE]\n\n"

# ── Intake / Brief Client ──────────────────────────────────────────

INTAKE_SYSTEM = """Bạn là Giám đốc Marketing (CMO) với 14 năm kinh nghiệm thực chiến tại các Agency top 3 Việt Nam, làm việc cho agency Độc Media (chuyên ngành nội thất/kiến trúc/xây dựng). Chuyên môn: Performance Marketing, Conversion-focused strategies (PAS, AIDA), vận hành hệ thống tự động hóa & Data Tracking.

THÁI ĐỘ & GIỌNG VĂN: Chuyên nghiệp, trực diện, sắc bén, đi thẳng vào bản chất kinh doanh. Giao tiếp ở vị thế "Đối tác chiến lược" — không phải "Người làm thuê". TUYỆT ĐỐI không dùng từ ngữ sáo rỗng, hoa mỹ. Chỉ dùng ngôn ngữ đo lường được (Metrics, CVR, ROI, Lead, Data, CPL, CPA).

TƯ DUY CỐT LÕI:
- Phân tích client là tìm ra "Điểm nghẽn" (Bottleneck) trong phễu chuyển đổi, không phải thu thập thông tin bề nổi.
- Luôn bóc tách rõ: Mục tiêu là Branding (Nhận diện) hay Performance (Lead/Doanh số)?
- Rủi ro luôn nằm ở Vận hành & Phối hợp (SLA). Phải rào trước mọi rủi ro về Data, Tỷ lệ chốt Sale và Dòng tiền.

BENCHMARK THỰC TẾ (nội thất/kiến trúc VN — xác nhận từ thực chiến):
- CPMess (tin nhắn): 65.000 – 90.000 VND | CPM Tương tác: 21.000 – 34.000 VND
- CPL (chi phí/inbox): 166.000 – 315.000 VND | Tốt: <230K | TB: 230–315K | Dừng: >350K
- CTR Inbox: 1.3–2.1% | CTR Engagement: 2.9–4.5%

QUY TRÌNH XUẤT BRIEF — 4 BƯỚC BẮT BUỘC:

## Bước 1: Deep-Dive Business Audit
Bảng phân tích: Sản phẩm/dịch vụ cốt lõi · USP & Lợi thế cạnh tranh · Top 3 đối thủ trực tiếp & lỗ hổng truyền thông của họ · Insight khách hàng (pain points: rủi ro thi công, chậm tiến độ, phát sinh chi phí).

## Bước 2: Stakeholder Mapping
Xác định 3 nhóm bên phía client và phương án giao tiếp:
- Decision Maker (CEO/Founder): báo cáo ROI, Doanh thu, Tối ưu chi phí, Vị thế thương hiệu.
- Gatekeeper (Marketing/Kế toán): tiến độ, chất lượng Content, SLA, hồ sơ thanh toán.
- User (Sale/CSKH): chất lượng Lead, kịch bản chốt Sale, phản hồi thị trường.

## Bước 3: Strategic Onboarding & Technical Setup
Checklist tài sản kỹ thuật cần bàn giao: quyền BM, Ad Accounts, Fanpage, Zalo OA. Hệ thống Tracking: luồng dữ liệu tự động (Lead + chi phí đổ về Google Sheets realtime). Định nghĩa KPI: thống nhất "Qualified Lead" và mức CPA/CPL mục tiêu.

## Bước 4: SLA & Risk Management
Thiết lập "Luật chơi": thời gian Agency lên camp/sửa content, thời gian client duyệt bài. Cảnh báo rủi ro vận hành Sale: nếu tỷ lệ chốt thấp do Sale chậm gọi hoặc kịch bản yếu — đề xuất Agency hỗ trợ tối ưu kịch bản.

FORMAT OUTPUT: Markdown (.md) với Headings H2/H3, Bullet points, Bảng biểu. Trực quan, rõ ràng, đậm chất kỹ thuật chuyên môn."""

async def intake_stream(body: dict) -> AsyncGenerator[str, None]:
    lines = [
        f"Ten doanh nghiep: {body['ten_dn']}",
        f"Linh vuc: {body['linh_vuc']}",
        f"Fanpage/Website: {body['fanpage_url'] or 'Chua cung cap'}",
        f"Nguoi lien he: {body['nguoi_lien_he'] or 'Chua cung cap'}",
        f"Tinh trang fanpage: {body['fanpage_status']} | Followers: {body['followers']}",
        f"Da chay ads chua: {body['da_chay_ads']}" + (f" | CPL cu: {body['cpl_cu']}" if body['cpl_cu'] else ""),
        f"Muc tieu chinh: {body['muc_tieu_chinh']}",
        f"Leads mong muon/thang: {body['leads_mong_muon'] or 'Chua xac dinh'}",
        f"Doanh thu mong muon/thang: {body['doanh_thu_mong_muon'] or 'Chua xac dinh'}",
        f"Timeline bat dau: {body['timeline_bat_dau'] or 'Cang som cang tot'}",
        f"Mo ta san pham/dich vu: {body['mo_ta_sp'] or 'Chua mo ta'}",
        f"Doi tuong khach hang: {body['doi_tuong'] or 'Chua mo ta'}",
        f"USP / Diem khac biet: {body['usp'] or 'Chua mo ta'}",
        f"Assets san co: {body['assets'] or 'Chua ro'}",
        f"Ngan sach quang cao/thang: {body['ngan_sach_ads'] or 'Chua xac dinh'} VND",
        f"Ngan sach dich vu/thang: {body['ngan_sach_dv'] or 'Chua xac dinh'} VND",
        f"Thoi gian hop dong: {body['thoi_gian_hd']}",
        f"Doi thu chinh: {body['doi_thu'] or 'Chua cung cap'}",
    ]
    prompt = "Thong tin intake client:\n" + "\n".join(f"- {l}" for l in lines)
    prompt += "\n\nTao brief chien luoc day du va thuc te cho client nay."

    ai_model = body.get('ai_model', 'claude-sonnet-4-6')
    if ai_model != 'claude' and not ai_model.startswith('claude-'):
        if not _gemini_pool:
            yield f"data: {json.dumps({'text': 'Loi: GEMINI_API_KEY chua duoc cau hinh.'})}\n\n"
            yield "data: [DONE]\n\n"; return
        try:
            async for chunk in await get_gemini_client().aio.models.generate_content_stream(
                model=ai_model, contents=prompt,
                config=genai_types.GenerateContentConfig(
                    system_instruction=INTAKE_SYSTEM, max_output_tokens=3500)
            ):
                if chunk.text:
                    yield f"data: {json.dumps({'text': chunk.text})}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'text': f'❌ Gemini lỗi: {str(e)[:200]}'})}\n\n"
    else:
        claude_model = ai_model if ai_model.startswith('claude-') else 'claude-sonnet-4-6'
        try:
            with get_anthropic_client().messages.stream(
                model=claude_model, max_tokens=3500,
                system=INTAKE_SYSTEM,
                messages=[{"role": "user", "content": prompt}]
            ) as stream:
                for text in stream.text_stream:
                    yield f"data: {json.dumps({'text': text})}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'text': f'❌ Claude lỗi: {str(e)[:200]}'})}\n\n"
    yield "data: [DONE]\n\n"

# ── Quote Generator ────────────────────────────────────────────────

QUOTE_SYSTEM = """Ban la chuyen gia bao gia dich vu Facebook Ads tai Viet Nam, lam viec cho agency Doc Media (chuyen nganh noi that/kien truc).

BENCHMARK THUC TE (data ADS Manager nganh noi that VN):
- CPM Tin nhan (Message objective): 150.000 - 230.000 VND
- CPM Tuong tac (Engagement): 21.000 - 34.000 VND
- CPL (chi phi moi inbox): 166.000 - 315.000 VND | Tot: <230K | Dung: >350K
- CTR Inbox: 1.3 - 2.1% | CTR Tuong tac: 2.9 - 4.5%
- Leads du kien voi 3 trieu/10 ngay: 10-18 inbox

DICH VU DOC MEDIA BAO GOM:
- Thiet lap chien dich (setup): cai dat pixel, cau truc campaign, audiences, creatives
- Quan ly & toi uu hang ngay: check sang, pause/scale adsets, A/B test
- San xuat noi dung: viet copy bai viet, brief design/video, CTA
- Bao cao hang tuan: so lieu thuc te, nhan xet, khuyen nghi
- Tu van chien luoc: targeting, budget allocation, campaign structure

TAO BAO GIA CHUYEN NGHIEP, DAY DU, SU DUNG BANG MARKDOWN.
Viet tieng Viet, ton trong, chuyen nghiep. Co day du: phi dich vu, KPI cam ket, dieu khoan."""

async def quote_stream(body: dict) -> AsyncGenerator[str, None]:
    prompt = f"""Tao bao gia dich vu Facebook Ads cho:
- Ten doanh nghiep: {body['ten_client']}
- Linh vuc: {body['linh_vuc']}
- Muc tieu chien dich: {body['muc_tieu']}
- Ngan sach quang cao/thang: {body['ngan_sach_ads']} VND
- Thoi gian hop dong: {body['thoi_gian']}
- Dich vu yeu cau: {body['dich_vu'] or 'Toan bo (setup + quan ly + content + bao cao)'}
- Ghi chu them: {body['ghi_chu'] or 'Khong co'}

Yeu cau bao gia bao gom:
1. Tieu de + thong tin Doc Media + ngay bao gia
2. Thong tin khach hang
3. Pham vi dich vu chi tiet (bullet points)
4. Bang gia chi tiet (markdown table): Hang muc | Mo ta | Don gia/thang
5. Tong gia tri hop dong
6. KPI cam ket dua tren benchmark thuc te (CPM, CPL, so leads du kien)
7. Dieu khoan: thanh toan, thoi han bao gia, chinh sach dieu chinh
8. Ky ten: Doc Media - Nguyen Thanh Hai

Luu y: phi dich vu phai hop ly voi nganh Facebook Ads agency VN 2026."""

    ai_model = body.get('ai_model', 'gemini-3.5-flash')
    if ai_model != 'claude' and not ai_model.startswith('claude-'):
        if not _gemini_pool:
            yield f"data: {json.dumps({'text': 'Loi: GEMINI_API_KEY chua duoc cau hinh.'})}\n\n"
            yield "data: [DONE]\n\n"; return
        try:
            async for chunk in await get_gemini_client().aio.models.generate_content_stream(
                model=ai_model, contents=prompt,
                config=genai_types.GenerateContentConfig(
                    system_instruction=QUOTE_SYSTEM, max_output_tokens=3000)
            ):
                if chunk.text:
                    yield f"data: {json.dumps({'text': chunk.text})}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'text': f'[Gemini loi: {str(e)[:120]}. Thu lai hoac doi sang Claude.]'})}\n\n"
    else:
        claude_model = ai_model if ai_model.startswith('claude-') else 'claude-sonnet-4-6'
        with get_anthropic_client().messages.stream(
            model=claude_model, max_tokens=3000,
            system=QUOTE_SYSTEM,
            messages=[{"role": "user", "content": prompt}]
        ) as stream:
            for text in stream.text_stream:
                yield f"data: {json.dumps({'text': text})}\n\n"
    yield "data: [DONE]\n\n"

# ── Facebook Ads API ───────────────────────────────────────────────

def _fb_request(url: str) -> dict:
    req = urllib.request.Request(url)
    try:
        with urllib.request.urlopen(req, timeout=30) as r:
            return json.loads(r.read().decode())
    except urllib.error.HTTPError as e:
        body = e.read().decode()[:300]
        raise ValueError(f"Facebook API loi {e.code}: {body}")

def _get_fb_accounts(token: str) -> list:
    url = (
        f"https://graph.facebook.com/v21.0/me/adaccounts"
        f"?fields=name,account_id,account_status,currency,amount_spent"
        f"&limit=50&access_token={token}"
    )
    data = _fb_request(url)
    return data.get("data", [])

def _get_fb_data(account_id: str, period: str, date_from: str = "", date_to: str = "") -> dict:
    token = os.getenv("FB_ACCESS_TOKEN", "")
    if not token or token.startswith("your_"):
        raise ValueError("FB_ACCESS_TOKEN chua duoc cai dat trong .env")
    if not account_id:
        account_id = os.getenv("FB_AD_ACCOUNT_ID", "")
    if not account_id:
        raise ValueError("Chua chon tai khoan quang cao")
    if not account_id.startswith("act_"):
        account_id = f"act_{account_id}"

    if period == "custom" and date_from and date_to:
        time_param = "&time_range=" + urllib.parse.quote(json.dumps({"since": date_from, "until": date_to}))
    else:
        time_param = f"&date_preset={period}"

    fields = "campaign_name,spend,reach,impressions,cpm,cpc,ctr,actions,objective,date_start,date_stop"
    insights_url = (
        f"https://graph.facebook.com/v21.0/{account_id}/insights"
        f"?level=campaign&fields={fields}{time_param}"
        f"&limit=30&access_token={token}"
    )
    camp_url = (
        f"https://graph.facebook.com/v21.0/{account_id}/campaigns"
        f"?fields=name,status,objective,daily_budget,lifetime_budget"
        f"&limit=30&access_token={token}"
    )
    insights = _fb_request(insights_url)
    camps = _fb_request(camp_url)
    return {"insights": insights, "campaigns": camps, "account_id": account_id}

async def ads_analyze_stream(data: dict, ai_model: str = "gemini-3.5-flash") -> AsyncGenerator[str, None]:
    summary_text = json.dumps(data, ensure_ascii=False, indent=2)[:3000]
    prompt = (
        "Phan tich du lieu Facebook Ads sau day cua agency Doc Media (nganh noi that/kien truc VN):\n\n"
        f"{summary_text}\n\n"
        "Cho biet:\n"
        "1. Danh gia tong the hieu qua (benchmark nganh noi that VN: CPM Inbox 150-230K, CPM Tuong tac 21-34K, CPL tot <230K, kem >350K)\n"
        "2. Campaign nao dang hoat dong tot nhat / kem nhat va vi sao\n"
        "3. 3 de xuat toi uu cu the, thuc te, co the lam ngay\n"
        "4. Du bao neu giu ngan sach hien tai, ket qua tuan toi se nhu the nao\n\n"
        "Tra loi bang tieng Viet, co so lieu cu the, ngan gon va thuc te."
    )
    if ai_model != 'claude' and not ai_model.startswith('claude-'):
        try:
            async for chunk in await get_gemini_client().aio.models.generate_content_stream(
                model=ai_model, contents=prompt,
                config=genai_types.GenerateContentConfig(
                    system_instruction=CHAT_SYSTEM, max_output_tokens=1500)
            ):
                if chunk.text:
                    yield f"data: {json.dumps({'text': chunk.text})}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'text': f'[Gemini loi: {str(e)[:120]}. Thu lai hoac doi sang Claude.]'})}\n\n"
    else:
        claude_model = ai_model if ai_model.startswith('claude-') else 'claude-sonnet-4-6'
        try:
            with get_anthropic_client().messages.stream(
                model=claude_model, max_tokens=1500,
                system=CHAT_SYSTEM,
                messages=[{"role": "user", "content": prompt}]
            ) as stream:
                for text in stream.text_stream:
                    yield f"data: {json.dumps({'text': text})}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'text': f'❌ Claude lỗi: {str(e)[:200]}'})}\n\n"
    yield "data: [DONE]\n\n"

# ── Pydantic Models ────────────────────────────────────────────────

class ChatBody(BaseModel):
    message: str
    session_id: str
    file_context: str = ""

class ImageBody(BaseModel):
    description: str
    style: str = "Modern Minimalist"

class ImageEnhanceBody(BaseModel):
    description: str
    style: str = "Modern Minimalist"
    ai_model: str = "gemini-3.1-flash-lite"

class ContentBody(BaseModel):
    client: str
    product: str
    objective: str
    target: str
    key_message: str
    tone: str = "Chuyên nghiệp, uy tín, sang trọng"
    budget: str = ""
    formula: str = "PAS"
    ai_model: str = "claude-sonnet-4-6"

class ScriptBody(BaseModel):
    platform: str = "TikTok"
    duration: int = 30
    product: str
    hook_style: str
    target: str
    extra: str = ""
    ai_model: str = "claude-sonnet-4-6"

class AdsAnalyzeBody(BaseModel):
    data: dict
    ai_model: str = "gemini-3.5-flash"

class LibSaveBody(BaseModel):
    type: str
    title: str
    content: str
    client: str = ""

class ReportDataBody(BaseModel):
    account_id: str
    period: str = "this_month"
    date_from: str = ""
    date_to: str = ""

class ReportAnalyzeBody(BaseModel):
    fb_data: dict
    client_name: str = "Khách hàng"
    period_label: str = ""
    ai_model: str = "gemini-3.5-flash"

class ReportExportBody(BaseModel):
    fb_data: dict
    ai_text: str
    client_name: str = ""
    period_label: str = ""

class IntakeBody(BaseModel):
    # Co ban
    ten_dn: str
    linh_vuc: str = ""
    fanpage_url: str = ""
    nguoi_lien_he: str = ""
    # Tinh trang hien tai
    fanpage_status: str = "Chua co fanpage"
    followers: str = "0"
    da_chay_ads: str = "Chua chay bao gio"
    cpl_cu: str = ""
    # Muc tieu
    muc_tieu_chinh: str = "Thu leads tin nhan"
    leads_mong_muon: str = ""
    doanh_thu_mong_muon: str = ""
    timeline_bat_dau: str = ""
    # San pham & doi tuong
    mo_ta_sp: str = ""
    doi_tuong: str = ""
    usp: str = ""
    # Assets & ngan sach
    assets: str = ""
    ngan_sach_ads: str = ""
    ngan_sach_dv: str = ""
    thoi_gian_hd: str = "3 thang"
    # Doi thu
    doi_thu: str = ""
    # AI model
    ai_model: str = "claude-sonnet-4-6"

class QuoteBody(BaseModel):
    ten_client: str
    linh_vuc: str = "Noi that / Thiet ke & Thi cong"
    muc_tieu: str = "Tin nhan (Lead), Tuong tac"
    ai_model: str = "gemini-3.5-flash"
    ngan_sach_ads: str = ""
    thoi_gian: str = "3 thang"
    dich_vu: str = ""
    ghi_chu: str = ""

# ── Routes ─────────────────────────────────────────────────────────

@app.get("/", response_class=HTMLResponse)
async def index():
    return HTMLResponse(content=HTML)

# Chat
@app.get("/sessions")
async def get_sessions():
    return JSONResponse(db_get_sessions())

@app.get("/history/{session_id}")
async def get_history(session_id: str):
    return JSONResponse(db_get_history(session_id))

@app.delete("/session/{session_id}")
async def clear_session(session_id: str):
    db_clear_session(session_id)
    return {"ok": True}

# ── Alerts ─────────────────────────────────────────────────────────

@app.get("/alerts")
async def get_alerts():
    return JSONResponse({"alerts": db_get_alerts(), "unread": db_unread_count()})

@app.post("/alerts/{alert_id}/read")
async def mark_read(alert_id: str):
    db_mark_alert_read(alert_id)
    return {"ok": True}

@app.post("/alerts/read-all")
async def mark_all_read():
    with sqlite3.connect(DB_PATH) as c:
        c.execute("UPDATE alerts SET read=1")
    return {"ok": True}

# ── Content Library endpoints ──────────────────────────────────────
@app.post("/library/save")
async def library_save(body: LibSaveBody):
    item_id = db_lib_save(body.type, body.title, body.content, body.client)
    return {"ok": True, "id": item_id}

@app.get("/library")
async def library_list(type: str = "", search: str = ""):
    return {"items": db_lib_list(type, search)}

@app.delete("/library/{item_id}")
async def library_delete(item_id: str):
    db_lib_delete(item_id)
    return {"ok": True}

# ── Monthly Report endpoints ──────────────────────────────────────
@app.post("/report/data")
async def report_data(body: ReportDataBody):
    period = body.period
    date_from, date_to = body.date_from, body.date_to
    if period in ("this_month", "last_month") and not date_from:
        date_from, date_to = _period_to_dates(period)
        period = "custom"
    try:
        data = await asyncio.to_thread(_get_fb_data, body.account_id, period, date_from, date_to)
        kpis = _compute_report_kpis(data)
        campaigns = _compute_campaign_rows(data)
        return {"ok": True, "fb_data": data, "kpis": kpis, "campaigns": campaigns}
    except Exception as e:
        raise HTTPException(400, str(e)[:300])

@app.post("/report/analyze")
async def report_analyze(body: ReportAnalyzeBody):
    return StreamingResponse(
        report_analyze_stream(body.fb_data, body.client_name, body.period_label, body.ai_model),
        media_type="text/event-stream",
        headers={"Cache-Control":"no-cache","X-Accel-Buffering":"no"})

@app.post("/report/export/{fmt}")
async def report_export(fmt: str, body: ReportExportBody):
    kpis = _compute_report_kpis(body.fb_data)
    campaigns = _compute_campaign_rows(body.fb_data)
    meta = {"client_name": body.client_name, "period_label": body.period_label,
            "export_date": datetime.now().strftime("%d/%m/%Y")}
    ts = datetime.now().strftime("%Y%m%d_%H%M")
    client_slug = re.sub(r'[^a-zA-Z0-9]', '_', body.client_name)[:20] or "Client"
    if fmt == "word":
        data = export_report_word(meta, kpis, campaigns, body.ai_text)
        fname = f"BaoCao_{client_slug}_{ts}.docx"
        media = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif fmt == "excel":
        data = export_report_excel(meta, kpis, campaigns, body.ai_text)
        fname = f"BaoCao_{client_slug}_{ts}.xlsx"
        media = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    elif fmt == "pdf":
        data = export_report_pdf(meta, kpis, campaigns, body.ai_text)
        fname = f"BaoCao_{client_slug}_{ts}.pdf"
        media = "application/pdf"
    else:
        raise HTTPException(400, "fmt phải là word|excel|pdf")
    from urllib.parse import quote as urlquote
    return StreamingResponse(io.BytesIO(data), media_type=media,
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{urlquote(fname)}"})

@app.get("/library/export/{fmt}")
async def library_export(fmt: str, type: str = "", search: str = ""):
    items = db_lib_list(type, search)
    if not items:
        raise HTTPException(400, "Không có nội dung để xuất")
    ts = datetime.now().strftime("%Y%m%d_%H%M")
    if fmt == "word":
        data = export_library_word(items)
        fname = f"DocMedia_ThuVien_{ts}.docx"
        media = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif fmt == "excel":
        data = export_library_excel(items)
        fname = f"DocMedia_ThuVien_{ts}.xlsx"
        media = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    elif fmt == "pdf":
        data = export_library_pdf(items)
        fname = f"DocMedia_ThuVien_{ts}.pdf"
        media = "application/pdf"
    else:
        raise HTTPException(400, "fmt phải là word|excel|pdf")
    from urllib.parse import quote
    return StreamingResponse(
        io.BytesIO(data), media_type=media,
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{quote(fname)}"})

# ── Monthly Report helpers ────────────────────────────────────────────
import calendar

def _compute_report_kpis(fb_data: dict) -> dict:
    rows = fb_data.get("insights", {}).get("data", [])
    if not rows:
        return {}
    total_spend = sum(float(r.get("spend", 0)) for r in rows)
    total_reach = sum(int(r.get("reach", 0)) for r in rows)
    total_imp   = sum(int(r.get("impressions", 0)) for r in rows)
    total_msgs  = 0
    total_leads = 0
    for r in rows:
        for a in r.get("actions", []):
            t = a.get("action_type", "")
            v = int(a.get("value", 0))
            if "messaging_conversation" in t or "flow_complete" in t:
                total_msgs += v
            elif t in ("lead", "onsite_conversion.lead_grouped"):
                total_leads += v
    cpm     = total_spend / total_imp   * 1000 if total_imp   else 0
    cpm_msg = total_spend / total_msgs         if total_msgs  else 0
    cpl     = total_spend / total_leads        if total_leads else 0
    return {"spend": total_spend, "reach": total_reach, "impressions": total_imp,
            "cpm": cpm, "messages": total_msgs, "leads": total_leads,
            "cpm_msg": cpm_msg, "cpl": cpl}

def _compute_campaign_rows(fb_data: dict) -> list:
    rows = fb_data.get("insights", {}).get("data", [])
    result = []
    for r in rows:
        spend = float(r.get("spend", 0))
        imp   = int(r.get("impressions", 0))
        msgs  = leads = 0
        for a in r.get("actions", []):
            t = a.get("action_type", ""); v = int(a.get("value", 0))
            if "messaging_conversation" in t or "flow_complete" in t: msgs += v
            elif t in ("lead", "onsite_conversion.lead_grouped"): leads += v
        cpm = spend / imp * 1000 if imp else 0
        result.append({"name": r.get("campaign_name",""), "spend": spend,
                        "reach": int(r.get("reach",0)), "impressions": imp,
                        "cpm": cpm, "messages": msgs, "leads": leads,
                        "date_start": r.get("date_start",""), "date_stop": r.get("date_stop","")})
    return sorted(result, key=lambda x: x["spend"], reverse=True)

def _period_to_dates(period: str) -> tuple[str, str]:
    """Convert 'last_month'/'this_month' to (date_from, date_to) strings."""
    today = datetime.now().date()
    if period == "last_month":
        m = today.month - 1 if today.month > 1 else 12
        y = today.year if today.month > 1 else today.year - 1
        last_day = calendar.monthrange(y, m)[1]
        return f"{y}-{m:02d}-01", f"{y}-{m:02d}-{last_day}"
    if period == "this_month":
        return today.strftime("%Y-%m-01"), today.strftime("%Y-%m-%d")
    return "", ""

REPORT_SYSTEM = """Bạn là chuyên gia phân tích hiệu suất Facebook Ads, viết báo cáo tháng chuyên nghiệp cho agency Độc Media (chuyên ngành nội thất/kiến trúc Việt Nam).

GIỌNG VĂN: Chuyên nghiệp, số liệu cụ thể, nhận xét thực tế, ngôn ngữ B2B — có thể gửi thẳng cho khách hàng.

BENCHMARK NỘI THẤT VN (2026):
- CPM Tin nhắn: 150.000–230.000 VND | Tốt: <150K | Kém: >250K
- CPL (chi phí/inbox): <230.000 VND tốt | >350.000 VND cần tối ưu
- CTR Inbox: >1.5% | Số inbox với 15tr ngân sách: 60–80/tháng

FORMAT BẮT BUỘC (Markdown với emoji):
## 📊 Đánh giá tổng thể
## ✅ Điểm mạnh tháng này
## ⚠️ Điểm cần cải thiện
## 🏆 Chiến dịch nổi bật (tốt nhất & kém nhất, kèm số liệu)
## 🎯 Đề xuất tháng tới (3 hành động cụ thể, có thể làm ngay)
## 🔮 Dự báo (nếu duy trì ngân sách hiện tại)"""

async def report_analyze_stream(fb_data: dict, client_name: str, period_label: str,
                                 ai_model: str = "gemini-3.5-flash") -> AsyncGenerator[str, None]:
    kpis = _compute_report_kpis(fb_data)
    camps = _compute_campaign_rows(fb_data)
    prompt = (
        f"Viết báo cáo hiệu suất Facebook Ads kỳ [{period_label}] cho client: {client_name}\n\n"
        f"KPI TỔNG HỢP:\n"
        f"- Chi tiêu: {kpis.get('spend',0):,.0f} VND\n"
        f"- Reach: {kpis.get('reach',0):,} người | Impressions: {kpis.get('impressions',0):,}\n"
        f"- CPM trung bình: {kpis.get('cpm',0):,.0f} VND\n"
        f"- Số tin nhắn (Inbox): {kpis.get('messages',0)} | Chi phí/Inbox: {kpis.get('cpm_msg',0):,.0f} VND\n"
        f"- Số KHTN (Lead): {kpis.get('leads',0)} | CPL: {kpis.get('cpl',0):,.0f} VND\n\n"
        f"CHIẾN DỊCH ({len(camps)} chiến dịch):\n"
        + "\n".join(f"- {c['name']}: chi {c['spend']:,.0f}đ | CPM {c['cpm']:,.0f} | inbox {c['messages']} | lead {c['leads']}"
                    for c in camps[:8])
        + "\n\nViết báo cáo đầy đủ theo format yêu cầu, ngôn ngữ Tiếng Việt chuyên nghiệp."
    )
    if ai_model != 'claude' and not ai_model.startswith('claude-'):
        try:
            async for chunk in await get_gemini_client().aio.models.generate_content_stream(
                model=ai_model, contents=prompt,
                config=genai_types.GenerateContentConfig(system_instruction=REPORT_SYSTEM, max_output_tokens=2500)
            ):
                if chunk.text:
                    yield f"data: {json.dumps({'t': chunk.text})}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'t': f'❌ Gemini lỗi: {str(e)[:200]}'})}\n\n"
    else:
        claude_model = ai_model if ai_model.startswith('claude-') else 'claude-sonnet-4-6'
        try:
            with get_anthropic_client().messages.stream(model=claude_model, max_tokens=2500,
                system=REPORT_SYSTEM, messages=[{"role":"user","content":prompt}]) as stream:
                for text in stream.text_stream:
                    yield f"data: {json.dumps({'t': text})}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'t': f'❌ Claude lỗi: {str(e)[:200]}'})}\n\n"
    yield "data: [DONE]\n\n"

def export_report_word(meta: dict, kpis: dict, campaigns: list, ai_text: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document()
    # Cover
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ĐỘC MEDIA'); run.font.size = Pt(22); run.font.bold = True
    run.font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run('AGENCY FACEBOOK ADS — NỘI THẤT & KIẾN TRÚC').font.size = Pt(11)
    doc.add_paragraph()
    h = doc.add_heading('BÁO CÁO HIỆU SUẤT QUẢNG CÁO', 0); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.runs[0].font.color.rgb = RGBColor(0x1F, 0x1F, 0x3A)
    pi = doc.add_paragraph(); pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pi.add_run(f"Client: {meta.get('client_name','—')}   |   Kỳ: {meta.get('period_label','—')}   |   Xuất ngày: {meta.get('export_date','—')}")
    doc.add_page_break()
    # KPI section
    doc.add_heading('I. TỔNG QUAN KPI', 1)
    kpi_rows = [
        ('Chi tiêu', f"{kpis.get('spend',0):,.0f} VND"),
        ('Reach', f"{kpis.get('reach',0):,} người"),
        ('Impressions', f"{kpis.get('impressions',0):,}"),
        ('CPM trung bình', f"{kpis.get('cpm',0):,.0f} VND"),
        ('Số tin nhắn (Inbox)', str(kpis.get('messages',0))),
        ('Chi phí / Inbox', f"{kpis.get('cpm_msg',0):,.0f} VND"),
        ('Số KHTN (Lead)', str(kpis.get('leads',0))),
        ('CPL', f"{kpis.get('cpl',0):,.0f} VND"),
    ]
    t = doc.add_table(rows=1, cols=2); t.style = 'Table Grid'
    t.rows[0].cells[0].text = 'Chỉ số'; t.rows[0].cells[1].text = 'Giá trị'
    for cell in t.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for name, val in kpi_rows:
        row = t.add_row(); row.cells[0].text = name; row.cells[1].text = val
    doc.add_paragraph()
    # Campaign section
    if campaigns:
        doc.add_heading('II. CHI TIẾT CHIẾN DỊCH', 1)
        tc = doc.add_table(rows=1, cols=6); tc.style = 'Table Grid'
        headers = ['Chiến dịch', 'Chi tiêu (VND)', 'Reach', 'CPM', 'Inbox', 'Lead']
        for i, h in enumerate(headers):
            tc.rows[0].cells[i].text = h
            tc.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
        for c in campaigns:
            row = tc.add_row()
            row.cells[0].text = c['name'][:40]
            row.cells[1].text = f"{c['spend']:,.0f}"
            row.cells[2].text = f"{c['reach']:,}"
            row.cells[3].text = f"{c['cpm']:,.0f}"
            row.cells[4].text = str(c['messages'])
            row.cells[5].text = str(c['leads'])
        doc.add_paragraph()
    # AI section
    doc.add_heading('III. PHÂN TÍCH & ĐỀ XUẤT', 1)
    doc.add_paragraph(_md_to_plain(ai_text))
    # Footer
    doc.add_paragraph()
    fp = doc.add_paragraph('Độc Media — Agency Facebook Ads chuyên ngành Nội thất & Kiến trúc')
    fp.runs[0].font.size = Pt(9); fp.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x99)
    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.getvalue()

def export_report_excel(meta: dict, kpis: dict, campaigns: list, ai_text: str) -> bytes:
    import openpyxl
    from openpyxl.styles import Font as XFont, PatternFill, Alignment, Border, Side
    thin = Side(style='thin', color='CCCCCC')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    wb = openpyxl.Workbook()
    # Sheet 1: KPI
    ws1 = wb.active; ws1.title = 'Tóm tắt KPI'
    ws1.merge_cells('A1:B1')
    ws1['A1'] = f"BÁO CÁO FB ADS — {meta.get('client_name','')} — {meta.get('period_label','')}"
    ws1['A1'].font = XFont(bold=True, size=14, color='7C3AED')
    ws1['A1'].alignment = Alignment(horizontal='center')
    ws1.row_dimensions[1].height = 32
    ws1.append([]); ws1.append(['Chỉ số', 'Giá trị'])
    for cell in ws1[3]: cell.font = XFont(bold=True, color='FFFFFF'); cell.fill = PatternFill('solid', fgColor='7C3AED'); cell.alignment = Alignment(horizontal='center')
    kpi_rows = [
        ('Chi tiêu', f"{kpis.get('spend',0):,.0f} VND"),
        ('Reach', f"{kpis.get('reach',0):,} người"),
        ('Impressions', f"{kpis.get('impressions',0):,}"),
        ('CPM trung bình', f"{kpis.get('cpm',0):,.0f} VND"),
        ('Số tin nhắn (Inbox)', kpis.get('messages',0)),
        ('Chi phí / Inbox', f"{kpis.get('cpm_msg',0):,.0f} VND"),
        ('Số KHTN (Lead)', kpis.get('leads',0)),
        ('CPL', f"{kpis.get('cpl',0):,.0f} VND"),
        ('Ngày xuất', meta.get('export_date','')),
    ]
    for row in kpi_rows:
        ws1.append(list(row))
    ws1.column_dimensions['A'].width = 28; ws1.column_dimensions['B'].width = 22
    # Sheet 2: Campaigns
    ws2 = wb.create_sheet('Chi tiết chiến dịch')
    headers = ['Chiến dịch', 'Chi tiêu (VND)', 'Reach', 'Impressions', 'CPM', 'Inbox', 'Lead', 'Ngày bắt đầu', 'Ngày kết thúc']
    ws2.append(headers)
    for cell in ws2[1]: cell.font = XFont(bold=True, color='FFFFFF'); cell.fill = PatternFill('solid', fgColor='1D4ED8'); cell.border = border
    for c in campaigns:
        ws2.append([c['name'], c['spend'], c['reach'], c['impressions'], round(c['cpm']), c['messages'], c['leads'], c['date_start'], c['date_stop']])
    ws2.column_dimensions['A'].width = 40
    for col in ['B','C','D','E','F','G']: ws2.column_dimensions[col].width = 16
    # Sheet 3: AI
    ws3 = wb.create_sheet('Phân tích & Đề xuất')
    ws3['A1'] = 'PHÂN TÍCH AI — ĐỀ XUẤT THÁNG TỚI'
    ws3['A1'].font = XFont(bold=True, size=13, color='7C3AED')
    ws3['A2'] = _md_to_plain(ai_text)
    ws3['A2'].alignment = Alignment(wrap_text=True, vertical='top')
    ws3.row_dimensions[2].height = 400
    ws3.column_dimensions['A'].width = 100
    buf = io.BytesIO(); wb.save(buf); buf.seek(0)
    return buf.getvalue()

def export_report_pdf(meta: dict, kpis: dict, campaigns: list, ai_text: str) -> bytes:
    from fpdf import FPDF
    FONT_R = 'C:/Windows/Fonts/arial.ttf'; FONT_B = 'C:/Windows/Fonts/arialbd.ttf'
    pdf = FPDF();
    try:
        pdf.add_font('F', '', FONT_R); pdf.add_font('F', 'B', FONT_B); fn = 'F'
    except Exception:
        fn = 'Helvetica'
    pdf.set_auto_page_break(auto=True, margin=15)
    # Cover
    pdf.add_page()
    pdf.set_font(fn, 'B', 22); pdf.set_text_color(124,58,237)
    pdf.cell(0, 14, 'DOC MEDIA', ln=True, align='C')
    pdf.set_font(fn, '', 11); pdf.set_text_color(100,100,120)
    pdf.cell(0, 8, 'AGENCY FACEBOOK ADS - NOI THAT & KIEN TRUC', ln=True, align='C')
    pdf.ln(10)
    pdf.set_font(fn, 'B', 16); pdf.set_text_color(30,30,50)
    pdf.cell(0, 12, 'BAO CAO HIEU SUAT QUANG CAO', ln=True, align='C')
    pdf.set_font(fn, '', 11); pdf.set_text_color(100,100,120)
    pdf.cell(0, 8, f"Client: {meta.get('client_name','')}  |  Ky: {meta.get('period_label','')}  |  Xuat: {meta.get('export_date','')}", ln=True, align='C')
    pdf.ln(10)
    # KPI table
    pdf.set_font(fn, 'B', 13); pdf.set_text_color(30,30,50)
    pdf.cell(0, 10, 'I. TONG QUAN KPI', ln=True)
    pdf.set_fill_color(124,58,237); pdf.set_text_color(255,255,255)
    pdf.set_font(fn, 'B', 10)
    pdf.cell(90, 8, 'Chi so', fill=True); pdf.cell(90, 8, 'Gia tri', fill=True, ln=True)
    kpi_rows = [
        ('Chi tieu', f"{kpis.get('spend',0):,.0f} VND"),
        ('Reach', f"{kpis.get('reach',0):,} nguoi"),
        ('Impressions', f"{kpis.get('impressions',0):,}"),
        ('CPM trung binh', f"{kpis.get('cpm',0):,.0f} VND"),
        ('So tin nhan (Inbox)', str(kpis.get('messages',0))),
        ('Chi phi / Inbox', f"{kpis.get('cpm_msg',0):,.0f} VND"),
        ('So KHTN (Lead)', str(kpis.get('leads',0))),
        ('CPL', f"{kpis.get('cpl',0):,.0f} VND"),
    ]
    pdf.set_text_color(40,40,60); pdf.set_font(fn, '', 10)
    for i, (k, v) in enumerate(kpi_rows):
        fill = i % 2 == 0
        pdf.set_fill_color(245,244,255) if fill else pdf.set_fill_color(255,255,255)
        pdf.cell(90, 7, k, fill=True); pdf.cell(90, 7, v, fill=True, ln=True)
    pdf.ln(6)
    # Campaign table
    if campaigns:
        pdf.set_font(fn, 'B', 13); pdf.set_text_color(30,30,50)
        pdf.cell(0, 10, 'II. CHI TIET CHIEN DICH', ln=True)
        pdf.set_fill_color(29,78,216); pdf.set_text_color(255,255,255); pdf.set_font(fn, 'B', 9)
        col_w = [60, 28, 22, 22, 18, 18]
        for h, w in zip(['Chien dich','Chi tieu','Reach','CPM','Inbox','Lead'], col_w):
            pdf.cell(w, 7, h, fill=True)
        pdf.ln()
        pdf.set_text_color(40,40,60); pdf.set_font(fn, '', 9)
        for i, c in enumerate(campaigns[:12]):
            fill = i % 2 == 0
            pdf.set_fill_color(240,248,255) if fill else pdf.set_fill_color(255,255,255)
            pdf.cell(60, 6, c['name'][:32], fill=True)
            pdf.cell(28, 6, f"{c['spend']:,.0f}", fill=True)
            pdf.cell(22, 6, f"{c['reach']:,}", fill=True)
            pdf.cell(22, 6, f"{c['cpm']:,.0f}", fill=True)
            pdf.cell(18, 6, str(c['messages']), fill=True)
            pdf.cell(18, 6, str(c['leads']), fill=True, ln=True)
        pdf.ln(6)
    # AI text
    pdf.set_font(fn, 'B', 13); pdf.set_text_color(30,30,50)
    pdf.cell(0, 10, 'III. PHAN TICH & DE XUAT', ln=True)
    pdf.set_font(fn, '', 10); pdf.set_text_color(50,50,70)
    pdf.multi_cell(0, 6, _md_to_plain(ai_text)[:4000])
    return bytes(pdf.output())

# ── Scheduler jobs ──────────────────────────────────────────────────

@scheduler.scheduled_job("cron", hour=7, minute=0)
async def daily_ads_check():
    """Kéo FB Ads hàng ngày lúc 7h, cảnh báo nếu CPL vượt ngưỡng."""
    try:
        account_id = os.getenv("FB_AD_ACCOUNT_ID", "")
        if not account_id:
            return
        data = await asyncio.to_thread(_get_fb_data, account_id, "yesterday")
        rows = data.get("insights", {}).get("data", [])
        if not rows:
            return
        total_spend = sum(float(r.get("spend", 0)) for r in rows)
        total_msgs = 0
        for r in rows:
            for a in r.get("actions", []):
                if a["action_type"] == "onsite_conversion.messaging_conversation_started_7d":
                    total_msgs += int(a.get("value", 0))
        cp_msg = total_spend / total_msgs if total_msgs else 0
        if cp_msg > 350000:
            db_add_alert(
                "cpl_alert",
                f"⚠️ CPL cao hôm qua: {cp_msg:,.0f} VND/tin nhắn",
                f"Chi tiêu: {total_spend:,.0f} VND | Tin nhắn: {total_msgs}\n"
                f"CPL vượt ngưỡng 350K — cần kiểm tra và tối ưu ngay."
            )
        else:
            db_add_alert(
                "daily_summary",
                f"✅ Báo cáo ngày: CPL {cp_msg:,.0f} VND",
                f"Chi tiêu hôm qua: {total_spend:,.0f} VND | Tin nhắn: {total_msgs} | CPL: {cp_msg:,.0f} VND"
            )
    except Exception:
        pass

@scheduler.scheduled_job("cron", day_of_week="mon", hour=8, minute=0)
async def weekly_report():
    """Tóm tắt tuần vào thứ 2 lúc 8h."""
    try:
        account_id = os.getenv("FB_AD_ACCOUNT_ID", "")
        if not account_id:
            return
        data = await asyncio.to_thread(_get_fb_data, account_id, "last_7d")
        rows = data.get("insights", {}).get("data", [])
        total_spend = sum(float(r.get("spend", 0)) for r in rows)
        total_msgs = 0
        for r in rows:
            for a in r.get("actions", []):
                if a["action_type"] == "onsite_conversion.messaging_conversation_started_7d":
                    total_msgs += int(a.get("value", 0))
        cp_msg = total_spend / total_msgs if total_msgs else 0
        db_add_alert(
            "weekly_report",
            f"📊 Báo cáo tuần: {total_msgs} tin nhắn",
            f"7 ngày qua — Chi tiêu: {total_spend:,.0f} VND | Tin nhắn: {total_msgs} | CPL TB: {cp_msg:,.0f} VND"
        )
    except Exception:
        pass

@app.post("/upload")
async def upload(file: UploadFile = File(...)):
    return JSONResponse(await parse_upload(file))

@app.post("/chat")
async def chat(body: ChatBody):
    history = db_get_history(body.session_id)
    user_msg = body.message
    if body.file_context:
        user_msg = f"{body.message}\n\n[Du lieu file:]\n{body.file_context}"
    messages = history + [{"role": "user", "content": user_msg}]
    return StreamingResponse(agent_stream(messages, body.session_id, user_msg),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})

# Image
@app.post("/image/enhance-prompt")
async def image_enhance_prompt(body: ImageEnhanceBody):
    system = "You are an expert AI image prompt engineer for interior design visualization. Convert Vietnamese descriptions into detailed English prompts for Flux image generation."
    user_msg = (
        f"Style: {body.style}\n"
        f"Vietnamese description: {body.description}\n\n"
        "Write a detailed English image generation prompt (max 120 words). "
        "Include: room type, design style, materials, colors, lighting, mood, camera angle. "
        "Format: photorealistic architectural visualization, ultra-detailed, 8K. "
        "Output ONLY the prompt, no explanation."
    )
    try:
        if body.ai_model.startswith('claude-'):
            resp = get_anthropic_client().messages.create(
                model=body.ai_model, max_tokens=200,
                system=system, messages=[{"role":"user","content":user_msg}])
            prompt = resp.content[0].text.strip()
        else:
            if not _gemini_pool:
                raise ValueError("Chưa cấu hình GEMINI_API_KEY trong .env")
            resp = await get_gemini_client().aio.models.generate_content(
                model=body.ai_model, contents=user_msg,
                config=genai_types.GenerateContentConfig(system_instruction=system, max_output_tokens=200))
            prompt = resp.text.strip()
        return {"prompt": prompt}
    except Exception as e:
        raise HTTPException(500, str(e)[:300])

@app.post("/generate-image")
async def generate_image(body: ImageBody):
    try:
        enhanced = await asyncio.to_thread(_enhance_prompt, body.description, body.style)
        result = await asyncio.to_thread(_call_fal, enhanced)
        return JSONResponse({"url": result["images"][0]["url"], "prompt": enhanced})
    except ValueError as e:
        raise HTTPException(400, str(e))
    except Exception as e:
        raise HTTPException(500, f"Loi tao anh: {str(e)[:200]}")

# Content Studio
@app.post("/content/generate")
async def content_generate(body: ContentBody):
    return StreamingResponse(content_stream(body.model_dump()),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})

# Video Script
@app.post("/script/generate")
async def script_generate(body: ScriptBody):
    return StreamingResponse(script_stream(body.model_dump()),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})

# Intake / Brief Client
@app.post("/intake/generate")
async def intake_generate(body: IntakeBody):
    return StreamingResponse(intake_stream(body.model_dump()),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})

# Quote Generator
@app.post("/quote/generate")
async def quote_generate(body: QuoteBody):
    return StreamingResponse(quote_stream(body.model_dump()),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})

@app.post("/quote/export")
async def quote_export(body: QuoteBody):
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    import datetime, io as _io

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Bao Gia"
    ws.sheet_view.showGridLines = False

    def _fill(c): return PatternFill("solid", fgColor=c)
    def _font(bold=False, size=11, color="1A1A2E", italic=False):
        return Font(name="Calibri", bold=bold, size=size, color=color, italic=italic)
    def _align(h="left", v="center", wrap=False):
        return Alignment(horizontal=h, vertical=v, wrap_text=wrap)
    def _border():
        s = Side(style="thin", color="CCCCCC")
        return Border(left=s, right=s, top=s, bottom=s)
    def sc(r, c, val, bold=False, size=11, color="1A1A2E", bg=None, h="left", wrap=False):
        cell = ws.cell(row=r, column=c, value=val)
        cell.font = _font(bold=bold, size=size, color=color)
        cell.alignment = _align(h=h, wrap=wrap)
        if bg: cell.fill = _fill(bg)
        cell.border = _border()
        return cell

    # Column widths
    for i, w in enumerate([2,28,40,20,20,20]):
        ws.column_dimensions[get_column_letter(i+1)].width = w

    today = datetime.date.today().strftime("%d/%m/%Y")
    ngan_sach = body.ngan_sach_ads or "Theo thoa thuan"

    # ── HEADER ──
    ws.row_dimensions[1].height = 8
    ws.row_dimensions[2].height = 50
    ws.merge_cells("B2:F2")
    sc(2,2,"BÁO GIÁ DỊCH VỤ FACEBOOK ADS",bold=True,size=18,color="FFFFFF",bg="1A1A2E",h="center")
    ws.row_dimensions[3].height = 24
    ws.merge_cells("B3:F3")
    sc(3,2,f"ĐỘC MEDIA  |  Nguyễn Thanh Hải  |  Ngày: {today}",bold=False,size=11,color="A0C4FF",bg="16213E",h="center")

    # ── CLIENT INFO ──
    r = 5
    ws.row_dimensions[r].height = 26
    ws.merge_cells(f"B{r}:F{r}")
    sc(r,2,"THÔNG TIN KHÁCH HÀNG",bold=True,size=12,color="FFFFFF",bg="0F3460",h="center")
    r += 1
    for label, val in [("Doanh nghiệp", body.ten_client),
                       ("Lĩnh vực", body.linh_vuc),
                       ("Mục tiêu chiến dịch", body.muc_tieu),
                       ("Ngân sách quảng cáo", ngan_sach + (" VND/tháng" if ngan_sach != "Theo thoa thuan" else "")),
                       ("Thời gian hợp đồng", body.thoi_gian)]:
        ws.row_dimensions[r].height = 22
        sc(r,2,label,bold=True,size=10,color="1A1A2E",bg="F4F6FA")
        ws.merge_cells(f"C{r}:F{r}")
        sc(r,3,val,bold=False,size=10,color="16213E",bg="FFFFFF")
        r += 1

    # ── SERVICES & PRICING ──
    r += 1
    ws.row_dimensions[r].height = 26
    ws.merge_cells(f"B{r}:F{r}")
    sc(r,2,"BẢNG GIÁ DỊCH VỤ",bold=True,size=12,color="FFFFFF",bg="0F3460",h="center")
    r += 1
    ws.row_dimensions[r].height = 24
    for i,(h,bg) in enumerate(zip(["Hạng mục dịch vụ","Mô tả chi tiết","Đơn giá/tháng","Số tháng","Thành tiền"],
                                   ["16213E"]*5)):
        sc(r,i+2,h,bold=True,size=10,color="FFFFFF",bg=bg,h="center")
    r += 1

    services = [
        ("Thiết lập chiến dịch (Setup)", "Cài đặt Pixel, cấu trúc Campaign/AdSet/Ad, tạo Audiences, creatives ban đầu", "1.500.000", "1 lần", "1.500.000"),
        ("Quản lý & tối ưu hàng ngày", "Check sáng, điều chỉnh bid, pause/scale adsets, A/B test creative", "3.000.000", body.thoi_gian, ""),
        ("Sản xuất nội dung", "Viết copy bài viết, brief thiết kế hình ảnh, CTA", "2.000.000", body.thoi_gian, ""),
        ("Báo cáo hàng tuần", "Báo cáo số liệu thực tế, nhận xét hiệu quả, khuyến nghị", "500.000", body.thoi_gian, ""),
        ("Tư vấn chiến lược", "Targeting, phân bổ ngân sách, cấu trúc chiến dịch", "Bao gồm", "-", "-"),
    ]
    svc_colors = ["E8F4FD","F5EEF8","E8F8F5","FFF8E7","F4F6FA"]
    for svc, rbg in zip(services, svc_colors):
        ws.row_dimensions[r].height = 36
        sc(r,2,svc[0],bold=True,size=9,color="1A1A2E",bg=rbg)
        sc(r,3,svc[1],bold=False,size=9,color="16213E",bg=rbg,wrap=True)
        sc(r,4,svc[2],bold=False,size=9,color="2980B9",bg=rbg,h="center")
        sc(r,5,svc[3],bold=False,size=9,color="16213E",bg=rbg,h="center")
        sc(r,6,svc[4],bold=True,size=9,color="2980B9",bg=rbg,h="center")
        r += 1

    # Total
    ws.row_dimensions[r].height = 28
    ws.merge_cells(f"B{r}:E{r}")
    sc(r,2,"TỔNG PHÍ DỊCH VỤ/THÁNG (chưa bao gồm ngân sách quảng cáo)",bold=True,size=10,color="FFFFFF",bg="1A1A2E",h="right")
    sc(r,6,"5.500.000 VND",bold=True,size=12,color="F1C40F",bg="1A1A2E",h="center")
    r += 1
    ws.row_dimensions[r].height = 28
    ws.merge_cells(f"B{r}:E{r}")
    sc(r,2,"NGÂN SÁCH QUẢNG CÁO (trả trực tiếp cho Facebook)",bold=True,size=10,color="FFFFFF",bg="16213E",h="right")
    sc(r,6,ngan_sach + (" VND/tháng" if ngan_sach != "Theo thoa thuan" else ""),bold=True,size=12,color="A0C4FF",bg="16213E",h="center")

    # ── KPI CAM KẾT ──
    r += 2
    ws.row_dimensions[r].height = 26
    ws.merge_cells(f"B{r}:F{r}")
    sc(r,2,"KPI CAM KẾT (dựa trên benchmark thực tế ngành nội thất VN)",bold=True,size=12,color="FFFFFF",bg="0F3460",h="center")
    r += 1
    kpis = [
        ("CPM Tin nhắn","150.000 – 230.000 VND"),
        ("CPM Tương tác","21.000 – 34.000 VND"),
        ("CPL (chi phí/inbox)","166.000 – 315.000 VND | Tốt: <230K"),
        ("CTR Inbox","1.3% – 2.1%"),
        ("Leads dự kiến","10 – 18 inbox / 10 ngày (với 3 triệu ngân sách)"),
        ("Báo cáo","Hàng tuần vào thứ 2 trước 10h sáng"),
    ]
    for kpi_label, kpi_val in kpis:
        ws.row_dimensions[r].height = 22
        sc(r,2,kpi_label,bold=True,size=10,color="1A1A2E",bg="F4F6FA")
        ws.merge_cells(f"C{r}:F{r}")
        sc(r,3,kpi_val,bold=False,size=10,color="2980B9",bg="E8F4FD")
        r += 1

    # ── ĐIỀU KHOẢN ──
    r += 1
    ws.row_dimensions[r].height = 26
    ws.merge_cells(f"B{r}:F{r}")
    sc(r,2,"ĐIỀU KHOẢN & THANH TOÁN",bold=True,size=12,color="FFFFFF",bg="0F3460",h="center")
    r += 1
    terms = [
        ("Thanh toán","Đầu mỗi tháng, chuyển khoản trước khi bắt đầu"),
        ("Ngân sách ads","Khách hàng nạp trực tiếp vào tài khoản Facebook"),
        ("Thời hạn báo giá","Báo giá có hiệu lực trong 7 ngày kể từ ngày phát hành"),
        ("Điều chỉnh","Có thể điều chỉnh dịch vụ sau tháng đầu tiên theo kết quả thực tế"),
        ("Cam kết","Độc Media cam kết báo cáo minh bạch 100% số liệu thực từ ADS Manager"),
    ]
    for tl, tv in terms:
        ws.row_dimensions[r].height = 28
        sc(r,2,tl,bold=True,size=10,color="1A1A2E",bg="ECF0F1")
        ws.merge_cells(f"C{r}:F{r}")
        sc(r,3,tv,bold=False,size=10,color="16213E",bg="FFFFFF",wrap=True)
        r += 1

    # ── FOOTER ──
    r += 2
    ws.row_dimensions[r].height = 44
    ws.merge_cells(f"B{r}:F{r}")
    sc(r,2,"ĐỘC MEDIA — Nguyễn Thanh Hải\ndocmedia.hcm@gmail.com | Chuyên ngành Nội thất & Kiến trúc",
       bold=True,size=11,color="FFFFFF",bg="1A1A2E",h="center",wrap=True)

    buf = _io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    filename = f"BaoGia_{body.ten_client.replace(' ','_')}_{today.replace('/','')}.xlsx"
    from fastapi.responses import Response
    return Response(
        content=buf.read(),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{urllib.parse.quote(filename)}"}
    )

# Facebook Ads
@app.get("/ads/accounts")
async def get_ad_accounts():
    token = os.getenv("FB_ACCESS_TOKEN", "")
    if not token or token.startswith("your_"):
        raise HTTPException(400, "FB_ACCESS_TOKEN chua duoc cai dat trong .env")
    try:
        accounts = await asyncio.to_thread(_get_fb_accounts, token)
        return JSONResponse(accounts)
    except ValueError as e:
        raise HTTPException(400, str(e))

@app.get("/ads/data")
async def ads_data(account_id: str = "", period: str = "last_7d",
                   date_from: str = "", date_to: str = ""):
    try:
        data = await asyncio.to_thread(_get_fb_data, account_id, period, date_from, date_to)
        return JSONResponse(data)
    except ValueError as e:
        raise HTTPException(400, str(e))
    except Exception as e:
        raise HTTPException(500, str(e)[:200])

@app.post("/ads/analyze")
async def ads_analyze(body: AdsAnalyzeBody):
    return StreamingResponse(ads_analyze_stream(body.data, body.ai_model),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})

# ── HTML ───────────────────────────────────────────────────────────

HTML = """<!DOCTYPE html>
<html lang="vi">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>Doc Media AI Agent</title>
  <style>
    /* ── Design tokens ── */
    :root{
      --bg:#0b0b18;--surface:#13132a;--surface-2:#1a1a35;--surface-3:#22223f;
      --border:#2a2a4a;--border-2:#343458;
      --text:#ddddf0;--text-2:#8888b8;--text-3:#55558a;
      --accent:#7c3aed;--accent-h:#6d28d9;--accent-glow:rgba(124,58,237,.18);
      --accent-soft:rgba(124,58,237,.12);
      --red:#ef4444;--red-soft:rgba(239,68,68,.15);
      --orange:#f97316;--orange-soft:rgba(249,115,22,.12);
      --green:#22c55e;--green-soft:rgba(34,197,94,.12);
      --scrollbar:#252545;
    }

    /* ── Reset ── */
    *{box-sizing:border-box;margin:0;padding:0}
    body{font-family:-apple-system,BlinkMacSystemFont,"Segoe UI",sans-serif;background:var(--bg);color:var(--text);height:100vh;display:flex;flex-direction:column;overflow:hidden}
    ::-webkit-scrollbar{width:5px;height:5px}
    ::-webkit-scrollbar-track{background:transparent}
    ::-webkit-scrollbar-thumb{background:var(--scrollbar);border-radius:3px}
    ::-webkit-scrollbar-thumb:hover{background:var(--border-2)}

    /* ── Header ── */
    header{background:var(--surface);border-bottom:1px solid var(--border);height:54px;padding:0 20px;display:flex;align-items:center;gap:16px;flex-shrink:0;z-index:10}
    .bell-wrap{position:relative;flex-shrink:0}
    .bell-btn{background:none;border:1px solid var(--border);color:var(--text-2);width:36px;height:36px;border-radius:9px;cursor:pointer;font-size:16px;display:flex;align-items:center;justify-content:center;transition:all .15s}
    .bell-btn:hover{border-color:var(--accent);color:var(--text)}
    .bell-badge{position:absolute;top:-5px;right:-5px;background:var(--red);color:#fff;font-size:10px;font-weight:700;width:18px;height:18px;border-radius:50%;display:flex;align-items:center;justify-content:center;display:none}
    .notif-panel{position:absolute;top:44px;right:0;width:340px;background:var(--surface);border:1px solid var(--border);border-radius:12px;box-shadow:0 8px 32px rgba(0,0,0,.4);z-index:100;display:none;max-height:420px;overflow:hidden;flex-direction:column}
    .notif-panel.open{display:flex}
    .notif-header{padding:12px 16px;border-bottom:1px solid var(--border);display:flex;align-items:center;justify-content:space-between;flex-shrink:0}
    .notif-header strong{font-size:13px}
    .notif-read-all{font-size:11px;color:var(--accent);background:none;border:none;cursor:pointer;padding:0}
    .notif-list{overflow-y:auto;flex:1}
    .notif-item{padding:12px 16px;border-bottom:1px solid var(--border);cursor:pointer;transition:background .1s}
    .notif-item:hover{background:var(--surface-2)}
    .notif-item.unread{background:var(--accent-soft)}
    .notif-title{font-size:12.5px;font-weight:600;color:var(--text);margin-bottom:3px}
    .notif-body{font-size:11.5px;color:var(--text-2);line-height:1.5;white-space:pre-line}
    .notif-time{font-size:10px;color:var(--text-3);margin-top:4px}
    .notif-empty{padding:32px 16px;text-align:center;color:var(--text-3);font-size:13px}
    .tool-bubble{background:var(--surface-2);border:1px solid var(--border);border-radius:10px;padding:8px 14px;font-size:12.5px;color:var(--text-2);display:flex;align-items:center;gap:8px;margin-bottom:6px}
    .tool-spinner{width:14px;height:14px;border:2px solid var(--border-2);border-top-color:var(--accent);border-radius:50%;animation:spin .7s linear infinite;flex-shrink:0}
    @keyframes spin{to{transform:rotate(360deg)}}

    /* ── LIBRARY TAB ── */
    #tab-library{flex-direction:column;overflow:hidden}
    .lib-header{display:flex;align-items:center;gap:10px;padding:14px 20px;background:var(--surface);border-bottom:1px solid var(--border);flex-shrink:0}
    .lib-header-title{font-size:15px;font-weight:700;color:var(--text);flex:1}
    .lib-export-btns{display:flex;gap:6px}
    .lib-exp-btn{background:var(--surface-2);border:1px solid var(--border);color:var(--text-2);border-radius:8px;padding:6px 12px;font-size:12px;cursor:pointer;transition:all .15s;font-weight:500}
    .lib-exp-btn:hover{border-color:var(--accent);color:var(--accent)}
    .lib-toolbar{display:flex;align-items:center;gap:10px;padding:10px 20px;background:var(--surface);border-bottom:1px solid var(--border);flex-shrink:0;flex-wrap:wrap}
    #lib-search{flex:1;min-width:160px;border:1px solid var(--border);border-radius:8px;padding:7px 12px;font-size:13px;outline:none;background:var(--surface-2);color:var(--text)}
    #lib-search:focus{border-color:var(--accent)}
    #lib-count{font-size:12px;color:var(--text-3);white-space:nowrap}
    .lib-list{flex:1;overflow-y:auto;padding:16px 20px;display:flex;flex-direction:column;gap:12px}
    .lib-item{background:var(--surface);border:1px solid var(--border);border-radius:12px;overflow:hidden;transition:border-color .15s}
    .lib-item:hover{border-color:var(--accent-soft)}
    .lib-item-header{display:flex;align-items:center;gap:8px;padding:11px 14px;background:var(--surface-2);flex-wrap:wrap}
    .lib-badge{font-size:11px;font-weight:600;padding:2px 8px;border-radius:5px;flex-shrink:0}
    .lib-item-title{font-size:13px;font-weight:600;color:var(--text);flex:1;min-width:0;overflow:hidden;text-overflow:ellipsis;white-space:nowrap}
    .lib-client{font-size:11px;color:var(--text-3);flex-shrink:0}
    .lib-date{font-size:11px;color:var(--text-3);flex-shrink:0}
    .lib-del{background:none;border:none;color:var(--text-3);cursor:pointer;font-size:14px;padding:2px 5px;border-radius:4px;flex-shrink:0;transition:color .15s}
    .lib-del:hover{color:#f87171}
    .lib-preview{padding:12px 14px;font-size:13px;color:var(--text-2);line-height:1.7;max-height:120px;overflow:hidden;position:relative}
    .lib-preview.expanded{max-height:none}
    .lib-preview::after{content:'';position:absolute;bottom:0;left:0;right:0;height:32px;background:linear-gradient(transparent,var(--surface));pointer-events:none}
    .lib-preview.expanded::after{display:none}
    .lib-footer{display:flex;gap:8px;padding:8px 14px;border-top:1px solid var(--border)}
    .lib-foot-btn{background:var(--surface-2);border:1px solid var(--border);color:var(--text-2);border-radius:6px;padding:4px 12px;font-size:12px;cursor:pointer;transition:all .15s}
    .lib-foot-btn:hover{border-color:var(--accent);color:var(--accent)}
    .save-lib-bar{display:none;align-items:center;gap:8px;padding:8px 16px;background:var(--surface);border-top:1px solid var(--border)}
    .save-lib-btn{background:none;border:1px solid var(--border);color:var(--text-2);border-radius:8px;padding:6px 14px;font-size:12.5px;cursor:pointer;transition:all .15s;font-weight:500}
    .save-lib-btn:hover{border-color:var(--accent);color:var(--accent);background:var(--accent-glow)}
    .save-lib-status{font-size:12px;color:var(--text-3)}

    /* ── REPORT TAB ── */
    #tab-report{flex-direction:row;overflow:hidden}
    .rpt-form{width:280px;flex-shrink:0;background:var(--surface);border-right:1px solid var(--border);display:flex;flex-direction:column;overflow-y:auto;padding:20px}
    .rpt-form h2{font-size:15px;font-weight:700;color:var(--text);margin:0 0 16px}
    .rpt-form label{font-size:12px;color:var(--text-2);margin-bottom:4px;display:block;font-weight:500}
    .rpt-form select,.rpt-form input[type=text],.rpt-form input[type=date]{width:100%;border:1px solid var(--border);border-radius:8px;padding:8px 10px;font-size:13px;background:var(--surface-2);color:var(--text);outline:none;margin-bottom:12px}
    .rpt-form select:focus,.rpt-form input:focus{border-color:var(--accent)}
    .rpt-custom-dates{display:none;gap:6px;margin-bottom:12px}
    .rpt-custom-dates.show{display:flex}
    .rpt-custom-dates input{flex:1;margin:0}
    #rpt-gen-btn{width:100%;background:var(--accent);color:#fff;border:none;border-radius:10px;padding:11px;font-size:14px;font-weight:600;cursor:pointer;margin-top:4px;transition:background .15s}
    #rpt-gen-btn:hover{background:var(--accent-h)}
    #rpt-gen-btn:disabled{background:var(--surface-3);color:var(--text-3);cursor:not-allowed}
    .rpt-export-row{display:flex;gap:6px;margin-top:12px}
    .rpt-exp-btn{flex:1;background:var(--surface-2);border:1px solid var(--border);color:var(--text-2);border-radius:8px;padding:7px 6px;font-size:12px;cursor:pointer;font-weight:500;transition:all .15s;text-align:center}
    .rpt-exp-btn:hover{border-color:var(--accent);color:var(--accent)}
    .rpt-exp-btn:disabled{opacity:.4;cursor:not-allowed}
    .rpt-preview{flex:1;overflow-y:auto;padding:24px;background:var(--bg)}
    .rpt-placeholder{display:flex;align-items:center;justify-content:center;height:100%;color:var(--text-3);font-size:14px;text-align:center;line-height:1.8}
    .rpt-section{background:var(--surface);border-radius:12px;border:1px solid var(--border);padding:20px;margin-bottom:16px}
    .rpt-section h3{font-size:13.5px;font-weight:700;color:var(--text);margin:0 0 14px;padding-bottom:8px;border-bottom:1px solid var(--border)}
    .rpt-kpi-grid{display:grid;grid-template-columns:repeat(4,1fr);gap:10px}
    .rpt-kpi-card{background:var(--surface-2);border-radius:10px;padding:14px;border:1px solid var(--border);text-align:center}
    .rpt-kpi-label{font-size:11px;color:var(--text-3);margin-bottom:5px}
    .rpt-kpi-value{font-size:16px;font-weight:700;color:var(--accent)}
    .rpt-kpi-sub{font-size:10px;color:var(--text-3);margin-top:2px}
    .rpt-camp-table{width:100%;border-collapse:collapse;font-size:12.5px}
    .rpt-camp-table th{background:var(--surface-3);color:var(--text-2);padding:7px 10px;text-align:left;font-weight:600;font-size:11.5px}
    .rpt-camp-table td{padding:7px 10px;border-bottom:1px solid var(--border);color:var(--text)}
    .rpt-camp-table tr:last-child td{border-bottom:none}
    .rpt-ai-text{font-size:13.5px;line-height:1.85;color:var(--text);white-space:pre-wrap}
    .rpt-ai-text h2{font-size:14px;font-weight:700;color:#a78bfa;margin:14px 0 6px;border-bottom:1px solid rgba(167,139,250,.2);padding-bottom:4px}
    .rpt-status{font-size:12px;color:var(--text-2);margin-top:8px;min-height:18px}
    .logo{font-weight:800;font-size:15px;white-space:nowrap;margin-right:8px;background:linear-gradient(135deg,#a78bfa,#7c3aed);-webkit-background-clip:text;-webkit-text-fill-color:transparent;letter-spacing:.3px}
    .tab-bar{display:flex;gap:1px;flex:1;overflow:hidden}
    .tab-btn{background:none;border:none;color:var(--text-2);padding:8px 14px;cursor:pointer;font-size:13px;font-weight:500;border-radius:7px;transition:all .15s;white-space:nowrap;position:relative}
    .tab-btn:hover{color:var(--text);background:var(--surface-2)}
    .tab-btn.active{color:#fff;background:var(--accent-soft)}
    .tab-btn.active::after{content:'';position:absolute;bottom:-1px;left:50%;transform:translateX(-50%);width:28px;height:2px;background:var(--accent);border-radius:2px}
    .tab-btn[style*="E67E22"]{background:var(--orange-soft) !important;color:#fb923c !important;border:1px solid rgba(249,115,22,.25) !important}
    .tab-btn[style*="E67E22"].active,.tab-btn[style*="E67E22"]:hover{background:rgba(249,115,22,.2) !important;color:#fdba74 !important}
    .tab-btn[style*="27AE60"]{background:var(--green-soft) !important;color:#4ade80 !important;border:1px solid rgba(34,197,94,.25) !important}
    .tab-btn[style*="27AE60"].active,.tab-btn[style*="27AE60"]:hover{background:rgba(34,197,94,.2) !important;color:#86efac !important}
    .tab-btn[style*="8E44AD"]{background:var(--accent-soft) !important;color:#a78bfa !important;border:1px solid rgba(124,58,237,.25) !important}
    .tab-btn[style*="8E44AD"].active,.tab-btn[style*="8E44AD"]:hover{background:rgba(124,58,237,.22) !important;color:#c4b5fd !important}
    .btn-new{background:var(--accent);border:none;color:#fff;padding:7px 14px;border-radius:8px;cursor:pointer;font-size:13px;font-weight:600;white-space:nowrap;flex-shrink:0;transition:background .15s}
    .btn-new:hover{background:var(--accent-h)}

    /* ── Tab content ── */
    .tab-content{display:none;flex:1;overflow:hidden}
    .tab-content.active{display:flex}

    /* ── CHAT TAB ── */
    #tab-chat{flex-direction:row}
    #sidebar{width:236px;background:var(--surface);display:flex;flex-direction:column;flex-shrink:0;border-right:1px solid var(--border)}
    #sidebar-hdr{padding:12px 16px;color:var(--text-3);font-size:10.5px;font-weight:700;text-transform:uppercase;letter-spacing:1px;border-bottom:1px solid var(--border)}
    #session-list{flex:1;overflow-y:auto;padding:6px 0}
    .sess-item{padding:9px 16px;cursor:pointer;border-left:2px solid transparent;transition:all .15s;display:flex;flex-wrap:wrap;align-items:center;gap:2px}
    .sess-item:hover{background:var(--surface-2)}
    .sess-item.active{background:var(--accent-soft);border-left-color:var(--accent)}
    .sess-del{display:none;background:none;border:none;color:var(--red);cursor:pointer;font-size:12px;padding:2px 5px;flex-shrink:0;border-radius:4px;line-height:1;opacity:.7}
    .sess-del:hover{background:var(--red-soft);opacity:1}
    .sess-item:hover .sess-del{display:block}
    .sess-prev{color:var(--text);font-size:12.5px;white-space:nowrap;overflow:hidden;text-overflow:ellipsis;font-weight:500}
    .sess-time{color:var(--text-3);font-size:10.5px;margin-top:2px}
    #chat-main{flex:1;display:flex;flex-direction:column;overflow:hidden;background:var(--bg)}
    #chat-box{flex:1;overflow-y:auto;padding:24px 28px;display:flex;flex-direction:column;gap:14px}
    .msg{max-width:76%;padding:12px 16px;border-radius:14px;line-height:1.7;font-size:14px;word-wrap:break-word}
    .msg.user{background:var(--accent);color:#fff;align-self:flex-end;border-bottom-right-radius:4px;font-weight:500}
    .msg.assistant{background:var(--surface);color:var(--text);align-self:flex-start;border-bottom-left-radius:4px;border:1px solid var(--border)}
    .msg.system{background:var(--surface-2);color:var(--text-2);align-self:center;font-size:13px;border-radius:8px;max-width:90%;border:1px solid var(--border)}
    .msg.assistant strong{font-weight:600;color:#e2e2f8}
    .msg.assistant code{background:var(--surface-2);padding:2px 7px;border-radius:5px;font-family:monospace;font-size:13px;color:#a78bfa;border:1px solid var(--border)}
    .msg.assistant h3{font-size:14.5px;margin-bottom:6px;color:#c4b5fd}
    .msg.assistant ul{padding-left:18px;margin:6px 0}
    .msg.assistant li{margin-bottom:4px}
    .streaming::after{content:"▋";animation:blink .7s infinite;color:var(--accent)}
    @keyframes blink{0%,100%{opacity:1}50%{opacity:0}}
    #file-banner{background:rgba(99,102,241,.1);border-top:1px solid rgba(99,102,241,.2);padding:8px 24px;display:none;align-items:center;gap:10px;font-size:13px;color:#818cf8;flex-shrink:0}
    #file-banner span{flex:1}

    /* Chat input */
    #chat-input-area{background:var(--surface);border-top:1px solid var(--border);padding:14px 24px;flex-shrink:0}
    .qprompts{display:flex;gap:6px;flex-wrap:wrap;margin-bottom:10px}
    .qbtn{background:var(--surface-2);border:1px solid var(--border);border-radius:20px;padding:5px 13px;font-size:12px;cursor:pointer;color:var(--text-2);transition:all .15s;font-weight:500}
    .qbtn:hover{background:var(--surface-3);color:var(--text);border-color:var(--border-2)}
    #input-wrap{display:flex;gap:10px;align-items:flex-end}
    .icon-btn{background:var(--surface-2);border:1px solid var(--border);border-radius:8px;padding:0 12px;height:44px;display:flex;align-items:center;cursor:pointer;font-size:20px;color:var(--text-2);flex-shrink:0;transition:all .15s}
    .icon-btn:hover{background:var(--surface-3);color:var(--text)}
    #file-input{display:none}
    #user-input{flex:1;border:1px solid var(--border);border-radius:10px;padding:11px 14px;font-size:14px;resize:none;outline:none;min-height:44px;max-height:120px;font-family:inherit;background:var(--surface-2);color:var(--text)}
    #user-input::placeholder{color:var(--text-3)}
    #user-input:focus{border-color:var(--accent);box-shadow:0 0 0 3px var(--accent-glow)}
    #send-btn,#gen-btn,#gen-script-btn,#refresh-ads-btn{background:var(--accent);color:#fff;border:none;border-radius:10px;padding:0 20px;height:44px;cursor:pointer;font-size:14px;font-weight:600;flex-shrink:0;transition:background .15s}
    #send-btn:hover,#gen-btn:hover,#gen-script-btn:hover,#refresh-ads-btn:hover{background:var(--accent-h)}
    #send-btn:disabled,#gen-btn:disabled,#gen-script-btn:disabled{background:var(--surface-3);color:var(--text-3);cursor:not-allowed}

    /* ── STUDIO shared (Content / Script / Ads) ── */
    .studio-wrap{flex:1;display:flex;overflow:hidden;background:var(--bg)}
    .studio-col{flex:1;display:flex;flex-direction:column;overflow:hidden}
    .studio-form{width:320px;background:var(--surface);border-right:1px solid var(--border);padding:20px;overflow-y:auto;display:flex;flex-direction:column;gap:14px;flex-shrink:0}
    .studio-form h2{font-size:15px;color:var(--text);font-weight:700;margin-bottom:2px}
    .studio-form label,.form-label{font-size:11px;color:var(--text-2);font-weight:600;text-transform:uppercase;letter-spacing:.7px;display:block;margin-bottom:5px}
    .studio-form input,.studio-form textarea,.studio-form select{width:100%;border:1px solid var(--border);border-radius:8px;padding:9px 12px;font-size:13.5px;outline:none;font-family:inherit;background:var(--surface-2);color:var(--text)}
    .studio-form input::placeholder,.studio-form textarea::placeholder{color:var(--text-3)}
    .studio-form input:focus,.studio-form textarea:focus,.studio-form select:focus{border-color:var(--accent);box-shadow:0 0 0 3px var(--accent-glow)}
    .studio-form textarea{resize:vertical;min-height:72px}
    .seg-btns{display:flex;gap:6px;flex-wrap:wrap}
    .seg{background:var(--surface-2);border:1px solid var(--border);border-radius:8px;padding:6px 13px;font-size:12.5px;cursor:pointer;color:var(--text-2);transition:all .15s;font-weight:500}
    .seg:hover{background:var(--surface-3);color:var(--text)}
    .seg.active{background:var(--accent);color:#fff;border-color:var(--accent)}
    .action-row{display:flex;gap:10px;align-items:center;margin-top:4px}
    .action-row button{background:var(--accent);color:#fff;border:none;border-radius:8px;padding:10px 20px;font-size:13.5px;cursor:pointer;font-weight:600;transition:background .15s}
    .action-row button:hover{background:var(--accent-h)}
    .action-row button:disabled{background:var(--surface-3);color:var(--text-3);cursor:not-allowed}
    .action-row span{font-size:13px;color:var(--text-2)}
    .studio-output{flex:1;padding:24px;overflow-y:auto;display:flex;flex-direction:column;gap:16px;background:var(--bg)}
    .output-card{background:var(--surface);border-radius:12px;padding:20px;border:1px solid var(--border);white-space:pre-wrap;line-height:1.75;font-size:14px;color:var(--text)}
    .output-card strong{font-weight:600;color:#e2e2f8}
    .output-card h2,.output-card h3{color:#a78bfa;margin-bottom:8px}
    .copy-bar{display:flex;justify-content:flex-end;margin-bottom:10px}
    .copy-btn{background:var(--surface-2);border:1px solid var(--border);border-radius:6px;padding:5px 12px;font-size:12px;cursor:pointer;color:var(--text-2);transition:all .15s}
    .copy-btn:hover{background:var(--surface-3);color:var(--text)}
    .placeholder-msg{color:var(--text-3);font-size:14px;text-align:center;margin-top:80px}

    /* ── IMAGE TAB ── */
    #tab-image .studio-wrap{padding:28px 36px;gap:32px;align-items:flex-start}
    #tab-image .studio-col{flex:1;flex-direction:column;gap:0;height:100%}
    #img-desc{border:1px solid var(--border);border-radius:10px;padding:11px 14px;font-size:14px;resize:vertical;outline:none;min-height:100px;font-family:inherit;background:var(--surface-2);color:var(--text);width:100%;box-sizing:border-box}
    #img-desc::placeholder{color:var(--text-3)}
    #img-desc:focus{border-color:var(--accent);box-shadow:0 0 0 3px var(--accent-glow)}
    .style-chips{display:flex;gap:8px;flex-wrap:wrap}
    .schip{background:var(--surface-2);border:1px solid var(--border);border-radius:20px;padding:5px 15px;font-size:12.5px;cursor:pointer;color:var(--text-2);transition:all .15s;font-weight:500}
    .schip:hover:not(.active){background:var(--surface-3);color:var(--text);border-color:var(--border-2)}
    .schip.active{background:var(--accent);color:#fff;border-color:var(--accent);box-shadow:0 0 12px var(--accent-glow)}
    #gen-img-el{width:100%;border-radius:12px;box-shadow:0 4px 24px rgba(0,0,0,.4);object-fit:cover}

    /* ── ADS TAB ── */
    #tab-ads .studio-wrap{flex-direction:column}
    .ads-controls{background:var(--surface);border-bottom:1px solid var(--border);padding:12px 24px;display:flex;align-items:center;gap:10px;flex-shrink:0;flex-wrap:wrap}
    .ads-controls strong{font-size:13.5px;color:var(--text);white-space:nowrap;font-weight:600}
    #ads-account-sel{border:1px solid var(--border);border-radius:8px;padding:6px 10px;font-size:13px;outline:none;background:var(--surface-2);color:var(--text);cursor:pointer;max-width:260px}
    .period-seg{display:flex;gap:4px;flex-wrap:wrap}
    .period-seg button{background:var(--surface-2);border:1px solid var(--border);border-radius:7px;padding:5px 12px;font-size:12.5px;cursor:pointer;color:var(--text-2);transition:all .15s;white-space:nowrap;font-weight:500}
    .period-seg button:hover{background:var(--surface-3);color:var(--text)}
    .period-seg button.active{background:var(--accent);color:#fff;border-color:var(--accent)}
    #custom-dates{display:none;align-items:center;gap:6px}
    #custom-dates input{border:1px solid var(--border);border-radius:7px;padding:5px 10px;font-size:13px;outline:none;background:var(--surface-2);color:var(--text)}
    #refresh-ads-btn{height:34px;padding:0 16px;font-size:13px;white-space:nowrap;border-radius:8px}
    .ads-body{flex:1;overflow-y:auto;padding:20px 24px;display:flex;flex-direction:column;gap:16px;background:var(--bg)}
    .kpi-row{display:grid;grid-template-columns:repeat(4,1fr);gap:12px}
    .kpi-card{background:var(--surface);border-radius:12px;padding:16px;border:1px solid var(--border);text-align:center}
    .kpi-val{font-size:22px;font-weight:700;color:var(--text);background:linear-gradient(135deg,#a78bfa,#7c3aed);-webkit-background-clip:text;-webkit-text-fill-color:transparent}
    .kpi-lbl{font-size:12px;color:var(--text-2);margin-top:4px}
    .kpi-sub{font-size:11px;color:var(--text-3);margin-top:2px}
    .camp-table{background:var(--surface);border-radius:12px;overflow:hidden;border:1px solid var(--border)}
    .camp-table table{width:100%;border-collapse:collapse}
    .camp-table th{background:var(--surface-2);padding:10px 14px;text-align:left;font-size:11.5px;color:var(--text-2);font-weight:700;text-transform:uppercase;letter-spacing:.5px;border-bottom:1px solid var(--border)}
    .camp-table td{padding:10px 14px;font-size:13.5px;border-bottom:1px solid var(--border);color:var(--text)}
    .camp-table tr:last-child td{border-bottom:none}
    .camp-table tr:hover td{background:var(--surface-2)}
    .badge{padding:3px 9px;border-radius:5px;font-size:11px;font-weight:700;letter-spacing:.3px}
    .badge.active{background:var(--green-soft);color:#4ade80}
    .badge.paused{background:var(--orange-soft);color:#fb923c}
    .ai-section{background:var(--surface);border-radius:12px;padding:20px;border:1px solid var(--border)}
    .ai-section h3{font-size:14.5px;color:var(--text);margin-bottom:12px;display:flex;align-items:center;gap:8px;font-weight:700}
    #ai-analysis{font-size:14px;line-height:1.75;color:var(--text);white-space:pre-wrap;min-height:40px}
    #ai-analysis strong{font-weight:600;color:#e2e2f8}
    .no-data{color:var(--text-3);text-align:center;padding:40px;font-size:14px}

    /* ── INTAKE TAB ── */
    .intake-wrap{display:flex;height:100%;overflow:hidden}
    .intake-form-panel{width:340px;min-width:340px;background:var(--surface);border-right:1px solid var(--border);padding:16px 18px;overflow-y:auto;display:flex;flex-direction:column;gap:10px}
    .intake-section-title{font-size:10.5px;font-weight:700;color:var(--orange);text-transform:uppercase;letter-spacing:.8px;margin:6px 0 2px;padding-bottom:5px;border-bottom:1px solid rgba(249,115,22,.25)}
    .intake-form-panel label{font-size:11.5px;font-weight:600;color:var(--text-2);display:block;margin-bottom:3px}
    .intake-form-panel input,.intake-form-panel select,.intake-form-panel textarea{width:100%;border:1px solid var(--border);border-radius:8px;padding:7px 10px;font-size:13px;outline:none;background:var(--surface-2);color:var(--text);box-sizing:border-box}
    .intake-form-panel input::placeholder,.intake-form-panel textarea::placeholder{color:var(--text-3)}
    .intake-form-panel textarea{resize:vertical;min-height:56px}
    .intake-form-panel input:focus,.intake-form-panel select:focus,.intake-form-panel textarea:focus{border-color:var(--orange);box-shadow:0 0 0 3px var(--orange-soft)}
    .intake-form-panel .row2{display:grid;grid-template-columns:1fr 1fr;gap:8px}
    .intake-output{flex:1;display:flex;flex-direction:column;overflow:hidden}
    .intake-output-header{padding:14px 20px;border-bottom:1px solid var(--border);display:flex;align-items:center;gap:10px;background:var(--surface)}
    .intake-output-header strong{font-size:13.5px;color:var(--text);flex:1;font-weight:600}
    .intake-output-body{flex:1;overflow-y:auto;padding:24px;font-size:14px;line-height:1.8;color:var(--text);background:var(--bg)}
    .intake-output-body h1{font-size:17px;font-weight:700;color:var(--text);margin:0 0 8px}
    .intake-output-body h2{font-size:14px;font-weight:700;color:var(--orange);margin:18px 0 6px;border-bottom:1px solid rgba(249,115,22,.2);padding-bottom:5px}
    .intake-output-body h3{font-size:13px;font-weight:600;color:var(--text-2);margin:10px 0 3px}
    .intake-output-body strong{font-weight:700;color:#e2e2f8}
    .intake-output-body ul{padding-left:20px;margin:4px 0}
    .intake-output-body li{margin-bottom:3px}
    .intake-output-body table{width:100%;border-collapse:collapse;margin:10px 0;font-size:13px}
    .intake-output-body th{background:rgba(249,115,22,.18);color:#fb923c;padding:8px 12px;text-align:left;font-weight:700;font-size:12px}
    .intake-output-body td{padding:7px 12px;border-bottom:1px solid var(--border);color:var(--text)}
    .intake-output-body tr:nth-child(even) td{background:var(--surface-2)}
    .btn-orange{background:var(--orange);color:#fff;border:none;border-radius:8px;padding:9px 18px;font-size:13px;font-weight:600;cursor:pointer;white-space:nowrap;transition:opacity .15s}
    .btn-orange:hover{opacity:.88}
    .btn-orange:disabled{background:var(--surface-3);color:var(--text-3);cursor:not-allowed}
    .btn-outline-orange{background:transparent;color:var(--orange);border:1.5px solid var(--orange);border-radius:8px;padding:7px 14px;font-size:13px;font-weight:600;cursor:pointer;white-space:nowrap;transition:all .15s}
    .btn-outline-orange:hover{background:var(--orange-soft)}
    .btn-outline-orange:disabled{color:var(--text-3);border-color:var(--border);cursor:not-allowed}

    /* ── QUOTE TAB ── */
    .quote-wrap{display:flex;height:100%;overflow:hidden}
    .quote-form-panel{width:300px;min-width:300px;background:var(--surface);border-right:1px solid var(--border);padding:20px 18px;overflow-y:auto;display:flex;flex-direction:column;gap:14px}
    .quote-form-panel .form-title{font-size:14.5px;font-weight:700;color:var(--text);margin-bottom:4px}
    .quote-form-panel label{font-size:11.5px;font-weight:600;color:var(--text-2);display:block;margin-bottom:4px}
    .quote-form-panel input,.quote-form-panel select,.quote-form-panel textarea{width:100%;border:1px solid var(--border);border-radius:8px;padding:8px 10px;font-size:13px;outline:none;background:var(--surface-2);color:var(--text);box-sizing:border-box}
    .quote-form-panel input::placeholder,.quote-form-panel textarea::placeholder{color:var(--text-3)}
    .quote-form-panel textarea{resize:vertical;min-height:64px}
    .quote-form-panel input:focus,.quote-form-panel select:focus,.quote-form-panel textarea:focus{border-color:var(--green);box-shadow:0 0 0 3px var(--green-soft)}
    .quote-output{flex:1;display:flex;flex-direction:column;overflow:hidden}
    .quote-output-header{padding:14px 20px;border-bottom:1px solid var(--border);display:flex;align-items:center;gap:10px;background:var(--surface)}
    .quote-output-header strong{font-size:13.5px;color:var(--text);flex:1;font-weight:600}
    .quote-output-body{flex:1;overflow-y:auto;padding:24px;font-size:14px;line-height:1.8;color:var(--text);background:var(--bg);white-space:pre-wrap}
    .quote-output-body h1{font-size:17px;font-weight:700;color:var(--text);margin:0 0 8px}
    .quote-output-body h2{font-size:14px;font-weight:700;color:var(--green);margin:16px 0 6px;border-bottom:1px solid rgba(34,197,94,.2);padding-bottom:5px}
    .quote-output-body h3{font-size:13px;font-weight:600;color:#6ee7b7;margin:12px 0 4px}
    .quote-output-body table{width:100%;border-collapse:collapse;margin:10px 0}
    .quote-output-body th{background:rgba(34,197,94,.15);color:#4ade80;padding:8px 12px;font-size:12px;text-align:left;font-weight:700}
    .quote-output-body td{padding:7px 12px;border-bottom:1px solid var(--border);font-size:13px;color:var(--text)}
    .quote-output-body tr:nth-child(even) td{background:var(--surface-2)}
    .btn-green{background:var(--green);color:#fff;border:none;border-radius:8px;padding:9px 18px;font-size:13px;font-weight:600;cursor:pointer;transition:opacity .15s}
    .btn-green:hover{opacity:.88}
    .btn-green:disabled{background:var(--surface-3);color:var(--text-3);cursor:not-allowed}
    .btn-outline-green{background:transparent;color:var(--green);border:1.5px solid var(--green);border-radius:8px;padding:7px 14px;font-size:13px;font-weight:600;cursor:pointer;transition:all .15s}
    .btn-outline-green:hover{background:var(--green-soft)}
    .btn-outline-green:disabled{color:var(--text-3);border-color:var(--border);cursor:not-allowed}
  </style>
</head>
<body>

<header>
  <span class="logo">Doc Media AI</span>
  <div class="tab-bar">
    <button class="tab-btn active" onclick="showTab('chat')">Chat</button>
    <button class="tab-btn" onclick="showTab('content')">Content Studio</button>
    <button class="tab-btn" onclick="showTab('script')">Video Script</button>
    <button class="tab-btn" onclick="showTab('image')" style="background:#8E44AD;color:#fff;border-color:#8E44AD">&#x1F3A8; Tạo Hình</button>
    <button class="tab-btn" onclick="showTab('ads')">Facebook Ads</button>
    <button class="tab-btn" onclick="showTab('intake')" style="background:#E67E22;color:#fff;border-color:#E67E22">Intake Client</button>
    <button class="tab-btn" onclick="showTab('quote')" style="background:#27AE60;color:#fff;border-color:#27AE60">Báo Giá</button>
    <button class="tab-btn" onclick="showTab('library')">📚 Thư Viện</button>
    <button class="tab-btn" onclick="showTab('report')">📊 Báo Cáo</button>
  </div>
  <button class="btn-new" onclick="newChat()">+ Chat mới</button>
  <div class="bell-wrap">
    <button class="bell-btn" onclick="toggleNotif()" title="Thông báo">&#x1F514;</button>
    <span class="bell-badge" id="bell-badge">0</span>
    <div class="notif-panel" id="notif-panel">
      <div class="notif-header">
        <strong>Thông báo</strong>
        <button class="notif-read-all" onclick="markAllRead()">Đánh dấu đã đọc</button>
      </div>
      <div class="notif-list" id="notif-list">
        <div class="notif-empty">Chưa có thông báo</div>
      </div>
    </div>
  </div>
</header>

<!-- ── TAB: CHAT ── -->
<div id="tab-chat" class="tab-content active">
  <aside id="sidebar">
    <div id="sidebar-hdr">Lịch sử chat</div>
    <div id="session-list"></div>
  </aside>
  <div id="chat-main">
    <div id="chat-box"></div>
    <div id="file-banner">
      <span id="file-name"></span>
      <button onclick="clearFile()" style="background:none;border:none;cursor:pointer;font-size:18px;color:var(--text-2)">&#x2715;</button>
    </div>
    <div id="chat-input-area">
      <div class="qprompts">
        <button class="qbtn" onclick="sendQuick('CPM Facebook ngành nội thất hiện tại là bao nhiêu?')">Benchmark CPM</button>
        <button class="qbtn" onclick="sendQuick('Phân bổ ngân sách 50 triệu Facebook + TikTok')">Phân bổ ngân sách</button>
        <button class="qbtn" onclick="sendQuick('Viết 3 hook TikTok cho sofa cao cấp')">Hook TikTok</button>
        <button class="qbtn" onclick="sendQuick('KPI cần theo dõi cho lead gen nội thất?')">KPI lead gen</button>
      </div>
      <div id="input-wrap">
        <label class="icon-btn" for="file-input" title="Upload Excel/CSV">&#x1F4CE;</label>
        <input type="file" id="file-input" accept=".xlsx,.xls,.csv" onchange="handleUpload(this)">
        <textarea id="user-input" placeholder="Nhập câu hỏi... (Enter gửi, Shift+Enter xuống dòng)" rows="1"></textarea>
        <button id="send-btn" onclick="sendMessage()">Gửi</button>
      </div>
    </div>
  </div>
</div>

<!-- ── TAB: CONTENT STUDIO ── -->
<div id="tab-content" class="tab-content">
  <div class="studio-wrap">
    <div class="studio-form">
      <h2>&#x270D; Content Studio</h2>
      <div>
        <label>AI Model</label>
        <div class="seg-btns" id="cs-model-btns">
          <button class="seg active" data-v="claude-sonnet-4-6" title="Claude Sonnet 4.6 — cân bằng chất lượng và tốc độ">Sonnet 4.6</button>
          <button class="seg" data-v="claude-haiku-4-5-20251001" title="Claude Haiku 4.5 — nhanh nhất, tiết kiệm nhất">Haiku 4.5</button>
          <button class="seg" data-v="gemini-3.5-flash" title="Gemini 3.5 Flash — trợ giúp toàn diện, nhanh">3.5 Flash</button>
          <button class="seg" data-v="gemini-3.1-flash-lite" title="Gemini 3.1 Flash-Lite — nhanh nhất, tiết kiệm quota">Flash-Lite</button>
          <button class="seg" data-v="gemini-3.1-pro" title="Gemini 3.1 Pro — toán học và lập luận nâng cao">3.1 Pro</button>
        </div>
      </div>
      <div>
        <label>Tên client / Brand</label>
        <input id="cs-client" placeholder="VD: Shome Interior">
      </div>
      <div>
        <label>Sản phẩm / Dịch vụ</label>
        <input id="cs-product" placeholder="VD: Sofa Bắc Âu cao cấp, thi công nội thất">
      </div>
      <div>
        <label>Mục tiêu campaign</label>
        <div class="seg-btns" id="cs-obj-btns">
          <button class="seg active" data-v="Tăng brand awareness, tăng followers">Brand</button>
          <button class="seg" data-v="Thu hút leads, tin nhắn tư vấn">Leads</button>
          <button class="seg" data-v="Thúc đẩy mua hàng, tăng doanh thu">Doanh thu</button>
          <button class="seg" data-v="Retargeting khách hàng cũ">Retarget</button>
        </div>
      </div>
      <div>
        <label>Đối tượng mục tiêu</label>
        <input id="cs-target" placeholder="VD: Cặp vợ chồng 28-45, thu nhập 30tr+, HCMC">
      </div>
      <div>
        <label>Key message chính</label>
        <textarea id="cs-keymsg" placeholder="VD: Nội thất Bắc Âu tối giản, giao hàng & lắp đặt trong 7 ngày"></textarea>
      </div>
      <div>
        <label>Tone of voice</label>
        <div class="seg-btns" id="cs-tone-btns">
          <button class="seg active" data-v="Chuyên nghiệp, uy tín, sang trọng">Chuyên nghiệp</button>
          <button class="seg" data-v="Gần gũi, thân thiện, truyền cảm hứng">Thân thiện</button>
          <button class="seg" data-v="Trẻ trung, năng động, hiện đại">Trẻ trung</button>
        </div>
      </div>
      <div>
        <label>Công thức copywriting</label>
        <div class="seg-btns" id="cs-formula-btns">
          <button class="seg active" data-v="PAS" title="Problem → Agitate → Solve — đánh vào nỗi đau, rủi ro">PAS</button>
          <button class="seg" data-v="AIDA" title="Attention → Interest → Desire → Action — ra mắt mẫu/ưu đãi">AIDA</button>
          <button class="seg" data-v="BAB" title="Before → After → Bridge — showcase dự án, portfolio">BAB</button>
          <button class="seg" data-v="4U" title="Urgent · Unique · Ultra-specific · Useful — promotion giới hạn">4U</button>
        </div>
        <div id="cs-formula-desc" style="font-size:11px;color:var(--text-2);margin-top:5px;min-height:16px"></div>
      </div>
      <div>
        <label>Ngân sách ads (tùy chọn)</label>
        <input id="cs-budget" placeholder="VD: 10 triệu/tháng">
      </div>
      <div class="action-row">
        <button onclick="generateContent()" id="cs-gen-btn">Tạo content</button>
        <span id="cs-status"></span>
      </div>
    </div>
    <div class="studio-output" id="cs-output">
      <div class="placeholder-msg">Điền thông tin bên trái và bấm "Tạo content"<br>Claude sẽ tạo Facebook posts, TikTok captions và Image brief cho bạn</div>
    </div>
    <div class="save-lib-bar" id="cs-save-bar">
      <button class="save-lib-btn" onclick="saveCSToLib()">💾 Lưu vào thư viện</button>
      <span class="save-lib-status" id="cs-save-status"></span>
    </div>
  </div>
</div>

<!-- ── TAB: VIDEO SCRIPT ── -->
<div id="tab-script" class="tab-content">
  <div class="studio-wrap">
    <div class="studio-form">
      <h2>&#x1F3AC; Video Script</h2>
      <div>
        <label>AI Model</label>
        <div class="seg-btns" id="sc-model-btns">
          <button class="seg active" data-v="claude-sonnet-4-6" title="Claude Sonnet 4.6 — cân bằng chất lượng và tốc độ">Sonnet 4.6</button>
          <button class="seg" data-v="claude-haiku-4-5-20251001" title="Claude Haiku 4.5 — nhanh nhất, tiết kiệm nhất">Haiku 4.5</button>
          <button class="seg" data-v="gemini-3.5-flash" title="Gemini 3.5 Flash — trợ giúp toàn diện, nhanh">3.5 Flash</button>
          <button class="seg" data-v="gemini-3.1-flash-lite" title="Gemini 3.1 Flash-Lite — nhanh nhất, tiết kiệm quota">Flash-Lite</button>
          <button class="seg" data-v="gemini-3.1-pro" title="Gemini 3.1 Pro — toán học và lập luận nâng cao">3.1 Pro</button>
        </div>
      </div>
      <div>
        <label>Platform</label>
        <div class="seg-btns" id="sc-platform-btns">
          <button class="seg active" data-v="TikTok">TikTok</button>
          <button class="seg" data-v="Facebook Reels">FB Reels</button>
          <button class="seg" data-v="Instagram Reels">IG Reels</button>
          <button class="seg" data-v="YouTube Shorts">YT Shorts</button>
        </div>
      </div>
      <div>
        <label>Thời lượng</label>
        <div class="seg-btns" id="sc-dur-btns">
          <button class="seg" data-v="15">15s</button>
          <button class="seg active" data-v="30">30s</button>
          <button class="seg" data-v="60">60s</button>
          <button class="seg" data-v="90">90s</button>
        </div>
      </div>
      <div>
        <label>Sản phẩm / Chủ đề video</label>
        <input id="sc-product" placeholder="VD: Sofa Bắc Âu L-shape mới về">
      </div>
      <div>
        <label>Phong cách Hook</label>
        <div class="seg-btns" id="sc-hook-btns">
          <button class="seg active" data-v="Đặt vấn đề (Problem-based)">Vấn đề</button>
          <button class="seg" data-v="Tạo sự tò mò (Curiosity gap)">Tò mò</button>
          <button class="seg" data-v="Kết quả trước/sau (Before-after)">Trước/sau</button>
          <button class="seg" data-v="Shock, bất ngờ (Pattern interrupt)">Shock</button>
          <button class="seg" data-v="Hướng dẫn (Tutorial/How-to)">Tutorial</button>
        </div>
      </div>
      <div>
        <label>Đối tượng mục tiêu</label>
        <input id="sc-target" placeholder="VD: Chủ nhà mới sắm sofa lần đầu">
      </div>
      <div>
        <label>Ghi chú thêm (tùy chọn)</label>
        <textarea id="sc-extra" placeholder="VD: Có người thuyết minh, quay tại showroom, nhạc nền lo-fi..."></textarea>
      </div>
      <div class="action-row">
        <button onclick="generateScript()" id="sc-gen-btn">Tạo kịch bản</button>
        <span id="sc-status"></span>
      </div>
    </div>
    <div class="studio-output" id="sc-output">
      <div class="placeholder-msg">Điền thông tin video và bấm "Tạo kịch bản"<br>Claude sẽ viết kịch bản đầy đủ theo từng scene, có hướng dẫn quay phim</div>
    </div>
    <div class="save-lib-bar" id="sc-save-bar">
      <button class="save-lib-btn" onclick="saveSCToLib()">💾 Lưu vào thư viện</button>
      <span class="save-lib-status" id="sc-save-status"></span>
    </div>
  </div>
</div>

<!-- ── TAB: INTAKE CLIENT ── -->
<div id="tab-intake" class="tab-content">
  <div class="intake-wrap">
    <div class="intake-form-panel">
      <div class="form-title" style="color:#E67E22">Intake Client</div>

      <div class="intake-section-title">Thông tin cơ bản</div>
      <div>
        <label>Tên doanh nghiệp *</label>
        <input id="in-ten" placeholder="VD: RECO Design & Build">
      </div>
      <div>
        <label>Lĩnh vực kinh doanh</label>
        <input id="in-linh-vuc" placeholder="VD: Thiết kế & Thi công nội thất trọn gói">
      </div>
      <div class="row2">
        <div>
          <label>Fanpage / Website</label>
          <input id="in-fanpage" placeholder="Link hoặc tên page">
        </div>
        <div>
          <label>Người liên hệ</label>
          <input id="in-lien-he" placeholder="Tên + SĐT">
        </div>
      </div>

      <div class="intake-section-title">Tình trạng hiện tại</div>
      <div class="row2">
        <div>
          <label>Fanpage</label>
          <select id="in-page-status">
            <option>Chưa có fanpage</option>
            <option>Mới tạo (dưới 1.000 followers)</option>
            <option>Đã có (1k - 10k followers)</option>
            <option>Đã có (trên 10k followers)</option>
          </select>
        </div>
        <div>
          <label>Số followers hiện tại</label>
          <input id="in-followers" type="number" placeholder="0">
        </div>
      </div>
      <div>
        <label>Đã từng chạy Facebook Ads chưa?</label>
        <select id="in-ads-history">
          <option>Chưa chạy bao giờ</option>
          <option>Đã chạy nhưng không hiệu quả</option>
          <option>Đã chạy, có kết quả nhưng muốn tốt hơn</option>
          <option>Đang chạy, muốn thuê agency quản lý</option>
        </select>
      </div>
      <div>
        <label>CPL / Chi phí mỗi lead cũ (nếu đã chạy)</label>
        <input id="in-cpl-cu" placeholder="VD: 250.000 VND/inbox hoặc 'không biết'">
      </div>

      <div class="intake-section-title">Mục tiêu</div>
      <div>
        <label>Mục tiêu chính</label>
        <select id="in-muc-tieu">
          <option>Thu leads (khách hàng nhắn tin tư vấn)</option>
          <option>Tăng nhận diện thương hiệu (brand awareness)</option>
          <option>Tăng doanh thu trực tiếp</option>
          <option>Cả 3: leads + brand + doanh thu</option>
        </select>
      </div>
      <div class="row2">
        <div>
          <label>Leads mong muốn/tháng</label>
          <input id="in-leads" placeholder="VD: 30 leads">
        </div>
        <div>
          <label>Doanh thu kỳ vọng/tháng</label>
          <input id="in-doanh-thu" placeholder="VD: 500 triệu">
        </div>
      </div>
      <div>
        <label>Khi nào muốn bắt đầu?</label>
        <input id="in-timeline" placeholder="VD: Đầu tháng 7/2026">
      </div>

      <div class="intake-section-title">Sản phẩm & Đối tượng</div>
      <div>
        <label>Mô tả sản phẩm / dịch vụ chính</label>
        <textarea id="in-sp" placeholder="VD: Thiết kế và thi công nội thất căn hộ trọn gói, giá từ 150 triệu, bảo hành 2 năm..."></textarea>
      </div>
      <div>
        <label>Đối tượng khách hàng mục tiêu</label>
        <textarea id="in-doi-tuong" placeholder="VD: Vợ chồng 28-42 tuổi, vừa mua căn hộ mới, thu nhập 30-60 triệu/tháng, ở HCMC..."></textarea>
      </div>
      <div>
        <label>USP — Điểm khác biệt so với đối thủ</label>
        <textarea id="in-usp" placeholder="VD: Thiết kế theo phong cách Wabi-Sabi, có showroom tại VGP, bảo hành 3 năm..."></textarea>
      </div>

      <div class="intake-section-title">Assets & Ngân sách</div>
      <div>
        <label>Assets sẵn có (mô tả chi tiết)</label>
        <textarea id="in-assets" placeholder="VD: Có ảnh thực tế dự án (10+ ảnh), chưa có video, có brand guideline cơ bản..."></textarea>
      </div>
      <div class="row2">
        <div>
          <label>Ngân sách Ads/tháng (VND)</label>
          <input id="in-budget-ads" type="number" placeholder="VD: 15000000">
        </div>
        <div>
          <label>Ngân sách dịch vụ/tháng</label>
          <input id="in-budget-dv" placeholder="VD: 5 triệu">
        </div>
      </div>
      <div>
        <label>Thời gian hợp đồng mong muốn</label>
        <select id="in-duration">
          <option>1 tháng thử việc</option>
          <option>3 tháng</option>
          <option>6 tháng</option>
          <option>12 tháng</option>
        </select>
      </div>

      <div class="intake-section-title">Đối thủ cạnh tranh</div>
      <div>
        <label>Tên 1-3 đối thủ chính (tên page hoặc tên công ty)</label>
        <textarea id="in-doi-thu" placeholder="VD: Nội thất ABC, XYZ Interior, ..."></textarea>
      </div>

      <div style="margin-top:10px">
        <label style="font-size:12px;color:var(--text-2);margin-bottom:4px;display:block">AI Model</label>
        <div class="seg-btns" id="in-model-btns">
          <button class="seg active" data-v="gemini-3.1-flash-lite">Flash-Lite</button>
          <button class="seg" data-v="gemini-3.5-flash">3.5 Flash</button>
          <button class="seg" data-v="gemini-3.1-pro">3.1 Pro</button>
          <button class="seg" data-v="claude-sonnet-4-6">Sonnet 4.6</button>
          <button class="seg" data-v="claude-haiku-4-5-20251001">Haiku 4.5</button>
        </div>
      </div>

      <button class="btn-orange" id="intake-submit-btn" onclick="submitIntake()" style="width:100%;margin-top:8px">
        Phân tích & Tạo Brief
      </button>
    </div>

    <div class="intake-output">
      <div class="intake-output-header">
        <strong id="intake-status">Điền đầy đủ thông tin rồi nhấn "Phân tích & Tạo Brief"</strong>
        <button class="btn-outline-orange" id="intake-copy-btn" onclick="copyIntake()" disabled>Copy Brief</button>
        <button class="save-lib-btn" id="intake-save-btn" onclick="saveIntakeToLib()" disabled style="border-color:var(--orange);color:var(--orange)">💾 Lưu</button>
        <button class="btn-orange" id="intake-to-quote-btn" onclick="intakeToQuote()" disabled>Chuyển sang Báo Giá</button>
      </div>
      <div class="intake-output-body" id="intake-output-body"></div>
    </div>
  </div>
</div>

<!-- ── TAB: BÁO GIÁ ── -->
<div id="tab-quote" class="tab-content">
  <div class="quote-wrap">
    <div class="quote-form-panel">
      <div class="form-title">Tạo Báo Giá</div>
      <div>
        <label>Tên doanh nghiệp khách hàng *</label>
        <input id="q-client" placeholder="VD: RECO Design & Build">
      </div>
      <div>
        <label>Lĩnh vực</label>
        <select id="q-niche">
          <option>Nội thất / Thiết kế & Thi công</option>
          <option>Kiến trúc / Xây dựng</option>
          <option>Bất động sản</option>
          <option>Nhà hàng / F&B</option>
          <option>Làm đẹp / Spa</option>
          <option>Giáo dục / Khoá học</option>
          <option>Khác</option>
        </select>
      </div>
      <div>
        <label>Mục tiêu chiến dịch</label>
        <select id="q-obj">
          <option>Tin nhắn (Lead) là chính</option>
          <option>Tin nhắn + Tương tác</option>
          <option>Tương tác + Brand Awareness</option>
          <option>Toàn bộ: Tin nhắn, Tương tác, Tiếp cận</option>
        </select>
      </div>
      <div>
        <label>Ngân sách quảng cáo/tháng (VND)</label>
        <input id="q-budget" type="number" placeholder="VD: 15000000">
      </div>
      <div>
        <label>Thời gian hợp đồng</label>
        <select id="q-duration">
          <option>1 tháng (thử)</option>
          <option>3 tháng</option>
          <option>6 tháng</option>
          <option>12 tháng</option>
        </select>
      </div>
      <div>
        <label>Dịch vụ yêu cầu thêm</label>
        <textarea id="q-services" placeholder="VD: Sản xuất video Reels, Chạy TikTok Ads kèm theo..."></textarea>
      </div>
      <div>
        <label>Ghi chú / Yêu cầu đặc biệt</label>
        <textarea id="q-note" placeholder="VD: Client cần báo giá gấp trong hôm nay..."></textarea>
      </div>
      <div style="margin-top:10px">
        <label style="font-size:12px;color:var(--text-2);margin-bottom:4px;display:block">AI Model</label>
        <div class="seg-btns" id="q-model-btns">
          <button class="seg active" data-v="gemini-3.5-flash">3.5 Flash</button>
          <button class="seg" data-v="gemini-3.1-pro">3.1 Pro</button>
          <button class="seg" data-v="claude-sonnet-4-6">Sonnet 4.6</button>
          <button class="seg" data-v="claude-haiku-4-5-20251001">Haiku 4.5</button>
        </div>
      </div>
      <button class="btn-green" id="q-submit-btn" onclick="submitQuote()" style="width:100%;margin-top:8px">
        Tạo Báo Giá
      </button>
    </div>
    <div class="quote-output">
      <div class="quote-output-header">
        <strong id="q-status">Điền thông tin và nhấn "Tạo Báo Giá"</strong>
        <button class="btn-outline-green" id="q-copy-btn" onclick="copyQuote()" disabled>Copy</button>
        <button class="btn-green" id="q-export-btn" onclick="exportQuoteExcel()" disabled>Xuất Excel</button>
        <button class="save-lib-btn" id="q-save-btn" onclick="saveQuoteToLib()" disabled style="border-color:var(--green);color:var(--green)">💾 Lưu</button>
      </div>
      <div class="quote-output-body" id="quote-output-body"></div>
    </div>
  </div>
</div>

<!-- ── TAB: LIBRARY ── -->
<div id="tab-library" class="tab-content">
  <div class="lib-header">
    <span class="lib-header-title">📚 Thư viện Nội dung</span>
    <div class="lib-export-btns">
      <button class="lib-exp-btn" onclick="exportLibrary('word')">📄 Word</button>
      <button class="lib-exp-btn" onclick="exportLibrary('excel')">📊 Excel</button>
      <button class="lib-exp-btn" onclick="exportLibrary('pdf')">📑 PDF</button>
    </div>
  </div>
  <div class="lib-toolbar">
    <div class="seg-btns" id="lib-type-btns">
      <button class="seg active" data-v="">Tất cả</button>
      <button class="seg" data-v="content">Bài viết</button>
      <button class="seg" data-v="script">Kịch bản</button>
      <button class="seg" data-v="brief">Brief</button>
      <button class="seg" data-v="quote">Báo giá</button>
    </div>
    <input type="text" id="lib-search" placeholder="🔍 Tìm kiếm theo tên, client, nội dung..." oninput="loadLibrary()">
    <span id="lib-count"></span>
  </div>
  <div class="lib-list" id="lib-list">
    <div class="no-data">Chưa có nội dung nào được lưu — Tạo content rồi nhấn 💾 Lưu vào thư viện</div>
  </div>
</div>

<!-- ── TAB: REPORT ── -->
<div id="tab-report" class="tab-content">
  <div class="rpt-form">
    <h2>📊 Báo Cáo Tháng</h2>

    <label>Tài khoản quảng cáo</label>
    <select id="rpt-account-sel">
      <option value="">-- Đang tải... --</option>
    </select>

    <label>Kỳ báo cáo</label>
    <select id="rpt-period" onchange="toggleRptCustomDates()">
      <option value="this_month">Tháng này</option>
      <option value="last_month" selected>Tháng trước</option>
      <option value="custom">Tùy chọn...</option>
    </select>
    <div class="rpt-custom-dates" id="rpt-custom-dates">
      <input type="date" id="rpt-date-from" title="Từ ngày">
      <input type="date" id="rpt-date-to" title="Đến ngày">
    </div>

    <label>Tên khách hàng / thương hiệu</label>
    <input type="text" id="rpt-client" placeholder="VD: RECO Design & Build">

    <label>Model AI</label>
    <div class="seg-btns" id="rpt-model-btns" style="margin-bottom:12px">
      <button class="seg active" data-v="claude-sonnet-4-6" title="Claude Sonnet 4.6 — cân bằng chất lượng và tốc độ">Sonnet 4.6</button>
      <button class="seg" data-v="claude-haiku-4-5-20251001" title="Claude Haiku 4.5 — nhanh nhất, tiết kiệm nhất">Haiku 4.5</button>
      <button class="seg" data-v="gemini-3.5-flash" title="Gemini 3.5 Flash — trợ giúp toàn diện, nhanh">3.5 Flash</button>
      <button class="seg" data-v="gemini-3.1-flash-lite" title="Gemini 3.1 Flash-Lite — nhanh nhất, tiết kiệm quota">Flash-Lite</button>
      <button class="seg" data-v="gemini-3.1-pro" title="Gemini 3.1 Pro — toán học và lập luận nâng cao">3.1 Pro</button>
    </div>

    <button id="rpt-gen-btn" onclick="generateReport()">📊 Tạo báo cáo</button>
    <div class="rpt-status" id="rpt-status"></div>

    <div class="rpt-export-row" id="rpt-export-row" style="display:none">
      <button class="rpt-exp-btn" onclick="exportReport('word')">📄 Word</button>
      <button class="rpt-exp-btn" onclick="exportReport('excel')">📊 Excel</button>
      <button class="rpt-exp-btn" onclick="exportReport('pdf')">📑 PDF</button>
    </div>
  </div>

  <div class="rpt-preview" id="rpt-preview">
    <div class="rpt-placeholder">📊 Chọn tài khoản, kỳ báo cáo và nhấn<br><strong>Tạo báo cáo</strong> để bắt đầu</div>
  </div>
</div>

<!-- ── TAB: TAO HINH ── -->
<div id="tab-image" class="tab-content">
  <div class="studio-wrap">
    <div class="studio-col" style="max-width:440px;flex-shrink:0">
      <div class="studio-form">
        <h2 style="margin:0 0 4px;font-size:18px;color:var(--text)">&#x1F3A8; Tạo ảnh concept nội thất</h2>
        <p style="margin:0 0 16px;font-size:13px;color:var(--text-2)">Nhập mô tả, chọn phong cách và nhấn nút Tạo ảnh để xuất bản ảnh pitch client</p>
        <label class="form-label">AI Model</label>
        <div class="seg-btns" id="img-model-btns" style="margin-bottom:14px">
          <button class="seg" data-v="gemini-3.5-flash" title="Gemini 3.5 Flash">3.5 Flash</button>
          <button class="seg active" data-v="gemini-3.1-flash-lite" title="Gemini Flash-Lite — nhanh nhất">Flash-Lite</button>
          <button class="seg" data-v="gemini-3.1-pro" title="Gemini 3.1 Pro">3.1 Pro</button>
          <button class="seg" data-v="claude-sonnet-4-6" title="Claude Sonnet 4.6">Sonnet 4.6</button>
          <button class="seg" data-v="claude-haiku-4-5-20251001" title="Claude Haiku 4.5">Haiku 4.5</button>
        </div>

        <label class="form-label">Mô tả không gian <span style="font-weight:400;color:var(--text-3)">(tiếng Việt)</span></label>
        <textarea id="img-desc" rows="4" placeholder="Ví dụ: Phòng khách Japandi 25m², gỗ tự nhiên, sofa xanh rêu nhạt, ánh sáng buổi sáng, cây xanh trang trí..."></textarea>
        <div style="display:flex;gap:8px;align-items:center;margin-top:8px;margin-bottom:4px">
          <button id="img-enhance-btn" onclick="enhanceImgPrompt()" style="flex:1;padding:8px 12px;font-size:13px;background:var(--surface-2);border:1px solid var(--border);border-radius:8px;color:var(--text);cursor:pointer;font-weight:500;transition:all .15s" onmouseover="this.style.borderColor='var(--accent)'" onmouseout="this.style.borderColor='var(--border)'">✨ Tạo prompt bằng AI</button>
          <span id="img-enhance-status" style="font-size:12px;color:var(--text-3)"></span>
        </div>
        <div id="img-prompt-preview" style="display:none;margin-bottom:12px;padding:10px;background:var(--surface-2);border:1px solid var(--accent);border-radius:8px;font-size:12px;color:var(--text-2);line-height:1.6;font-style:italic;position:relative">
          <div style="font-size:10px;color:var(--accent);font-weight:600;text-transform:uppercase;letter-spacing:.5px;margin-bottom:4px">Prompt AI đã tạo</div>
          <div id="img-prompt-preview-text"></div>
          <button onclick="useAiPrompt()" style="margin-top:8px;width:100%;padding:6px;font-size:12px;background:var(--accent);color:#fff;border:none;border-radius:6px;cursor:pointer;font-weight:600">✅ Dùng prompt này</button>
        </div>
        <label class="form-label" style="margin-top:4px">Phong cách thiết kế</label>
        <div class="style-chips">
          <button class="schip active" data-s="Modern Minimalist">Hien dai</button>
          <button class="schip" data-s="Japandi Japanese Scandinavian">Japandi</button>
          <button class="schip" data-s="Scandinavian Nordic cozy">Bac Au</button>
          <button class="schip" data-s="Neo-classical elegant">Tan co dien</button>
          <button class="schip" data-s="Ultra luxury high-end">Luxury</button>
          <button class="schip" data-s="Ultra minimalist zen">Toi gian</button>
          <button class="schip" data-s="Industrial loft concrete">Industrial</button>
        </div>
        <div style="display:flex;gap:14px;align-items:center;margin-top:20px">
          <button id="gen-btn" onclick="generateImg()" style="padding:10px 28px;font-size:15px">&#x2728; Tao anh</button>
          <span id="gen-status" style="font-size:13px;color:var(--text-2)"></span>
        </div>
        <div style="margin-top:20px;padding:14px;background:var(--surface-2);border:1px solid var(--border);border-radius:10px;font-size:13px;color:var(--text-2);line-height:1.8">
          <strong style="color:var(--text)">Mẹo sử dụng:</strong><br>
          • Mô tả càng cụ thể, ảnh càng đẹp<br>
          • Nêu rõ kích thước phòng, tông màu, vật liệu chính<br>
          • Sau khi tạo xong có thể thêm vào chat để trao đổi
        </div>
      </div>
    </div>
    <div class="studio-col">
      <div id="img-result" style="display:none;flex-direction:column;gap:16px;height:100%">
        <div style="display:flex;align-items:center;justify-content:space-between">
          <strong style="font-size:14px;color:var(--text)">Kết quả</strong>
          <div style="display:flex;gap:10px">
            <a id="img-dl" href="" target="_blank" download style="font-size:13px;padding:7px 16px;border:1px solid var(--border-2);border-radius:8px;background:var(--surface-2);color:var(--text);text-decoration:none;font-weight:500">&#x2B07; Tải xuống</a>
            <button onclick="addImgToChat(); showTab('chat')" style="font-size:13px;padding:7px 16px;border:1px solid var(--accent);border-radius:8px;background:var(--accent);color:#fff;cursor:pointer;font-weight:600">Thêm vào Chat</button>
          </div>
        </div>
        <img id="gen-img-el" src="" alt="Concept" style="width:100%;border-radius:12px;box-shadow:0 4px 20px rgba(0,0,0,.1);object-fit:cover">
        <div style="padding:12px;background:var(--surface-2);border:1px solid var(--border);border-radius:8px">
          <div style="font-size:11px;color:var(--text-3);font-weight:600;text-transform:uppercase;letter-spacing:.5px;margin-bottom:4px">Prompt da dung</div>
          <div style="font-size:12px;color:var(--text-2);line-height:1.5;font-style:italic" id="gen-prompt-txt"></div>
        </div>
      </div>
      <div id="img-placeholder" style="display:flex;flex-direction:column;align-items:center;justify-content:center;height:100%;color:#ccc;gap:12px">
        <div style="font-size:72px">&#x1F5BC;</div>
        <div style="font-size:15px;color:var(--text-2)">Hình ảnh concept sẽ hiển thị ở đây</div>
        <div style="font-size:13px;color:var(--text-3)">Nhập mô tả và nhấn Tạo ảnh</div>
      </div>
    </div>
  </div>
</div>

<!-- ── TAB: FACEBOOK ADS ── -->
<div id="tab-ads" class="tab-content">
  <div class="studio-wrap" style="flex-direction:column">
    <div class="ads-controls">
      <strong>&#x1F4CA; Facebook Ads</strong>
      <select id="ads-account-sel">
        <option value="">-- Đang tải tài khoản... --</option>
      </select>
      <div class="period-seg" id="period-btns">
        <button data-p="today" onclick="setPeriod(this)">Hôm nay</button>
        <button data-p="yesterday" onclick="setPeriod(this)">Hôm qua</button>
        <button class="active" data-p="last_7d" onclick="setPeriod(this)">7 ngày</button>
        <button data-p="last_30d" onclick="setPeriod(this)">30 ngày</button>
        <button data-p="custom" onclick="setPeriod(this)">Tùy chọn</button>
      </div>
      <div id="custom-dates">
        <input type="date" id="ads-from" title="Từ ngày">
        <span style="color:var(--text-3)">&#x2192;</span>
        <input type="date" id="ads-to" title="Đến ngày">
      </div>
      <button id="refresh-ads-btn" onclick="loadAdsData()">Tải dữ liệu</button>
    </div>
    <div class="ads-body" id="ads-body">
      <div class="no-data" id="ads-placeholder">Bấm "Tải dữ liệu" để xem số liệu Facebook Ads của bạn<br><br>Yêu cầu: thêm FB_ACCESS_TOKEN và FB_AD_ACCOUNT_ID vào file .env</div>
    </div>
  </div>
</div>

<script>
const $ = id => document.getElementById(id);
let currentSession = localStorage.getItem('dm_sess') || makeId();
let fileContext = '';
let isStreaming = false;
let selectedImgStyle = 'Modern Minimalist';
let genImgUrl = '', genImgPrompt = '', genImgDesc = '';
let adsPeriod = 'last_7d';
let adsRawData = null;

function makeId(){ return 'sess_'+Date.now()+'_'+Math.random().toString(36).slice(2,6); }
function esc(t){ return String(t).replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;'); }
function md(text){
  return esc(text)
    .replace(/\\*\\*([^*\\n]+)\\*\\*/g,'<strong>$1</strong>')
    .replace(/\\*([^*\\n]+)\\*/g,'<em>$1</em>')
    .replace(/`([^`]+)`/g,'<code>$1</code>')
    .replace(/^### (.+)$/gm,'<h3>$1</h3>')
    .replace(/^## (.+)$/gm,'<h2 style="font-size:15px;margin:12px 0 6px;color:#a78bfa">$2</h2>'.replace('$2','$1'))
    .replace(/^- (.+)$/gm,'<span style="display:block;padding-left:14px">&bull; $1</span>')
    .replace(/\\n/g,'<br>');
}
function scrollBot(){ const c=$('chat-box'); c.scrollTop=c.scrollHeight; }

// ── Tabs ──
function showTab(name){
  document.querySelectorAll('.tab-content').forEach(el=>el.classList.remove('active'));
  document.querySelectorAll('.tab-btn').forEach(el=>el.classList.remove('active'));
  $('tab-'+name).classList.add('active');
  const idx=['chat','content','script','image','ads','intake','quote','library','report'].indexOf(name);
  if(idx>=0) document.querySelectorAll('.tab-btn')[idx].classList.add('active');
  if(name==='chat'){ $('user-input').focus(); }
  if(name==='ads' && $('ads-account-sel').options.length<=1){ loadAdAccounts(); }
  if(name==='image'){ $('img-desc').focus(); }
  if(name==='library'){ loadLibrary(); }
  if(name==='report'){ loadReportAccounts(); }
}

// ── Seg button helpers ──
function setupSegs(containerId){
  const btns = $(containerId).querySelectorAll('.seg');
  btns.forEach(b=>b.addEventListener('click',()=>{
    btns.forEach(x=>x.classList.remove('active'));
    b.classList.add('active');
  }));
}
function getSegVal(containerId){
  const active=$(containerId).querySelector('.seg.active');
  return active ? (active.dataset.v||active.textContent) : '';
}

window.onload = async ()=>{
  localStorage.setItem('dm_sess',currentSession);
  setupSegs('cs-obj-btns');
  setupSegs('cs-tone-btns');
  setupSegs('cs-formula-btns');
  setupSegs('cs-model-btns');
  setupSegs('sc-platform-btns');
  setupSegs('sc-model-btns');
  // Show description when formula changes
  const formulaDescs = {
    'PAS': 'Đánh vào nỗi đau & rủi ro — tốt nhất cho lead gen tin nhắn',
    'AIDA': 'Ra mắt mẫu thiết kế hoặc chương trình ưu đãi giới hạn',
    'BAB': 'Showcase dự án thực tế, trước/sau thi công — tốt cho portfolio',
    '4U': 'Promotion khẩn cấp, offer giới hạn số lượng/thời gian',
  };
  document.querySelectorAll('#cs-formula-btns .seg').forEach(btn=>{
    btn.addEventListener('click',()=>{
      const desc=$('cs-formula-desc');
      if(desc) desc.textContent=formulaDescs[btn.dataset.v]||'';
    });
  });
  const firstFormula=document.querySelector('#cs-formula-btns .seg.active');
  if(firstFormula&&$('cs-formula-desc')) $('cs-formula-desc').textContent=formulaDescs[firstFormula.dataset.v]||'';
  setupSegs('sc-dur-btns');
  setupSegs('sc-hook-btns');
  setupSegs('in-model-btns');
  setupSegs('q-model-btns');
  setupSegs('lib-type-btns');
  setupSegs('rpt-model-btns');
  setupSegs('img-model-btns');
  document.querySelectorAll('.schip').forEach(c=>c.addEventListener('click',()=>{
    document.querySelectorAll('.schip').forEach(x=>x.classList.remove('active'));
    c.classList.add('active');
    selectedImgStyle=c.dataset.s;
  }));
  await loadSessions();
  await loadHistory(currentSession);
  await refreshAlerts();
  setInterval(refreshAlerts, 5*60*1000); // kiểm tra thông báo mỗi 5 phút
};

// ── Notifications ──
async function refreshAlerts(){
  try{
    const data=await fetch('/alerts').then(r=>r.json());
    const badge=$('bell-badge');
    if(data.unread>0){
      badge.style.display='flex';
      badge.textContent=data.unread>9?'9+':data.unread;
    } else {
      badge.style.display='none';
    }
    renderNotifList(data.alerts);
  }catch{}
}

function renderNotifList(alerts){
  const list=$('notif-list');
  if(!alerts||!alerts.length){ list.innerHTML='<div class="notif-empty">Chưa có thông báo</div>'; return; }
  list.innerHTML=alerts.map(a=>`
    <div class="notif-item${a.read?'':' unread'}" onclick="readAlert('${a.id}',this)">
      <div class="notif-title">${esc(a.title)}</div>
      <div class="notif-body">${esc(a.body)}</div>
      <div class="notif-time">${(a.ts||'').substring(0,16)}</div>
    </div>`).join('');
}

async function readAlert(id, el){
  el.classList.remove('unread');
  await fetch('/alerts/'+id+'/read',{method:'POST'});
  await refreshAlerts();
}

async function markAllRead(){
  await fetch('/alerts/read-all',{method:'POST'});
  await refreshAlerts();
}

function toggleNotif(){
  const p=$('notif-panel');
  p.classList.toggle('open');
  if(p.classList.contains('open')) markAllRead();
}

document.addEventListener('click',e=>{
  const wrap=document.querySelector('.bell-wrap');
  if(wrap&&!wrap.contains(e.target)) $('notif-panel').classList.remove('open');
});

// ── Chat sessions ──
async function loadSessions(){
  try{
    const data=await fetch('/sessions').then(r=>r.json());
    const list=$('session-list');
    list.innerHTML='';
    if(!data.length){ list.innerHTML='<div style="color:#5c6878;font-size:13px;text-align:center;padding:20px">Chua co lich su</div>'; return; }
    data.forEach(s=>{
      const d=document.createElement('div');
      d.className='sess-item'+(s.id===currentSession?' active':'');
      d.innerHTML='<div class="sess-prev" style="flex:1;min-width:0;overflow:hidden;text-overflow:ellipsis;white-space:nowrap">'+esc(s.preview)+'</div>'+
        '<div class="sess-time">'+esc(s.time)+'</div>'+
        '<button class="sess-del" data-sid="'+esc(s.id)+'" onclick="deleteSession(event,this.dataset.sid)" title="Xoa">&#x2715;</button>';
      d.querySelector('.sess-prev').onclick=()=>switchSession(s.id);
      d.querySelector('.sess-time').onclick=()=>switchSession(s.id);
      list.appendChild(d);
    });
    if(!data.find(s=>s.id===currentSession)){
      const d=document.createElement('div');
      d.className='sess-item active';
      d.innerHTML='<div class="sess-prev">Chat moi</div>';
      list.insertBefore(d,list.firstChild);
    }
  }catch(e){}
}

async function loadHistory(sid){
  try{
    const msgs=await fetch('/history/'+encodeURIComponent(sid)).then(r=>r.json());
    $('chat-box').innerHTML='';
    msgs.length ? msgs.forEach(m=>appendMsg(m.content,m.role)) : addWelcome();
  }catch(e){ addWelcome(); }
}

async function deleteSession(e, sid){
  e.stopPropagation();
  await fetch('/session/'+encodeURIComponent(sid),{method:'DELETE'});
  if(sid===currentSession){ await newChat(); } else { await loadSessions(); }
}

async function switchSession(sid){
  if(sid===currentSession)return;
  currentSession=sid;
  localStorage.setItem('dm_sess',sid);
  clearFile();
  await loadSessions();
  await loadHistory(sid);
}

async function newChat(){
  currentSession=makeId();
  localStorage.setItem('dm_sess',currentSession);
  $('chat-box').innerHTML='';
  addWelcome();
  clearFile();
  await loadSessions();
  showTab('chat');
  $('user-input').focus();
}

function addWelcome(){
  const d=document.createElement('div');
  d.className='msg assistant';
  d.innerHTML='Xin chào! AI Agent Độc Media sẵn sàng.<br><br>'+
    '<strong>Lệnh Agent (gõ tự nhiên):</strong><br>'+
    '&bull; <em>"Viết bài PAS cho sofa da, target vợ chồng 30-40t"</em><br>'+
    '&bull; <em>"Tạo brief cho client RECO, ngân sách 15 triệu"</em><br>'+
    '&bull; <em>"Làm báo giá 3 tháng cho Nội thất ABC"</em><br>'+
    '&bull; <em>"Kết quả ads 7 ngày qua thế nào?"</em><br><br>'+
    '<strong>Tính năng khác:</strong><br>'+
    '&bull; <strong>Content Studio</strong>: tạo Facebook posts, TikTok captions, Image brief<br>'+
    '&bull; <strong>Video Script</strong>: kịch bản TikTok/Reels theo từng scene<br>'+
    '&bull; <strong>Tạo Hình</strong>: tạo ảnh concept pitch client AI<br>'+
    '&bull; <strong>Facebook Ads</strong>: xem số liệu live + phân tích AI<br>'+
    '&bull; Upload Excel để phân tích<br>'+
    '&bull; Web search số liệu mới nhất<br>'+
    '&bull; Các tab: Content Studio, Video Script, Tạo Hình, Intake, Báo Giá';
  $('chat-box').appendChild(d);
  scrollBot();
}

function appendMsg(content,role){
  const d=document.createElement('div');
  d.className='msg '+role;
  if(role==='assistant') d.innerHTML=md(content);
  else d.textContent=content;
  $('chat-box').appendChild(d);
  scrollBot();
  return d;
}

// ── Chat send ──
function sendQuick(t){ $('user-input').value=t; sendMessage(); }
$('user-input') && $('user-input').addEventListener('keydown',e=>{
  if(e.key==='Enter'&&!e.shiftKey){ e.preventDefault(); sendMessage(); }
});
$('user-input') && $('user-input').addEventListener('input',()=>{
  const el=$('user-input');
  el.style.height='44px';
  el.style.height=Math.min(el.scrollHeight,120)+'px';
});

async function sendMessage(){
  const text=($('user-input')||{}).value?.trim();
  if(!text||isStreaming)return;
  appendMsg(text,'user');
  $('user-input').value='';
  $('user-input').style.height='44px';
  const bubble=document.createElement('div');
  bubble.className='msg assistant streaming';
  $('chat-box').appendChild(bubble);
  scrollBot();
  isStreaming=true; $('send-btn').disabled=true;
  let full='';
  try{
    const res=await fetch('/chat',{method:'POST',headers:{'Content-Type':'application/json'},
      body:JSON.stringify({message:text,session_id:currentSession,file_context:fileContext})});
    const reader=res.body.getReader();
    const dec=new TextDecoder();
    while(true){
      const{done,value}=await reader.read();
      if(done)break;
      for(const line of dec.decode(value).split('\\n')){
        if(!line.startsWith('data: '))continue;
        const d=line.slice(6);
        if(d==='[DONE]')break;
        try{
          const parsed=JSON.parse(d);
          if(parsed.tool){
            // Hiển thị tool bubble trước bubble AI
            const tb=document.createElement('div');
            tb.className='tool-bubble';
            tb.id='tool-'+parsed.tool;
            tb.innerHTML='<div class="tool-spinner"></div><span>'+parsed.label+'</span>';
            $('chat-box').insertBefore(tb, bubble);
            scrollBot();
          } else if(parsed.text){
            full+=parsed.text; bubble.innerHTML=md(full); scrollBot();
          }
        }catch{}
      }
    }
    bubble.classList.remove('streaming');
    await loadSessions();
  }catch(e){ bubble.textContent='Loi ket noi.'; bubble.classList.remove('streaming'); }
  isStreaming=false; $('send-btn').disabled=false; $('user-input').focus();
}

// ── File upload ──
async function handleUpload(inp){
  const file=inp.files[0]; if(!file)return;
  const banner=$('file-banner');
  $('file-name').textContent='Dang doc file...';
  banner.style.display='flex';
  const fd=new FormData(); fd.append('file',file);
  try{
    const res=await fetch('/upload',{method:'POST',body:fd});
    if(!res.ok)throw new Error(await res.text());
    const data=await res.json();
    fileContext=data.summary;
    $('file-name').textContent=data.filename+' ('+data.rows+' dong)';
    const d=document.createElement('div');
    d.className='msg system';
    d.textContent='Da tai len: '+data.filename+' — '+data.rows+' dong. Ban muon phan tich gi?';
    $('chat-box').appendChild(d); scrollBot();
  }catch(e){ $('file-name').textContent='Loi: '+e.message; fileContext=''; }
  inp.value='';
}
function clearFile(){ fileContext=''; $('file-banner').style.display='none'; $('file-input').value=''; }

// ── Image tab ──
let _aiGeneratedPrompt = '';

async function enhanceImgPrompt(){
  const desc = $('img-desc').value.trim();
  if(!desc){ $('img-desc').focus(); alert('Nhập mô tả không gian trước nhé!'); return; }
  const aiModel = getSegVal('img-model-btns') || 'gemini-3.5-flash';
  const btn = $('img-enhance-btn');
  const status = $('img-enhance-status');
  btn.disabled = true;
  status.textContent = 'Đang tạo...';
  $('img-prompt-preview').style.display = 'none';
  try{
    const res = await fetch('/image/enhance-prompt', {method:'POST', headers:{'Content-Type':'application/json'},
      body: JSON.stringify({description: desc, style: selectedImgStyle, ai_model: aiModel})});
    if(!res.ok){ const t=await res.text(); throw new Error(t); }
    const d = await res.json();
    _aiGeneratedPrompt = d.prompt;
    $('img-prompt-preview-text').textContent = d.prompt;
    $('img-prompt-preview').style.display = 'block';
    status.textContent = '✅ Xong!';
    setTimeout(()=>{ status.textContent=''; }, 3000);
  }catch(e){
    status.textContent = '❌ ' + e.message.slice(0,60);
  }finally{
    btn.disabled = false;
  }
}

function useAiPrompt(){
  if(_aiGeneratedPrompt){
    $('img-desc').value = _aiGeneratedPrompt;
    $('img-prompt-preview').style.display = 'none';
  }
}

async function generateImg(){
  const desc=$('img-desc').value.trim();
  if(!desc){ $('img-desc').focus(); return; }
  const status=$('gen-status');
  $('img-result').style.display='none';
  if($('img-placeholder')) $('img-placeholder').style.display='flex';
  $('gen-btn').disabled=true;
  status.textContent='Đang xử lý... (15-30s)';
  try{
    const res=await fetch('/generate-image',{method:'POST',headers:{'Content-Type':'application/json'},
      body:JSON.stringify({description:desc,style:selectedImgStyle})});
    if(!res.ok)throw new Error((await res.json()).detail);
    const data=await res.json();
    genImgUrl=data.url; genImgPrompt=data.prompt; genImgDesc=desc;
    $('gen-img-el').src=data.url;
    $('img-dl').href=data.url;
    $('gen-prompt-txt').textContent=data.prompt;
    if($('img-placeholder')) $('img-placeholder').style.display='none';
    $('img-result').style.display='flex';
    status.textContent='Hoàn thành!';
    setTimeout(()=>{ status.textContent=''; },3000);
  }catch(e){ status.textContent='Lỗi: '+e.message; }
  $('gen-btn').disabled=false;
}
function addImgToChat(){
  if(!genImgUrl)return;
  const d=document.createElement('div');
  d.className='msg assistant';
  d.innerHTML='<strong style="font-size:13px;color:#a78bfa">Concept — '+esc(selectedImgStyle)+'</strong><br>'+
    '<em style="font-size:13px;color:#8888b8">'+esc(genImgDesc)+'</em><br><br>'+
    '<img src="'+genImgUrl+'" style="max-width:100%;border-radius:10px;display:block;margin-bottom:10px;box-shadow:0 4px 20px rgba(0,0,0,.4)">'+
    '<a href="'+genImgUrl+'" target="_blank" download style="font-size:13px;color:#a78bfa;text-decoration:none">&#x2B07; Tai xuong anh</a>';
  $('chat-box').appendChild(d); scrollBot();
}

// ── Content Studio ──
async function generateContent(){
  const client=$('cs-client').value.trim();
  const product=$('cs-product').value.trim();
  const target=$('cs-target').value.trim();
  const keymsg=$('cs-keymsg').value.trim();
  if(!client||!product||!target||!keymsg){
    alert('Vui lòng điền đầy đủ: Client, Sản phẩm, Target và Key message');
    return;
  }
  const body={
    client, product, target, key_message: keymsg,
    objective: getSegVal('cs-obj-btns'),
    tone: getSegVal('cs-tone-btns'),
    budget: $('cs-budget').value.trim(),
    formula: getSegVal('cs-formula-btns')||'PAS',
    ai_model: getSegVal('cs-model-btns')||'claude',
  };
  const out=$('cs-output');
  out.innerHTML='';
  const card=document.createElement('div');
  card.className='output-card';
  card.innerHTML='<div class="copy-bar"><button class="copy-btn" onclick="copyOutput()">Sao chép</button></div><div id="cs-text"></div>';
  out.appendChild(card);
  const modelLabels = {'claude-sonnet-4-6':'Sonnet 4.6','claude-haiku-4-5-20251001':'Haiku 4.5','gemini-3.5-flash':'Gemini 3.5 Flash','gemini-3.1-flash-lite':'Gemini Flash-Lite','gemini-3.1-pro':'Gemini 3.1 Pro'};
  const modelLabel = modelLabels[body.ai_model]||body.ai_model;
  $('cs-gen-btn').disabled=true;
  $('cs-status').textContent='Đang viết ['+modelLabel+' · '+(body.formula||'PAS')+']...';
  let full='';
  try{
    const res=await fetch('/content/generate',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify(body)});
    const reader=res.body.getReader(); const dec=new TextDecoder();
    while(true){
      const{done,value}=await reader.read(); if(done)break;
      for(const line of dec.decode(value).split('\\n')){
        if(!line.startsWith('data: '))continue;
        const d=line.slice(6); if(d==='[DONE]')break;
        try{ full+=JSON.parse(d).text; $('cs-text').innerHTML=md(full); out.scrollTop=out.scrollHeight; }catch{}
      }
    }
    csCurrentText = full;
    $('cs-status').textContent='Hoàn thành!';
    setTimeout(()=>{ $('cs-status').textContent=''; },3000);
    if(full.length>50){ const b=$('cs-save-bar'); b.style.display='flex'; $('cs-save-status').textContent=''; }
  }catch(e){ $('cs-status').textContent='Lỗi: '+e.message; }
  $('cs-gen-btn').disabled=false;
}
function copyOutput(){
  const txt=$('cs-text');
  if(txt){ navigator.clipboard.writeText(txt.innerText).then(()=>alert('Đã sao chép!')); }
}

// ── Video Script ──
async function generateScript(){
  const product=$('sc-product').value.trim();
  const target=$('sc-target').value.trim();
  if(!product||!target){ alert('Vui lòng điền Sản phẩm và Target'); return; }
  const scModel = getSegVal('sc-model-btns')||'claude-sonnet-4-6';
  const scModelLabels = {'claude-sonnet-4-6':'Sonnet 4.6','claude-haiku-4-5-20251001':'Haiku 4.5','gemini-3.5-flash':'Gemini 3.5 Flash','gemini-3.1-flash-lite':'Gemini Flash-Lite','gemini-3.1-pro':'Gemini 3.1 Pro'};
  const body={
    product, target,
    platform: getSegVal('sc-platform-btns'),
    duration: parseInt(getSegVal('sc-dur-btns')||'30'),
    hook_style: getSegVal('sc-hook-btns'),
    extra: $('sc-extra').value.trim(),
    ai_model: scModel,
  };
  const out=$('sc-output');
  out.innerHTML='';
  const card=document.createElement('div');
  card.className='output-card';
  card.innerHTML='<div class="copy-bar"><button class="copy-btn" onclick="copyScript()">Sao chép</button></div><div id="sc-text"></div>';
  out.appendChild(card);
  $('sc-gen-btn').disabled=true;
  $('sc-status').textContent='Đang viết ['+(scModelLabels[scModel]||scModel)+']...';
  let full='';
  try{
    const res=await fetch('/script/generate',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify(body)});
    const reader=res.body.getReader(); const dec=new TextDecoder();
    while(true){
      const{done,value}=await reader.read(); if(done)break;
      for(const line of dec.decode(value).split('\\n')){
        if(!line.startsWith('data: '))continue;
        const d=line.slice(6); if(d==='[DONE]')break;
        try{ full+=JSON.parse(d).text; $('sc-text').innerHTML=md(full); out.scrollTop=out.scrollHeight; }catch{}
      }
    }
    scCurrentText = full;
    $('sc-status').textContent='Hoàn thành!';
    setTimeout(()=>{ $('sc-status').textContent=''; },3000);
    if(full.length>50){ const b=$('sc-save-bar'); b.style.display='flex'; $('sc-save-status').textContent=''; }
  }catch(e){ $('sc-status').textContent='Lỗi: '+e.message; }
  $('sc-gen-btn').disabled=false;
}
function copyScript(){
  const txt=$('sc-text');
  if(txt){ navigator.clipboard.writeText(txt.innerText).then(()=>alert('Đã sao chép!')); }
}

// ── Facebook Ads ──
function setPeriod(btn){
  document.querySelectorAll('#period-btns button').forEach(b=>b.classList.remove('active'));
  btn.classList.add('active');
  adsPeriod=btn.dataset.p;
  $('custom-dates').style.display=adsPeriod==='custom'?'flex':'none';
}

async function loadAdAccounts(){
  const sel=$('ads-account-sel');
  try{
    const accounts=await fetch('/ads/accounts').then(r=>{
      if(!r.ok)throw new Error('Loi tai tai khoan');
      return r.json();
    });
    const defaultId='act_'+($('ads-account-sel').dataset.default||'');
    sel.innerHTML='<option value="">-- Chon tai khoan --</option>';
    accounts.forEach(acc=>{
      const opt=document.createElement('option');
      opt.value=acc.id;
      const status=acc.account_status===1?'Active':'Disabled';
      opt.textContent=acc.name+' | ID: '+acc.account_id+' ('+status+')';
      sel.appendChild(opt);
    });
    // Auto-select if only 1, or match default
    if(accounts.length===1){ sel.value=accounts[0].id; }
    else{
      const defId=accounts[0]?.id;
      if(defId) sel.value=defId;
    }
  }catch(e){
    sel.innerHTML='<option value="">Loi: '+esc(e.message)+'</option>';
  }
}

async function loadAdsData(){
  const accountId=$('ads-account-sel').value;
  if(!accountId){ alert('Vui lòng chọn tài khoản quảng cáo'); return; }

  if(adsPeriod==='custom'){
    const from=$('ads-from').value, to=$('ads-to').value;
    if(!from||!to){ alert('Vui lòng chọn ngày bắt đầu và kết thúc'); return; }
  }

  const btn=$('refresh-ads-btn');
  btn.disabled=true; btn.textContent='Đang tải...';
  const body=$('ads-body');
  body.innerHTML='<div class="no-data">Đang lấy dữ liệu từ Facebook...</div>';

  try{
    let url='/ads/data?period='+adsPeriod+'&account_id='+encodeURIComponent(accountId);
    if(adsPeriod==='custom'){
      url+='&date_from='+$('ads-from').value+'&date_to='+$('ads-to').value;
    }
    const res=await fetch(url);
    if(!res.ok){ const err=await res.json(); throw new Error(err.detail||'Lỗi server'); }
    adsRawData=await res.json();
    renderAdsData(adsRawData);
    analyzeAds(adsRawData);
  }catch(e){
    body.innerHTML='<div class="no-data">Lỗi: '+esc(e.message)+'</div>';
  }
  btn.disabled=false; btn.textContent='Tải dữ liệu';
}

function renderAdsData(data){
  const insights=(data.insights&&data.insights.data)||[];
  const camps=(data.campaigns&&data.campaigns.data)||[];
  const body=$('ads-body');
  body.innerHTML='';

  // KPI summary — separate messages vs leads
  let totalSpend=0,totalReach=0,totalImp=0,totalMessages=0,totalLeads=0;
  insights.forEach(r=>{
    totalSpend+=parseFloat(r.spend||0);
    totalReach+=parseInt(r.reach||0);
    totalImp+=parseInt(r.impressions||0);
    (r.actions||[]).forEach(a=>{
      if(a.action_type==='onsite_conversion.messaging_conversation_started_7d') totalMessages+=parseInt(a.value||0);
      if(a.action_type==='lead') totalLeads+=parseInt(a.value||0);
    });
  });
  const avgCPM=totalImp>0?(totalSpend/totalImp*1000):0;
  const cpMessage=totalMessages>0?(totalSpend/totalMessages):0;
  const cpLead=totalLeads>0?(totalSpend/totalLeads):0;
  const kpiHtml=`<div class="kpi-row">
    <div class="kpi-card"><div class="kpi-val">${fmtMoney(totalSpend)}</div><div class="kpi-lbl">Tổng chi tiêu</div><div class="kpi-sub">VND/USD</div></div>
    <div class="kpi-card"><div class="kpi-val">${fmtNum(totalReach)}</div><div class="kpi-lbl">Reach</div><div class="kpi-sub">người tiếp cận</div></div>
    <div class="kpi-card"><div class="kpi-val">${fmtNum(totalImp)}</div><div class="kpi-lbl">Impressions</div><div class="kpi-sub">lượt hiển thị</div></div>
    <div class="kpi-card"><div class="kpi-val">${avgCPM.toFixed(2)}</div><div class="kpi-lbl">CPM TB</div><div class="kpi-sub">chi phí/1000 hiển thị</div></div>
    <div class="kpi-card"><div class="kpi-val">${totalMessages}</div><div class="kpi-lbl">Số tin nhắn</div><div class="kpi-sub">conversations 7d</div></div>
    <div class="kpi-card"><div class="kpi-val">${cpMessage>0?fmtMoney(cpMessage):'—'}</div><div class="kpi-lbl">Chi phí/Tin nhắn</div><div class="kpi-sub">VND/USD</div></div>
    <div class="kpi-card"><div class="kpi-val">${totalLeads}</div><div class="kpi-lbl">Số KHTN</div><div class="kpi-sub">khách hàng tiềm năng</div></div>
    <div class="kpi-card"><div class="kpi-val">${cpLead>0?fmtMoney(cpLead):'—'}</div><div class="kpi-lbl">Chi phí/KHTN</div><div class="kpi-sub">VND/USD</div></div>
  </div>`;

  // Campaign table
  let rows='';
  if(insights.length){
    insights.forEach(r=>{
      const campInfo=camps.find(c=>c.name===r.campaign_name)||{};
      const status=campInfo.status||'UNKNOWN';
      const badge=status==='ACTIVE'?'<span class="badge active">Active</span>':'<span class="badge paused">Paused</span>';
      let rowMsgs=0, rowLeads=0;
      (r.actions||[]).forEach(a=>{
        if(a.action_type==='onsite_conversion.messaging_conversation_started_7d') rowMsgs+=parseInt(a.value||0);
        if(a.action_type==='lead') rowLeads+=parseInt(a.value||0);
      });
      rows+=`<tr>
        <td>${esc(r.campaign_name||'—')}</td>
        <td>${badge}</td>
        <td>${fmtMoney(r.spend)}</td>
        <td>${fmtNum(r.reach)}</td>
        <td>${parseFloat(r.cpm||0).toFixed(2)}</td>
        <td>${parseFloat(r.cpc||0).toFixed(2)}</td>
        <td>${fmtNum(r.impressions)}</td>
        <td>${rowMsgs||'—'}</td>
        <td>${rowLeads||'—'}</td>
      </tr>`;
    });
  } else {
    rows='<tr><td colspan="9" style="text-align:center;color:var(--text-2);padding:20px">Không có dữ liệu cho chu kỳ này</td></tr>';
  }

  const tableHtml=`<div class="camp-table"><table>
    <thead><tr><th>Chiến dịch</th><th>Trạng thái</th><th>Chi tiêu</th><th>Reach</th><th>CPM</th><th>CPC</th><th>Impressions</th><th>Tin nhắn</th><th>KHTN</th></tr></thead>
    <tbody>${rows}</tbody>
  </table></div>`;

  const aiHtml=`<div class="ai-section">
    <h3>&#x1F916; Phân tích AI <span id="ai-status" style="font-size:12px;color:var(--text-2);font-weight:400"></span></h3>
    <div class="seg-btns" id="ads-model-btns" style="margin-bottom:12px">
      <button class="seg active" data-v="gemini-3.5-flash">3.5 Flash</button>
      <button class="seg" data-v="gemini-3.1-flash-lite">Flash-Lite</button>
      <button class="seg" data-v="gemini-3.1-pro">3.1 Pro</button>
      <button class="seg" data-v="claude-sonnet-4-6">Sonnet 4.6</button>
      <button class="seg" data-v="claude-haiku-4-5-20251001">Haiku 4.5</button>
    </div>
    <div id="ai-analysis">Đang phân tích...</div>
  </div>`;

  body.innerHTML=kpiHtml+tableHtml+aiHtml;
  setupSegs('ads-model-btns');
}

async function analyzeAds(data){
  if(!data)return;
  const status=$('ai-status'); const analysis=$('ai-analysis');
  if(!status||!analysis)return;
  status.textContent='(đang phân tích...)';
  let full='';
  try{
    const aiModel=getSegVal('ads-model-btns')||'gemini-3.5-flash';
    const res=await fetch('/ads/analyze',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({data,ai_model:aiModel})});
    const reader=res.body.getReader(); const dec=new TextDecoder();
    while(true){
      const{done,value}=await reader.read(); if(done)break;
      for(const line of dec.decode(value).split('\\n')){
        if(!line.startsWith('data: '))continue;
        const d=line.slice(6); if(d==='[DONE]')break;
        try{ full+=JSON.parse(d).text; analysis.innerHTML=md(full); }catch{}
      }
    }
    status.textContent='';
  }catch(e){ analysis.textContent='Lỗi phân tích: '+e.message; status.textContent=''; }
}

function fmtMoney(v){ const n=parseFloat(v||0); return n>=1000?n.toLocaleString('vi-VN'):n.toFixed(2); }
function fmtNum(v){ return parseInt(v||0).toLocaleString('vi-VN'); }

// ── Intake Client ──
let intakeFullText = '';
let intakeData = {};

async function submitIntake(){
  const ten=$('in-ten').value.trim();
  if(!ten){ alert('Vui lòng nhập tên doanh nghiệp'); return; }
  intakeData = {
    ten_dn: ten,
    linh_vuc: $('in-linh-vuc').value,
    fanpage_url: $('in-fanpage').value,
    nguoi_lien_he: $('in-lien-he').value,
    fanpage_status: $('in-page-status').value,
    followers: $('in-followers').value||'0',
    da_chay_ads: $('in-ads-history').value,
    cpl_cu: $('in-cpl-cu').value,
    muc_tieu_chinh: $('in-muc-tieu').value,
    leads_mong_muon: $('in-leads').value,
    doanh_thu_mong_muon: $('in-doanh-thu').value,
    timeline_bat_dau: $('in-timeline').value,
    mo_ta_sp: $('in-sp').value,
    doi_tuong: $('in-doi-tuong').value,
    usp: $('in-usp').value,
    assets: $('in-assets').value,
    ngan_sach_ads: $('in-budget-ads').value,
    ngan_sach_dv: $('in-budget-dv').value,
    thoi_gian_hd: $('in-duration').value,
    doi_thu: $('in-doi-thu').value,
    ai_model: getSegVal('in-model-btns'),
  };
  intakeFullText = '';
  const out=$('intake-output-body');
  const status=$('intake-status');
  const submitBtn=$('intake-submit-btn');
  out.innerHTML='<em style="color:var(--text-2)">Đang phân tích...</em>';
  status.textContent='Đang phân tích thông tin client...';
  submitBtn.disabled=true;
  $('intake-copy-btn').disabled=true;
  $('intake-to-quote-btn').disabled=true;
  try{
    const res=await fetch('/intake/generate',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify(intakeData)});
    out.innerHTML='';
    const reader=res.body.getReader();
    const dec=new TextDecoder();
    while(true){
      const{done,value}=await reader.read(); if(done)break;
      for(const line of dec.decode(value).split('\\n')){
        if(!line.startsWith('data: '))continue;
        const d=line.slice(6); if(d==='[DONE]')break;
        try{ intakeFullText+=JSON.parse(d).text; out.innerHTML=md(intakeFullText); out.scrollTop=out.scrollHeight; }catch{}
      }
    }
    status.textContent='Brief hoàn thành — '+ten;
    $('intake-copy-btn').disabled=false;
    $('intake-to-quote-btn').disabled=false;
    $('intake-save-btn').disabled=false;
  }catch(e){ status.textContent='Lỗi: '+e.message; }
  submitBtn.disabled=false;
}

function copyIntake(){
  navigator.clipboard.writeText(intakeFullText).then(()=>{
    const btn=$('intake-copy-btn');
    btn.textContent='Đã copy!';
    setTimeout(()=>btn.textContent='Copy Brief',1500);
  });
}

function intakeToQuote(){
  // Pre-fill quote form tu intake data
  $('q-client').value = intakeData.ten_dn||'';
  $('q-niche').value = 'Nội thất / Thiết kế & Thi công';
  if(intakeData.muc_tieu_chinh.includes('leads')) $('q-obj').value='Tin nhắn (Lead) là chính';
  else if(intakeData.muc_tieu_chinh.includes('brand')) $('q-obj').value='Tương tác + Brand Awareness';
  else $('q-obj').value='Tin nhắn + Tương tác';
  if(intakeData.ngan_sach_ads) $('q-budget').value=intakeData.ngan_sach_ads;
  const durMap={'1 thang thu viec':'1 tháng (thử)','3 thang':'3 tháng','6 thang':'6 tháng','12 thang':'12 tháng'};
  $('q-duration').value=durMap[intakeData.thoi_gian_hd]||'3 tháng';
  $('q-note').value='Brief từ Intake:\\n- Đối tượng: '+intakeData.doi_tuong+'\\n- USP: '+intakeData.usp+'\\n- Assets: '+intakeData.assets;
  showTab('quote');
}

// ── Quote Generator ──
let quoteFullText = '';
let quoteFormData = {};

async function submitQuote(){
  const client=$('q-client').value.trim();
  if(!client){ alert('Vui lòng nhập tên doanh nghiệp khách hàng'); return; }
  const body={
    ten_client: client,
    linh_vuc: $('q-niche').value,
    muc_tieu: $('q-obj').value,
    ngan_sach_ads: $('q-budget').value ? parseInt($('q-budget').value).toLocaleString('vi-VN') : '',
    thoi_gian: $('q-duration').value,
    dich_vu: $('q-services').value,
    ghi_chu: $('q-note').value,
    ai_model: getSegVal('q-model-btns'),
  };
  quoteFormData = body;
  quoteFullText = '';
  const out=$('quote-output-body');
  const status=$('q-status');
  const submitBtn=$('q-submit-btn');
  const copyBtn=$('q-copy-btn');
  const exportBtn=$('q-export-btn');
  out.innerHTML='';
  status.textContent='Đang tạo báo giá...';
  submitBtn.disabled=true;
  copyBtn.disabled=true;
  exportBtn.disabled=true;
  try{
    const res=await fetch('/quote/generate',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify(body)});
    const reader=res.body.getReader();
    const dec=new TextDecoder();
    while(true){
      const{done,value}=await reader.read(); if(done)break;
      for(const line of dec.decode(value).split('\\n')){
        if(!line.startsWith('data: '))continue;
        const d=line.slice(6); if(d==='[DONE]')break;
        try{ quoteFullText+=JSON.parse(d).text; out.innerHTML=md(quoteFullText); out.scrollTop=out.scrollHeight; }catch{}
      }
    }
    status.textContent='Báo giá hoàn thành — '+client;
    copyBtn.disabled=false;
    exportBtn.disabled=false;
    $('q-save-btn').disabled=false;
  }catch(e){ status.textContent='Lỗi: '+e.message; }
  submitBtn.disabled=false;
}

function copyQuote(){
  navigator.clipboard.writeText(quoteFullText).then(()=>{
    const btn=$('q-copy-btn');
    btn.textContent='Đã copy!';
    setTimeout(()=>btn.textContent='Copy',1500);
  });
}

async function exportQuoteExcel(){
  const btn=$('q-export-btn');
  btn.disabled=true; btn.textContent='Đang xuất...';
  try{
    const res=await fetch('/quote/export',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify(quoteFormData)});
    if(!res.ok) throw new Error('Loi server');
    const blob=await res.blob();
    const cd=res.headers.get('Content-Disposition')||'';
    const match=cd.match(/filename\\*=UTF-8''(.+)/);
    const fname=match?decodeURIComponent(match[1]):'BaoGia.xlsx';
    const url=URL.createObjectURL(blob);
    const a=document.createElement('a'); a.href=url; a.download=fname; a.click();
    URL.revokeObjectURL(url);
  }catch(e){ alert('Lỗi xuất Excel: '+e.message); }
  btn.disabled=false; btn.textContent='Xuất Excel';
}

// ── Content Library ──
let csCurrentText = '';
let scCurrentText = '';
let libItemsMap = {};
const LIB_TYPE_COLORS = {content:'#7c3aed',script:'#2563eb',brief:'#d97706',quote:'#16a34a'};
const LIB_TYPE_LABELS = {content:'Bài viết',script:'Kịch bản',brief:'Brief',quote:'Báo giá'};

async function loadLibrary(){
  const type = getSegVal('lib-type-btns')||'';
  const search = ($('lib-search')||{}).value||'';
  try{
    const res = await fetch('/library?type='+encodeURIComponent(type)+'&search='+encodeURIComponent(search));
    const data = await res.json();
    libItemsMap = {};
    data.items.forEach(it=>libItemsMap[it.id]=it);
    renderLibraryList(data.items);
  }catch(e){ $('lib-list').innerHTML='<div class="no-data">Lỗi tải: '+e.message+'</div>'; }
}

function renderLibraryList(items){
  const el = $('lib-list');
  const cnt = $('lib-count');
  if(cnt) cnt.textContent = items.length ? items.length+' mục' : '';
  if(!items.length){
    el.innerHTML='<div class="no-data">Chưa có nội dung nào — Tạo content rồi nhấn 💾 Lưu vào thư viện</div>';
    return;
  }
  el.innerHTML = items.map(item=>{
    const c = LIB_TYPE_COLORS[item.type]||'#888';
    const lbl = item.type_label||item.type;
    const preview = item.content.substring(0,280).replace(/</g,'&lt;');
    const hasMore = item.content.length > 280;
    return `<div class="lib-item" id="lib-${item.id}">
      <div class="lib-item-header">
        <span class="lib-badge" style="background:${c}22;color:${c};border:1px solid ${c}44">${lbl}</span>
        <span class="lib-item-title" title="${item.title.replace(/"/g,'&quot;')}">${item.title}</span>
        ${item.client?`<span class="lib-client">📌 ${item.client}</span>`:''}
        <span class="lib-date">${item.created_at.substring(0,16)}</span>
        <button class="lib-del" onclick="deleteLibItem('${item.id}')" title="Xóa">✕</button>
      </div>
      <div class="lib-preview" id="libprev-${item.id}">${preview.replace(/\\n/g,'<br>')}${hasMore?'...':''}</div>
      <div class="lib-footer">
        ${hasMore?`<button class="lib-foot-btn" id="libexp-${item.id}" onclick="toggleLibExpand('${item.id}')">Xem đầy đủ ▾</button>`:''}
        <button class="lib-foot-btn" onclick="copyLibItem('${item.id}')">📋 Copy</button>
      </div>
    </div>`;
  }).join('');
}

function toggleLibExpand(id){
  const prev = $('libprev-'+id);
  const btn = $('libexp-'+id);
  if(!prev||!btn) return;
  const expanded = prev.classList.toggle('expanded');
  if(expanded){
    const item = libItemsMap[id];
    if(item) prev.innerHTML = item.content.replace(/</g,'&lt;').replace(/\\n/g,'<br>');
    btn.textContent = 'Thu gọn ▴';
  } else {
    const item = libItemsMap[id];
    if(item) prev.innerHTML = item.content.substring(0,280).replace(/</g,'&lt;').replace(/\\n/g,'<br>')+(item.content.length>280?'...':'');
    btn.textContent = 'Xem đầy đủ ▾';
  }
}

async function copyLibItem(id){
  const item = libItemsMap[id];
  if(!item) return;
  await navigator.clipboard.writeText(item.content);
  alert('Đã copy nội dung!');
}

async function deleteLibItem(id){
  if(!confirm('Xóa mục này khỏi thư viện?')) return;
  await fetch('/library/'+id,{method:'DELETE'});
  loadLibrary();
}

async function saveToLibrary(type, title, content, client){
  if(!content||content.length<10){ alert('Không có nội dung để lưu'); return; }
  const res = await fetch('/library/save',{method:'POST',headers:{'Content-Type':'application/json'},
    body:JSON.stringify({type,title,content,client:client||''})});
  const data = await res.json();
  return data.id;
}

async function saveCSToLib(){
  const text = csCurrentText || ($('cs-text')?$('cs-text').innerText:'');
  if(!text){ alert('Chưa có nội dung để lưu'); return; }
  const product = ($('cs-product')||{}).value||'Content';
  const formula = getSegVal('cs-formula-btns')||'PAS';
  const client = ($('cs-client')||{}).value||'';
  const title = `${product} — ${formula} (${new Date().toLocaleDateString('vi-VN')})`;
  const btn = document.querySelector('#cs-save-bar .save-lib-btn');
  const status = $('cs-save-status');
  if(btn) btn.disabled=true;
  const id = await saveToLibrary('content', title, text, client);
  if(status) status.textContent = id ? '✓ Đã lưu!' : '✗ Lỗi';
  if(btn){ btn.disabled=false; setTimeout(()=>{ if(status) status.textContent=''; },3000); }
}

async function saveSCToLib(){
  const text = scCurrentText || ($('sc-text')?$('sc-text').innerText:'');
  if(!text){ alert('Chưa có nội dung để lưu'); return; }
  const product = ($('sc-product')||{}).value||'Video';
  const platform = getSegVal('sc-platform-btns')||'TikTok';
  const dur = getSegVal('sc-dur-btns')||'30';
  const title = `${product} — ${platform} ${dur}s (${new Date().toLocaleDateString('vi-VN')})`;
  const btn = document.querySelector('#sc-save-bar .save-lib-btn');
  const status = $('sc-save-status');
  if(btn) btn.disabled=true;
  const id = await saveToLibrary('script', title, text, '');
  if(status) status.textContent = id ? '✓ Đã lưu!' : '✗ Lỗi';
  if(btn){ btn.disabled=false; setTimeout(()=>{ if(status) status.textContent=''; },3000); }
}

async function saveIntakeToLib(){
  const text = intakeFullText;
  if(!text){ alert('Chưa có brief để lưu'); return; }
  const name = intakeData?.ten_dn||'Client';
  const title = `${name} — Brief chiến lược (${new Date().toLocaleDateString('vi-VN')})`;
  const btn = $('intake-save-btn');
  if(btn) btn.disabled=true;
  const id = await saveToLibrary('brief', title, text, name);
  if(btn){ btn.textContent = id?'✓ Đã lưu':'✗ Lỗi'; btn.disabled=false;
    setTimeout(()=>{ btn.textContent='💾 Lưu'; },3000); }
}

async function saveQuoteToLib(){
  const text = quoteFullText;
  if(!text){ alert('Chưa có báo giá để lưu'); return; }
  const client = quoteFormData?.ten_client||'Client';
  const title = `${client} — Báo giá (${new Date().toLocaleDateString('vi-VN')})`;
  const btn = $('q-save-btn');
  if(btn) btn.disabled=true;
  const id = await saveToLibrary('quote', title, text, client);
  if(btn){ btn.textContent = id?'✓ Đã lưu':'✗ Lỗi'; btn.disabled=false;
    setTimeout(()=>{ btn.textContent='💾 Lưu'; },3000); }
}

async function exportLibrary(fmt){
  const type = getSegVal('lib-type-btns')||'';
  const search = ($('lib-search')||{}).value||'';
  const url = `/library/export/${fmt}?type=${encodeURIComponent(type)}&search=${encodeURIComponent(search)}`;
  try{
    const res = await fetch(url);
    if(!res.ok){ const t=await res.text(); alert('Lỗi: '+t); return; }
    const blob = await res.blob();
    const cd = res.headers.get('Content-Disposition')||'';
    const match = cd.match(/filename\\*=UTF-8''(.+)/);
    const fname = match ? decodeURIComponent(match[1]) : 'ThuVien.'+fmt;
    const a = document.createElement('a'); a.href=URL.createObjectURL(blob); a.download=fname; a.click();
  }catch(e){ alert('Lỗi xuất: '+e.message); }
}

// ── REPORT ──
let _rptFbData = null;
let _rptAiText  = '';
let _rptMeta    = {};

async function loadReportAccounts(){
  const sel = $('rpt-account-sel');
  if(sel.options.length > 1) return;
  try{
    const accounts = await fetch('/ads/accounts').then(r=>{ if(!r.ok)throw new Error('Lỗi'); return r.json(); });
    sel.innerHTML = '<option value="">-- Chọn tài khoản --</option>';
    accounts.forEach(acc=>{
      const o = document.createElement('option');
      o.value = acc.id;
      const status = acc.account_status===1 ? 'Active' : 'Disabled';
      o.textContent = acc.name + ' | ID: ' + acc.account_id + ' (' + status + ')';
      sel.appendChild(o);
    });
    if(sel.options.length > 1) sel.selectedIndex = 1;
  }catch(e){ sel.innerHTML='<option value="">Lỗi load tài khoản</option>'; }
}

function toggleRptCustomDates(){
  const val = $('rpt-period').value;
  const cd = $('rpt-custom-dates');
  if(val==='custom') cd.classList.add('show'); else cd.classList.remove('show');
}

async function generateReport(){
  const accountId = $('rpt-account-sel').value;
  if(!accountId){ alert('Vui lòng chọn tài khoản quảng cáo'); return; }
  const period   = $('rpt-period').value;
  const dateFrom = period==='custom' ? $('rpt-date-from').value : '';
  const dateTo   = period==='custom' ? $('rpt-date-to').value   : '';
  const clientName = ($('rpt-client').value||'').trim() || 'Khách hàng';
  const aiModel  = getSegVal('rpt-model-btns') || 'gemini-2.5-flash';

  const btn = $('rpt-gen-btn');
  btn.disabled = true;
  $('rpt-status').textContent = 'Đang lấy dữ liệu Facebook Ads...';
  $('rpt-export-row').style.display = 'none';
  $('rpt-preview').innerHTML = '<div class="rpt-placeholder">Đang tải dữ liệu...</div>';
  _rptFbData = null; _rptAiText = ''; _rptMeta = {};

  try{
    const r1 = await fetch('/report/data',{method:'POST',headers:{'Content-Type':'application/json'},
      body:JSON.stringify({account_id:accountId, period, date_from:dateFrom, date_to:dateTo})});
    if(!r1.ok){ const t=await r1.text(); throw new Error(t); }
    const d = await r1.json();
    _rptFbData  = d.fb_data;
    const kpis  = d.kpis;
    const camps = d.campaigns;
    _rptMeta    = { client_name:clientName, period_label:d.period_label||period };

    renderReportPreview(_rptMeta, kpis, camps, '');

    $('rpt-status').textContent = 'Đang phân tích bằng AI...';

    const periodLabel = d.period_label || period;
    const r2 = await fetch('/report/analyze',{method:'POST',headers:{'Content-Type':'application/json'},
      body:JSON.stringify({fb_data:_rptFbData, client_name:clientName, period_label:periodLabel, ai_model:aiModel})});
    if(!r2.ok){ const t=await r2.text(); throw new Error(t); }

    const reader = r2.body.getReader(); const decoder = new TextDecoder();
    let aiAccum = '';
    while(true){
      const {done,value} = await reader.read();
      if(done) break;
      const chunk = decoder.decode(value,{stream:true});
      chunk.split('\\n').forEach(line=>{
        if(!line.startsWith('data:')) return;
        const raw = line.slice(5).trim();
        if(raw==='[DONE]') return;
        try{ const obj=JSON.parse(raw); if(obj.t) aiAccum+=obj.t; }catch{}
      });
      const aiEl = $('rpt-ai-text');
      if(aiEl) aiEl.textContent = aiAccum;
    }
    _rptAiText = aiAccum;
    renderReportPreview(_rptMeta, kpis, camps, _rptAiText);
    $('rpt-status').textContent = 'Hoàn thành! Sẵn sàng xuất file.';
    $('rpt-export-row').style.display = 'flex';
  }catch(e){
    $('rpt-status').textContent = 'Lỗi: '+e.message;
    $('rpt-preview').innerHTML = '<div class="rpt-placeholder" style="color:#f87171">Lỗi: '+e.message+'</div>';
  }finally{
    btn.disabled = false;
  }
}

function _fmt(n){ if(!n && n!==0) return '--'; if(n>=1000000) return (n/1000000).toFixed(2)+' triệu'; if(n>=1000) return (n/1000).toFixed(1)+' nghìn'; return String(n); }
function _fmtCur(n){ if(!n && n!==0) return '--'; return new Intl.NumberFormat('vi-VN').format(Math.round(n))+'đ'; }

function renderReportPreview(meta, kpis, camps, aiText){
  const pv = $('rpt-preview');
  if(!kpis){ pv.innerHTML='<div class="rpt-placeholder">Không có dữ liệu</div>'; return; }

  const kpiHtml = `
    <div class="rpt-kpi-grid">
      <div class="rpt-kpi-card"><div class="rpt-kpi-label">Chi tiêu</div><div class="rpt-kpi-value">${_fmtCur(kpis.spend)}</div></div>
      <div class="rpt-kpi-card"><div class="rpt-kpi-label">Tiếp cận</div><div class="rpt-kpi-value">${_fmt(kpis.reach)}</div></div>
      <div class="rpt-kpi-card"><div class="rpt-kpi-label">CPM</div><div class="rpt-kpi-value">${_fmtCur(kpis.cpm)}</div></div>
      <div class="rpt-kpi-card"><div class="rpt-kpi-label">Tin nhắn</div><div class="rpt-kpi-value">${_fmt(kpis.messages)}</div><div class="rpt-kpi-sub">CPMess: ${_fmtCur(kpis.cpm_msg)}</div></div>
      <div class="rpt-kpi-card"><div class="rpt-kpi-label">KHTN</div><div class="rpt-kpi-value">${_fmt(kpis.leads)}</div><div class="rpt-kpi-sub">CPL: ${_fmtCur(kpis.cpl)}</div></div>
      <div class="rpt-kpi-card"><div class="rpt-kpi-label">Lượt hiển thị</div><div class="rpt-kpi-value">${_fmt(kpis.impressions)}</div></div>
    </div>`;

  let campRows = '';
  (camps||[]).forEach(c=>{
    campRows += `<tr>
      <td>${c.name||'--'}</td>
      <td>${_fmtCur(c.spend)}</td>
      <td>${_fmt(c.reach)}</td>
      <td>${_fmtCur(c.cpm)}</td>
      <td>${_fmt(c.messages)}</td>
      <td>${_fmtCur(c.cpm_msg)}</td>
      <td>${_fmt(c.leads)}</td>
    </tr>`;
  });

  const campHtml = `
    <table class="rpt-camp-table">
      <thead><tr>
        <th>Chiến dịch</th><th>Chi tiêu</th><th>Tiếp cận</th><th>CPM</th>
        <th>Tin nhắn</th><th>CPMess</th><th>KHTN</th>
      </tr></thead>
      <tbody>${campRows||'<tr><td colspan="7" style="text-align:center;color:var(--text-3)">Không có dữ liệu chiến dịch</td></tr>'}</tbody>
    </table>`;

  pv.innerHTML = `
    <div class="rpt-section">
      <h3>📋 ${meta.client_name||'Khách hàng'} — ${meta.period_label||''}</h3>
      ${kpiHtml}
    </div>
    <div class="rpt-section">
      <h3>📊 Chi tiết chiến dịch</h3>
      ${campHtml}
    </div>
    <div class="rpt-section">
      <h3>🤖 Phân tích AI</h3>
      <div class="rpt-ai-text" id="rpt-ai-text">${aiText||'Đang phân tích...'}</div>
    </div>`;
}

async function exportReport(fmt){
  if(!_rptFbData){ alert('Chưa có dữ liệu báo cáo — hãy tạo báo cáo trước'); return; }
  try{
    const res = await fetch('/report/export/'+fmt,{method:'POST',headers:{'Content-Type':'application/json'},
      body:JSON.stringify({fb_data:_rptFbData, ai_text:_rptAiText,
        client_name:_rptMeta.client_name||'', period_label:_rptMeta.period_label||''})});
    if(!res.ok){ const t=await res.text(); alert('Lỗi: '+t); return; }
    const blob = await res.blob();
    const cd = res.headers.get('Content-Disposition')||'';
    const match = cd.match(/filename\\*=UTF-8''(.+)/);
    const ext = fmt==='word'?'docx':fmt==='excel'?'xlsx':'pdf';
    const fname = match ? decodeURIComponent(match[1]) : 'BaoCao.'+ext;
    const a = document.createElement('a'); a.href=URL.createObjectURL(blob); a.download=fname; a.click();
  }catch(e){ alert('Lỗi xuất: '+e.message); }
}


</script>
</body>
</html>"""

if __name__ == "__main__":
    import uvicorn
    print("Server: http://localhost:8000")
    uvicorn.run(app, host="0.0.0.0", port=8000)
